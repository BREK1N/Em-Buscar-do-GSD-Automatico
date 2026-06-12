from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# --- Default font ---
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

def set_heading(paragraph, level, text):
    paragraph.style = f'Heading {level}'
    paragraph.clear()
    run = paragraph.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0x6E)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0x6E)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2E, 0x4D, 0x8A)
    run.font.name = 'Arial'
    run.font.bold = True

def add_paragraph(content='', bold_parts=None, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    # If content is a list of (text, bold) tuples, render them
    if isinstance(content, list):
        for part_text, is_bold in content:
            run = p.add_run(part_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.bold = is_bold
    else:
        run = p.add_run(content)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    return p

def add_bullet(content, bold_parts=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    if isinstance(content, list):
        for part_text, is_bold in content:
            run = p.add_run(part_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.bold = is_bold
    else:
        run = p.add_run(content)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Header shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1E406E')
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(10)
    # Column widths
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])
    doc.add_paragraph()
    return table

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E406E')
    pBdr.append(bottom)
    pPr.append(pBdr)

# =====================================================================
# TÍTULO PRINCIPAL
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run('Relatório do Dia — 12/06/2026')
run.font.name = 'Arial'
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x40, 0x6E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(16)
run2 = sub.add_run('Projeto GSD Automático — Sessão de desenvolvimento')
run2.font.name = 'Arial'
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

add_divider()
doc.add_paragraph()

# =====================================================================
# SEÇÃO 1
# =====================================================================
p = doc.add_paragraph()
set_heading(p, 1, '1. Correção de namespace ESI — NoReverseMatch')

add_paragraph([('Problema: ', True),
               ('Ao acessar ', False),
               ('/home/', True),
               (', o Django lançava ', False),
               ("NoReverseMatch: 'ESI' is not a registered namespace", True),
               ('. O app ', False),
               ('ESI', True),
               (' existia completo (models, views, urls, templates, migrations) mas nunca tinha sido incluído no arquivo de URLs principal do projeto.', False)])

add_paragraph([('Solução: ', True),
               ('Uma linha adicionada em ', False),
               ('GsdAutomatico/urls.py', True),
               (':', False)])

add_code("path('esi/', include('ESI.urls', namespace='ESI')),")

add_divider()
doc.add_paragraph()

# =====================================================================
# SEÇÃO 2
# =====================================================================
p = doc.add_paragraph()
set_heading(p, 1, '2. Migração do Relsisam — de SCIM para Secao_operacoes')

add_paragraph([('Contexto: ', True),
               ('O Relsisam (painel de disponibilidade diária) tinha sido criado por engano dentro do app SCIM. Foi solicitada a migração para Secao_operacoes, que é onde faz sentido semanticamente.', False)])

doc.add_paragraph()
p2 = doc.add_paragraph()
set_heading(p2, 2, 'O que foi movido')

add_table(
    ['Ação', 'Arquivo'],
    [
        ['Removida RelsisamView', 'SCIM/views.py'],
        ['Removida rota relsisam/', 'SCIM/urls.py'],
        ['Removido link do sidebar', 'SCIM/templates/SCIM/base_scim.html'],
        ['Adicionada RelsisamView', 'Secao_operacoes/views.py'],
        ['Adicionada rota relsisam/', 'Secao_operacoes/urls.py'],
        ['Criado template', 'Secao_operacoes/templates/Secao_operacoes/relsisam.html'],
        ['Adicionado link no sidebar', 'Secao_operacoes/templates/Secao_operacoes/base.html'],
    ],
    col_widths=[6, 10]
)

add_paragraph([('Resultado: ', True),
               ('A URL migrou de ', False),
               ('/scim/relsisam/', True),
               (' para ', False),
               ('/secao_operacoes/relsisam/', True),
               ('. Os modelos TipoCurso e CursoEfetivo permaneceram na SCIM (apenas leitura cruzada). Nenhuma migração de banco de dados foi necessária.', False)])

add_divider()
doc.add_paragraph()

# =====================================================================
# SEÇÃO 3
# =====================================================================
p = doc.add_paragraph()
set_heading(p, 1, '3. Redesenho completo do Relsisam')

add_paragraph([('Contexto: ', True),
               ('O Relsisam original mostrava disponibilidade por tipo de curso. A necessidade real é um painel de controle de efetivo diário por setor, com situação em tempo real de cada militar, alimentado 100% automaticamente pelas escalas e missões já existentes no sistema.', False)])

doc.add_paragraph()
p3a = doc.add_paragraph()
set_heading(p3a, 2, '3.1 Fontes de dados (100% automático)')

add_table(
    ['Situação', 'Origem / Modelo'],
    [
        ['De missão (SOP)', 'Missao — cmt_missao, motorista, equipe onde data_missao = hoje'],
        ['De missão (ESI)', 'EscalaMissaoESI via missao__data_missao = hoje'],
        ['De serviço', 'TurnoEscala onde data = hoje'],
        ['Licença / Afastamento / TDO / Outro', 'Novo modelo SituacaoEspecialEfetivo'],
        ['Livre', 'Calculado: total do setor − todas as situações acima'],
    ],
    col_widths=[6, 10]
)

doc.add_paragraph()
p3b = doc.add_paragraph()
set_heading(p3b, 2, '3.2 Novo modelo — SituacaoEspecialEfetivo')

add_paragraph('Adicionado em Secao_operacoes/models.py com os seguintes campos e características:')

add_bullet([('Tipos disponíveis: ', True),
            ('Licença/Férias, Afastamento Médico, TDO/Curso, Outro', False)],
           bold_parts=True)
add_bullet([('Campos: ', True),
            ('efetivo, tipo, data_inicio, data_fim (nullable = situação em aberto), observacao, registrado_por', False)],
           bold_parts=True)
add_bullet([('Migração: ', True),
            ('0021_situacaoespecialefetivo — gerada e aplicada via Docker', False)],
           bold_parts=True)
add_bullet([('Admin: ', True),
            ('Registrada em Secao_operacoes/admin.py com filtros por tipo, busca por nome e hierarquia de datas', False)],
           bold_parts=True)

doc.add_paragraph()
p3c = doc.add_paragraph()
set_heading(p3c, 2, '3.3 Relsisam redesenhado — /secao_operacoes/relsisam/')

add_bullet('KPIs globais no topo: efetivo total, livres, de serviço, de missão, situação especial e missões do dia')
add_bullet('Um card por setor com contadores coloridos: livres (verde), de serviço (azul), de missão (laranja), outros (cinza)')
add_bullet('Cada card possui seções expansíveis com a lista de nomes dos militares por situação')
add_bullet('Auto-refresh automático a cada 60 segundos')
add_bullet('Botão "Lançar situação" com link direto para o formulário de situações especiais')

doc.add_paragraph()
p3d = doc.add_paragraph()
set_heading(p3d, 2, '3.4 CRUD de situações especiais — /secao_operacoes/situacoes-especiais/')

add_paragraph('Quatro views novas + quatro rotas + dois templates criados:')

add_table(
    ['Rota', 'Função'],
    [
        ['situacoes-especiais/', 'Lista de situações ativas + histórico recente (últimas 30)'],
        ['situacoes-especiais/nova/', 'Formulário com autocomplete de militar (reutiliza API existente)'],
        ['situacoes-especiais/<pk>/encerrar/', 'Define data_fim = hoje, encerrando a situação'],
        ['situacoes-especiais/<pk>/excluir/', 'Remove o registro permanentemente'],
    ],
    col_widths=[7, 9]
)

add_paragraph([('Observação: ', True),
               ('O autocomplete do formulário reutiliza a API já existente em ', False),
               ('/secao_operacoes/api/efetivo/', True),
               (' — sem nenhuma nova endpoint necessária.', False)])

add_divider()
doc.add_paragraph()

# =====================================================================
# RESUMO FINAL
# =====================================================================
p = doc.add_paragraph()
set_heading(p, 1, 'Resumo Geral do Dia')

add_table(
    ['#', 'Tarefa', 'Arquivos tocados'],
    [
        ['1', 'Fix NoReverseMatch — namespace ESI', '1 arquivo'],
        ['2', 'Migração do Relsisam SCIM → Secao_operacoes', '7 arquivos'],
        ['3', 'Redesenho total do Relsisam (efetivo por setor)', '8 arquivos + 1 migração de banco'],
    ],
    col_widths=[1, 9, 6]
)

add_paragraph([('Total de arquivos modificados/criados no dia: ', True), ('16', False)], bold_parts=True)
add_paragraph([('Migrações de banco aplicadas: ', True), ('1  (0021_situacaoespecialefetivo)', False)], bold_parts=True)

# =====================================================================
# SAVE
# =====================================================================
output_path = r'C:/Users/Migue/Desktop/Relatorio_12-06-2026.docx'
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
