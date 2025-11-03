import pandas as pd
import io
import json
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.views.generic import ListView, CreateView, UpdateView, DeleteView, DetailView
from django.urls import reverse_lazy, reverse
from django.db.models import Q, Max, Case, When, Value, IntegerField, Count
from django.db import transaction
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_POST, require_GET
from .models import Militar, PATD, Configuracao, Anexo
# --- ALTERAÇÃO: Adicionar ComandanteAprovarForm ---
from .forms import MilitarForm, PATDForm, AtribuirOficialForm, AceitarAtribuicaoForm, ComandanteAprovarForm
from langchain_community.document_loaders import PyPDFLoader
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel, Field
import os
import tempfile
from dotenv import load_dotenv
import logging
from datetime import datetime, timedelta
from django.conf import settings
import locale
import docx
import re
# --- CORREÇÃO DA IMPORTAÇÃO ---
from django.contrib.staticfiles.storage import staticfiles_storage
# --- FIM DA CORREÇÃO ---
from django.utils import timezone
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.mixins import UserPassesTestMixin
# --- ALTERAÇÃO: Importar authenticate ---
from django.contrib.auth import authenticate


# --- INÍCIO DA MODIFICAÇÃO: Importar novas classes e funções ---
from .analise_transgressao import (
    enquadra_item,
    verifica_agravante_atenuante,
    sugere_punicao,
    model, # Assuming model is defined here
    analisar_e_resumir_defesa,
    reescrever_ocorrencia,
    texto_relatorio,
    AnaliseTransgressao, # Importar a classe principal
    MilitarAcusado,      # Importar a sub-classe
    analisar_documento_pdf # Importar a função de análise atualizada
)
# --- FIM DA MODIFICAÇÃO ---

from difflib import SequenceMatcher # Importado para a verificação de similaridade
from django.utils.decorators import method_decorator
from num2words import num2words # Importação para converter números em texto
from functools import wraps # Importado para criar o decorator
import threading # Importado para tarefas em background
from .permissions import has_comandante_access, has_ouvidoria_access
import base64
from django.core.files.base import ContentFile
from uuid import uuid4
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from bs4 import BeautifulSoup, NavigableString
import traceback # Importar traceback para log detalhado


# --- Funções e Mixins de Permissão ---
def comandante_redirect(view_func):
    """
    Decorator for views that checks that the user is NOT a comandante,
    redirecting to the comandante dashboard if they are.
    """
    @wraps(view_func)
    def _wrapped_view(request, *args, **kwargs):
        if has_comandante_access(request.user) and not request.user.is_superuser:
            return redirect('Ouvidoria:comandante_dashboard')
        return view_func(request, *args, **kwargs)
    return _wrapped_view

# --- NOVO DECORATOR ---
def oficial_responsavel_required(view_func):
    """
    Decorator que verifica se o usuário logado é o oficial responsável pela PATD.
    """
    @wraps(view_func)
    def _wrapped_view(request, pk, *args, **kwargs):
        patd = get_object_or_404(PATD, pk=pk)

        # Verifica se o usuário tem um perfil militar e se ele é o oficial responsável
        if (hasattr(request.user, 'profile') and
            request.user.profile.militar and
            request.user.profile.militar == patd.oficial_responsavel):
            return view_func(request, pk, *args, **kwargs)
        else:
            messages.error(request, "Acesso negado. Apenas o oficial apurador designado pode executar esta ação.")
            return redirect('Ouvidoria:patd_detail', pk=pk)
    return _wrapped_view

class OuvidoriaAccessMixin(UserPassesTestMixin):
    """Mixin para Class-Based Views para verificar a permissão de acesso à Ouvidoria."""
    def test_func(self):
        return has_ouvidoria_access(self.request.user)

class ComandanteAccessMixin(UserPassesTestMixin):
    """Mixin para Class-Based Views para verificar a permissão de acesso do Comandante."""
    def test_func(self):
        return has_comandante_access(self.request.user)

ouvidoria_required = user_passes_test(has_ouvidoria_access)
comandante_required = user_passes_test(has_comandante_access)


# Configuração de logging para depuração
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
load_dotenv()

# Define o locale para português do Brasil para formatar as datas
try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
except locale.Error:
    logger.warning("Locale pt_BR.UTF-8 não encontrado. A data pode não ser formatada corretamente.")

@login_required
@oficial_responsavel_required
@require_POST
def regenerar_ocorrencia(request, pk):
    patd = get_object_or_404(PATD, pk=pk)
    try:
        nova_ocorrencia = reescrever_ocorrencia(patd.transgressao)
        patd.ocorrencia_reescrita = nova_ocorrencia
        patd.comprovante = nova_ocorrencia  # Atualiza o comprovante também
        patd.save(update_fields=['ocorrencia_reescrita', 'comprovante'])
        return JsonResponse({'status': 'success', 'novo_texto': nova_ocorrencia})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@oficial_responsavel_required
@require_POST
def regenerar_resumo_defesa(request, pk):
    patd = get_object_or_404(PATD, pk=pk)
    if not patd.alegacao_defesa:
        return JsonResponse({'status': 'error', 'message': 'Não há texto de defesa para resumir.'}, status=400)
    try:
        novo_resumo = analisar_e_resumir_defesa(patd.alegacao_defesa)
        patd.alegacao_defesa_resumo = novo_resumo
        patd.save(update_fields=['alegacao_defesa_resumo'])
        return JsonResponse({'status': 'success', 'novo_texto': novo_resumo})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@oficial_responsavel_required
@require_POST
def regenerar_texto_relatorio(request, pk):
    patd = get_object_or_404(PATD, pk=pk)
    try:
        novo_relatorio = texto_relatorio(patd.transgressao, patd.alegacao_defesa)
        patd.texto_relatorio = novo_relatorio
        patd.save(update_fields=['texto_relatorio'])
        return JsonResponse({'status': 'success', 'novo_texto': novo_relatorio})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@oficial_responsavel_required
@require_POST
def regenerar_punicao(request, pk):
    patd = get_object_or_404(PATD, pk=pk)
    try:
        punicao_obj = sugere_punicao(
            transgressao=patd.transgressao,
            agravantes=patd.circunstancias.get('agravantes', []),
            atenuantes=patd.circunstancias.get('atenuantes', []),
            itens=patd.itens_enquadrados,
            observacao="Regeneração de punição"
        )
        nova_punicao_sugerida = punicao_obj.punicao.get('punicao', 'Erro na sugestão.')
        patd.punicao_sugerida = nova_punicao_sugerida
        patd.save(update_fields=['punicao_sugerida'])
        return JsonResponse({'status': 'success', 'novo_texto': nova_punicao_sugerida})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

def _check_and_advance_reconsideracao_status(patd_pk):
    """
    Função centralizada que verifica, dentro de uma transação para garantir a consistência dos dados,
    se a reconsideração tem conteúdo e assinatura. Se ambas as condições forem verdadeiras,
    o status é avançado.
    """
    try:
        with transaction.atomic():
            # Bloqueia a linha da PATD para evitar condições de corrida
            patd = PATD.objects.select_for_update().get(pk=patd_pk)

            has_content = bool(patd.texto_reconsideracao or patd.anexos.filter(tipo='reconsideracao').exists())
            has_signature = bool(patd.assinatura_reconsideracao)

            logger.info(f"PATD {patd.pk}: Current status: {patd.status}, Has Content? {has_content}, Has Signature? {has_signature}")

            if patd.status == 'em_reconsideracao' and has_content:
                patd.status = 'aguardando_comandante_base'
                patd.save(update_fields=['status'])
                logger.info(f"PATD {patd.pk} status advanced to 'aguardando_comandante_base'.")
            else:
                logger.info(f"PATD {patd.pk}: Conditions not met to advance status.")
    except PATD.DoesNotExist:
        logger.error(f"PATD {patd_pk} not found during status check.")
    except Exception as e:
        logger.error(f"Error in _check_and_advance_reconsideracao_status for PATD {patd_pk}: {e}")

# Removida a classe AnaliseTransgressao daqui, pois foi movida para analise_transgressao.py

def get_next_patd_number():
    """Gera o próximo número sequencial para a PATD."""
    max_num = PATD.objects.aggregate(max_num=Max('numero_patd'))['max_num']
    return (max_num or 0) + 1

# =============================================================================
# FUNÇÕES AUXILIARES MOVIMDAS PARA CIMA PARA CORRIGIR O ERRO
# =============================================================================
def format_militar_string(militar, with_spec=False):
    """
    Formata o nome do militar, colocando o nome de guerra em negrito (com **).
    Lida com casos complexos como 'D. PAULA'.
    """
    if not militar:
        return ""

    nome_completo = militar.nome_completo
    nome_guerra = militar.nome_guerra

    # Limpa o nome de guerra de pontuações e divide em partes
    guerra_parts = re.sub(r'[^\w\s]', '', nome_guerra).upper().split()

    # Divide o nome completo em palavras para manipulação
    nome_completo_words = nome_completo.split()

    # Itera sobre cada parte do nome de guerra
    for part in guerra_parts:
        if len(part) == 1:  # Se for uma inicial (ex: 'D')
            last_match_index = -1
            # Encontra o índice da ÚLTIMA palavra no nome completo que começa com essa inicial
            for i, word in enumerate(nome_completo_words):
                if word.upper().startswith(part):
                    last_match_index = i

            # Se encontrou uma correspondência, aplica o negrito apenas na inicial
            if last_match_index != -1:
                word_to_format = nome_completo_words[last_match_index]
                # Coloca em negrito apenas a primeira letra
                formatted_word = f"**{word_to_format[0]}**{word_to_format[1:]}"
                nome_completo_words[last_match_index] = formatted_word

        else:  # Se for uma palavra completa (ex: 'PAULA')
            # Procura a palavra exata no nome completo e aplica o negrito
            for i, word in enumerate(nome_completo_words):
                # Remove pontuação da palavra do nome completo para comparação
                clean_word = re.sub(r'[^\w\s]', '', word)
                if clean_word.upper() == part:
                    nome_completo_words[i] = f"**{word}**"
                    break # Para de procurar após encontrar a primeira correspondência

    formatted_name = ' '.join(nome_completo_words)
    posto = getattr(militar, 'posto', '')

    if with_spec:
        especializacao = getattr(militar, 'especializacao', '')
        return f"{posto} {especializacao} {formatted_name}".strip()
    else:
        return f"{posto} {formatted_name}".strip()

# =============================================================================
# Otimização da Geração de Documentos
# =============================================================================

def get_anexo_content_as_html(anexo):
    """
    Lê um ficheiro Anexo e retorna o seu conteúdo como uma string HTML.
    Suporta imagens, PDFs, e ficheiros DOCX.
    """
    try:
        file_path = anexo.arquivo.path
        file_url = anexo.arquivo.url
        file_name = os.path.basename(file_path)

        if not os.path.exists(file_path):
            return f"<p><strong>{file_name}</strong>: Erro - Ficheiro não encontrado no servidor.</p>"

        ext = os.path.splitext(file_name)[1].lower()

        if ext in ['.png', '.jpg', '.jpeg', '.gif']:
            with open(file_path, 'rb') as f:
                encoded_string = base64.b64encode(f.read()).decode('utf-8')
                return f'<h4>Anexo: {file_name}</h4><img src="data:image/{ext[1:]};base64,{encoded_string}" style="max-width: 100%; height: auto;" alt="{file_name}"><hr>'

        elif ext == '.pdf':
            try:
                loader = PyPDFLoader(file_path)
                pages = loader.load_and_split()
                content = "\n".join([page.page_content for page in pages])
                return f'<h4>Anexo: {file_name}</h4><pre style="white-space: pre-wrap; word-wrap: break-word;">{content}</pre><hr>'
            except Exception as e:
                logger.error(f"Erro ao ler PDF {file_name}: {e}")
                return f'<p><strong>{file_name}</strong>: Não foi possível extrair o texto do PDF. <a href="{file_url}" target="_blank">Fazer download</a></p><hr>'

        elif ext == '.docx':
            try:
                doc = docx.Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
                return f'<h4>Anexo: {file_name}</h4><pre style="white-space: pre-wrap; word-wrap: break-word;">{content}</pre><hr>'
            except Exception as e:
                logger.error(f"Erro ao ler DOCX {file_name}: {e}")
                return f'<p><strong>{file_name}</strong>: Não foi possível ler o conteúdo do DOCX. <a href="{file_url}" target="_blank">Fazer download</a></p><hr>'

        elif ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                return f'<h4>Anexo: {file_name}</h4><pre style="white-space: pre-wrap; word-wrap: break-word;">{content}</pre><hr>'

        else:
            # Para outros tipos de ficheiro, apenas fornece um link
            return f'<p><strong>Anexo: {file_name}</strong> (Tipo de ficheiro não suportado para visualização) - <a href="{file_url}" target="_blank">Fazer download</a></p><hr>'

    except Exception as e:
        logger.error(f"Erro ao processar anexo {anexo.id}: {e}")
        return f"<p>Erro ao carregar o anexo {os.path.basename(anexo.arquivo.name)}.</p>"


def _get_document_context(patd):
    """
    Função centralizada para coletar e formatar todos os dados
    necessários para qualquer documento.
    """
    config = Configuracao.load()
    comandante_gsd = config.comandante_gsd
    comandante_bagl = config.comandante_bagl
    now = timezone.now()

    # Formatações de Data
    data_inicio = patd.data_inicio
    data_patd_fmt = data_inicio.strftime('%d%m%Y')
    data_ocorrencia_fmt = patd.data_ocorrencia.strftime('%d/%m/%Y') if patd.data_ocorrencia else "[Data não informada]"
    data_oficio_fmt = patd.data_oficio.strftime('%d/%m/%Y') if patd.data_oficio else "[Data do ofício não informada]"
    data_ciencia_fmt = patd.data_ciencia.strftime('%d/%m/%Y') if patd.data_ciencia else "[Data não informada]"
    data_alegacao_fmt = patd.data_alegacao.strftime('%d/%m/%Y') if patd.data_alegacao else "[Data não informada]"

    # Formatação de Itens Enquadrados e Circunstâncias
    itens_enquadrados_str = ", ".join([str(item.get('numero', '')) for item in patd.itens_enquadrados]) if patd.itens_enquadrados else ""
    atenuantes_str = ", ".join(patd.circunstancias.get('atenuantes', [])) if patd.circunstancias else "Nenhuma"
    agravantes_str = ", ".join(patd.circunstancias.get('agravantes', [])) if patd.circunstancias else "Nenhuma"

    # --- LÓGICA DE FORMATAÇÃO DA PUNIÇÃO ATUALIZADA ---
    # Formatação da Punição Completa
    punicao_final_str = patd.punicao or "[Punição não definida]"
    if patd.dias_punicao and patd.punicao:
        # A lógica agora constrói a string completa, ex: "Seis (06) de detenção"
        punicao_final_str = f"{patd.dias_punicao} de {patd.punicao}"

    # Cálculo do Prazo Final (Deadline) para Preclusão
    deadline_str = "[Prazo não iniciado]"
    if patd.data_ciencia:
        dias_uteis_a_adicionar = config.prazo_defesa_dias
        data_final = patd.data_ciencia
        dias_adicionados = 0
        while dias_adicionados < dias_uteis_a_adicionar:
            data_final += timedelta(days=1)
            if data_final.weekday() < 5:
                dias_adicionados += 1
        deadline = data_final + timedelta(minutes=config.prazo_defesa_minutos)
        deadline_str = deadline.strftime('%d/%m/%Y às %H:%M')

    data_publicacao_fmt = patd.data_publicacao_punicao.strftime('%d/%m/%Y às %H:%M') if patd.data_publicacao_punicao else "[Data não informada]"
    data_reconsideracao_fmt = patd.data_reconsideracao.strftime('%d/%m/%Y') if patd.data_reconsideracao else "[Data não informada]"

    # Lógica para Oficial Apurador
    oficial_definido = patd.status not in ['definicao_oficial', 'aguardando_aprovacao_atribuicao']

    context = {
        # Placeholders Comuns
        # --- CORREÇÃO: Usar staticfiles_storage.url ---
        '{Brasao da Republica}': f'<img src="{staticfiles_storage.url("img/brasao.png")}" alt="Brasão da República" style="width: 100px; height: auto;">',
        # --- FIM DA CORREÇÃO ---
        '{N PATD}': str(patd.numero_patd),
        '{DataPatd}': data_patd_fmt,
        '{dia}': now.strftime('%d'),
        '{Mês}': now.strftime('%B').capitalize(),
        '{Ano}': now.strftime('%Y'),

        # Dados do Militar Arrolado
        '{Militar Arrolado}': format_militar_string(patd.militar),
        '{Saram Militar Arrolado}': str(getattr(patd.militar, 'saram', '[Não informado]')),
        '{Setor Militar Arrolado}': getattr(patd.militar, 'setor', '[Não informado]'),

        # Dados do Oficial Apurador
        '{Oficial Apurador}': format_militar_string(patd.oficial_responsavel) if oficial_definido else '[Aguardando Oficial confirmar]',
        '{Posto/Especialização Oficial Apurador}': format_militar_string(patd.oficial_responsavel, with_spec=True) if oficial_definido else "[Aguardando Oficial confirmar]",
        '{Saram Oficial Apurador}': str(getattr(patd.oficial_responsavel, 'saram', 'N/A')) if oficial_definido else "[Aguardando Oficial confirmar]",
        '{Setor Oficial Apurador}': getattr(patd.oficial_responsavel, 'setor', 'N/A') if oficial_definido else "[Aguardando Oficial confirmar]",
        '{Assinatura Oficial Apurador}': '{Assinatura_Imagem_Oficial_Apurador}' if oficial_definido and patd.assinatura_oficial else ('[Aguardando Oficial confirmar]' if not oficial_definido else '{Botao Assinar Oficial}'),

        # Dados do Comandante
        '{Comandante /Posto/Especialização}': format_militar_string(comandante_gsd, with_spec=True) if comandante_gsd else "[Comandante GSD não definido]",
        '{Comandante_bagl_botao}': format_militar_string(comandante_bagl, with_spec=True) if comandante_bagl else "[Comandante BAGL não definido]",


        # Dados da Transgressão
        '{data da Ocorrencia}': data_ocorrencia_fmt,
        '{Ocorrencia reescrita}': patd.ocorrencia_reescrita or patd.transgressao,
        '{protocolo comaer}': patd.protocolo_comaer,
        '{Oficio Transgressao}': patd.oficio_transgressao,
        '{data_oficio}': data_oficio_fmt,
        '{comprovante}': patd.comprovante or '',

        # Dados da Apuração
        '{Itens enquadrados}': itens_enquadrados_str,
        '{Atenuante}': atenuantes_str,
        '{agravantes}': agravantes_str,
        '{transgressao_afirmativa}': patd.transgressao_afirmativa or '',
        '{natureza_transgressao}': patd.natureza_transgressao or '',

        # Dados da Defesa
        '{data ciência}': data_ciencia_fmt,
        '{Data da alegação}': data_alegacao_fmt,
        '{Alegação_defesa_resumo}': patd.alegacao_defesa_resumo or '',

        # Placeholders de Punição
        '{punicao_completa}': punicao_final_str,
        '{punicao}': punicao_final_str,
        '{punição_botao}': f"{patd.nova_punicao_dias} de {patd.nova_punicao_tipo}" if patd.nova_punicao_dias else '{Botao Definir Nova Punicao}',
        '{dias_punicao}': "",
        '{comportamento}': patd.comportamento or "[Não avaliado]",
        '{data_publicacao_punicao}': data_publicacao_fmt,

        # Placeholders de Assinatura
        # --- ALTERAÇÃO: Inicialmente, o placeholder do CMD fica como texto ---
        '{Assinatura Comandante do GSD}': '[Assinatura Pendente]',
        '{Assinatura Alegacao Defesa}': '{Assinatura Alegacao Defesa}' if patd.assinatura_alegacao_defesa else '{Botao Assinar Defesa}',
        '{Assinatura Reconsideracao}': '{Assinatura Reconsideracao}' if patd.assinatura_reconsideracao else '{Botao Assinar Reconsideracao}',

        # Testemunhas
        '{Testemunha 1}': format_militar_string(patd.testemunha1) if patd.testemunha1 else '[Testemunha não definida]',
        '{Testemunha 2}': format_militar_string(patd.testemunha2) if patd.testemunha2 else '[Testemunha não definida]',
        '{Assinatura Testemunha 1}': '{Assinatura_Imagem_Testemunha_1}' if patd.assinatura_testemunha1 else ('{Botao Assinar Testemunha 1}' if patd.testemunha1 else '[Sem assinatura]'),
        '{Assinatura Testemunha 2}': '{Assinatura_Imagem_Testemunha_2}' if patd.assinatura_testemunha2 else ('{Botao Assinar Testemunha 2}' if patd.testemunha2 else '[Sem assinatura]'),

        # Específicos
        '{Data Final Prazo}': deadline_str,
        '{texto_relatorio}': patd.texto_relatorio or '',
        '{Texto_reconsideracao}': patd.texto_reconsideracao or '',
        '{Data_reconsideracao}': data_reconsideracao_fmt,
    }

    # Adiciona os dados das assinaturas AO CONTEXTO APENAS SE APLICÁVEL
    # Assinatura do Oficial Apurador
    if oficial_definido and patd.assinatura_oficial:
        context['assinatura_oficial_data'] = patd.assinatura_oficial.url # Usar URL se for FileField

    # Assinatura do Comandante - SÓ ADICIONA SE A PATD JÁ FOI APROVADA
    status_aprovados_e_posteriores = [
        'aguardando_assinatura_npd', 'periodo_reconsideracao', 'em_reconsideracao',
        'aguardando_comandante_base', 'aguardando_preenchimento_npd_reconsideracao',
        'aguardando_publicacao', 'finalizado'
    ]
    if comandante_gsd and comandante_gsd.assinatura and patd.status in status_aprovados_e_posteriores:
        context['assinatura_comandante_data'] = comandante_gsd.assinatura
        # Atualiza o placeholder no contexto para usar a imagem
        context['{Assinatura Comandante do GSD}'] = '{Assinatura_Imagem_Comandante_GSD}'

    return context

def _render_document_from_template(template_name, context):
    """
    Função genérica para renderizar um documento .docx a partir de um template,
    preservando o alinhamento e convertendo para HTML.
    AGORA SUPORTA O PLACEHOLDER {nova_pagina}
    """
    try:
        doc_path = os.path.join(settings.BASE_DIR, 'pdf', template_name)
        document = docx.Document(doc_path)

        alignment_map = {
            None: 'left',
            0: 'left',
            1: 'center',
            2: 'right',
            3: 'justify'
        }

        html_content = []

        for p in document.paragraphs:
            # --- NOVA VERIFICAÇÃO DE QUEBRA DE PÁGINA ---
            if '{nova_pagina}' in p.text:
                html_content.append('<div class="manual-page-break"></div>')
                continue # Pula para o próximo parágrafo
            # --- FIM DA VERIFICAÇÃO ---

            inline_text = p.text
            for placeholder, value in context.items():
                if placeholder in inline_text:
                    inline_text = inline_text.replace(str(placeholder), str(value))

            # **INÍCIO DA CORREÇÃO**
            # Tenta obter o alinhamento direto do parágrafo
            effective_alignment = p.paragraph_format.alignment
            # Se não houver alinhamento direto, herda do estilo
            if effective_alignment is None and p.style and p.style.paragraph_format:
                effective_alignment = p.style.paragraph_format.alignment
            # **FIM DA CORREÇÃO**

            alignment = alignment_map.get(effective_alignment, 'left')

            html_content.append(f'<p style="text-align: {alignment};">{inline_text}</p>')

        return ''.join(html_content)

    except FileNotFoundError:
        error_msg = f'<p style="color: red;">ERRO: Template "{template_name}" não encontrado.</p>'
        logger.error(error_msg)
        return error_msg
    except Exception as e:
        error_msg = f'<p style="color: red;">ERRO ao processar o template "{template_name}": {e}</p>'
        logger.error(error_msg)
        return error_msg


def get_document_pages(patd):
    """
    Gera uma LISTA de páginas de documento em HTML a partir dos templates.
    Cada item na lista representa um documento/seção separada.
    """
    doc_context = _get_document_context(patd)
    document_pages = []

    # 1. Documento Principal
    document_pages.append(_render_document_from_template('PATD_Coringa.docx', doc_context))

    # 2. Alegação de Defesa
    if patd.alegacao_defesa or patd.anexos.filter(tipo='defesa').exists():
        alegacao_context = doc_context.copy()
        alegacao_context['{Alegação de defesa}'] = patd.alegacao_defesa if patd.alegacao_defesa else "[Ver documentos anexos]"
        html_content = _render_document_from_template('PATD_Alegacao_DF.docx', alegacao_context)
        html_content += "{ANEXOS_DEFESA_PLACEHOLDER}"
        document_pages.append(html_content)

    # 3. Termo de Preclusão
    status_preclusao_e_posteriores = [
        'preclusao', 'apuracao_preclusao', 'aguardando_punicao',
        'aguardando_assinatura_npd', 'finalizado', 'aguardando_punicao_alterar',
        'analise_comandante', 'periodo_reconsideracao', 'em_reconsideracao',
        'aguardando_publicacao', 'aguardando_preenchimento_npd_reconsideracao',
        'aguardando_comandante_base'
    ]
    if not patd.alegacao_defesa and not patd.anexos.filter(tipo='defesa').exists() and patd.status in status_preclusao_e_posteriores:
        document_pages.append(_render_document_from_template('PRECLUSAO.docx', doc_context))

    # 4. Relatório de Apuração
    if patd.justificado:
        document_pages.append(_render_document_from_template('RELATORIO_JUSTIFICADO.docx', doc_context))
    elif patd.punicao_sugerida:
        document_pages.append(_render_document_from_template('RELATORIO_DELTA.docx', doc_context))

    # 5. Nota de Punição Disciplinar (NPD)
    status_npd_e_posteriores = [
        'aguardando_assinatura_npd', 'finalizado', 'periodo_reconsideracao',
        'em_reconsideracao', 'aguardando_publicacao',
        'aguardando_preenchimento_npd_reconsideracao', 'aguardando_comandante_base'
    ]
    if patd.status in status_npd_e_posteriores:
        document_pages.append(_render_document_from_template('MODELO_NPD.docx', doc_context))

    # 6. Reconsideração
    status_reconsideracao_e_posteriores = [
        'em_reconsideracao', 'aguardando_publicacao', 'finalizado',
        'aguardando_preenchimento_npd_reconsideracao', 'aguardando_comandante_base'
    ]
    if patd.status in status_reconsideracao_e_posteriores:
         reconsideracao_context = doc_context.copy()
         if not patd.texto_reconsideracao and not patd.anexos.filter(tipo='reconsideracao').exists():
             reconsideracao_context['{Texto_reconsideracao}'] = '{Botao Adicionar Reconsideracao}'
         else:
             reconsideracao_context['{Texto_reconsideracao}'] = patd.texto_reconsideracao or "[Ver documentos anexos]"
         html_content = _render_document_from_template('MODELO_RECONSIDERACAO.docx', reconsideracao_context)
         html_content += "{ANEXOS_RECONSIDERACAO_PLACEHOLDER}"
         document_pages.append(html_content)

    # 7. Anexos e NPD da reconsideração
    status_npd_reconsideracao_e_posteriores = [
        'aguardando_preenchimento_npd_reconsideracao', 'aguardando_publicacao', 'finalizado'
    ]
    if patd.status in status_npd_reconsideracao_e_posteriores:
        html_content = "{ANEXO_OFICIAL_RECONSIDERACAO_PLACEHOLDER}"
        html_content += _render_document_from_template('MODELO_NPD_RECONSIDERACAO.docx', doc_context)
        document_pages.append(html_content)

    return document_pages


def _check_preclusao_signatures(patd):
    """
    Verifica se as assinaturas das testemunhas para o documento de preclusão
    foram coletadas, APENAS se as testemunhas tiverem sido designadas.
    Retorna True se as assinaturas estiverem completas, False caso contrário.
    """
    if patd.testemunha1 and not patd.assinatura_testemunha1:
        return False
    if patd.testemunha2 and not patd.assinatura_testemunha2:
        return False

    return True

def _check_and_finalize_patd(patd):
    """
    Verifica se todas as assinaturas necessárias para a NPD foram coletadas
    e, em caso afirmativo, avança o PATD para o período de reconsideração.
    """
    if patd.status != 'aguardando_assinatura_npd':
        return False

    document_pages = get_document_pages(patd)
    raw_document_text = "".join(document_pages)

    required_mil_signatures = raw_document_text.count('{Assinatura Militar Arrolado}')
    provided_mil_signatures = sum(1 for s in (patd.assinaturas_militar or []) if s)
    if provided_mil_signatures < required_mil_signatures:
        return False

    if not patd.testemunha1 or not patd.assinatura_testemunha1:
        return False

    if not patd.testemunha2 or not patd.assinatura_testemunha2:
        return False

    patd.status = 'periodo_reconsideracao'
    patd.data_publicacao_punicao = timezone.now()
    patd.save()
    return True

def _try_advance_status_from_justificativa(patd):
    """
    Verifica se a PATD no status 'aguardando_justificativa' pode avançar
    para 'em_apuracao'. Isso só deve ocorrer se tanto a alegação de defesa
    quanto todas as assinaturas necessárias estiverem presentes.
    """
    if patd.status != 'aguardando_justificativa':
        return False

    # Verifica se há texto de defesa OU anexos de defesa
    has_defesa = bool(patd.alegacao_defesa or patd.anexos.filter(tipo='defesa').exists())
    if not has_defesa:
        return False # Não avança se não há defesa registrada

    # Verifica assinaturas de ciência (já devem estar ok para chegar aqui, mas verificamos de novo)
    document_pages = get_document_pages(patd)
    raw_document_text = "".join(document_pages)
    required_signatures = raw_document_text.count('{Assinatura Militar Arrolado}')
    provided_signatures = sum(1 for s in (patd.assinaturas_militar or []) if s)

    if provided_signatures < required_signatures:
        logger.warning(f"PATD {patd.pk}: Tentativa de avançar de 'aguardando_justificativa', mas assinaturas de ciência incompletas ({provided_signatures}/{required_signatures}).")
        return False # Não avança se as assinaturas de ciência não estiverem completas

    # Se chegou aqui, a defesa existe e as assinaturas de ciência estão completas
    patd.status = 'em_apuracao'
    logger.info(f"PATD {patd.pk}: Avançando status de 'aguardando_justificativa' para 'em_apuracao'.")
    return True # Indica que o status foi alterado (será salvo na chamada principal)


# =============================================================================
# Views e Lógica da Aplicação
# =============================================================================

@login_required
@ouvidoria_required
def atribuir_oficial(request, pk):
    patd = get_object_or_404(PATD, pk=pk)
    if request.method == 'POST':
        form = AtribuirOficialForm(request.POST, instance=patd)
        if form.is_valid():
            form.save()
            messages.success(request, f'Oficial {patd.oficial_responsavel.nome_guerra} foi atribuído. Aguardando aceitação.')
            return redirect('Ouvidoria:patd_detail', pk=pk)
    else:
        form = AtribuirOficialForm(instance=patd)
    return render(request, 'atribuir_oficial.html', {'form': form, 'patd': patd})

@login_required
def patd_atribuicoes_pendentes(request):
    if not hasattr(request.user, 'profile') or not request.user.profile.militar:
        messages.warning(request, "Seu usuário não está associado a um militar.")
        return redirect('Ouvidoria:index')

    militar_logado = request.user.profile.militar
    active_tab = request.GET.get('tab', 'aprovar') # 'aprovar' is the default tab

    count_aprovar = PATD.objects.filter(
        oficial_responsavel=militar_logado,
        status='aguardando_aprovacao_atribuicao'
    ).count()

    status_list_apuracao = ['em_apuracao', 'apuracao_preclusao', 'aguardando_punicao', 'aguardando_punicao_alterar']
    count_apuracao = PATD.objects.filter(
        oficial_responsavel=militar_logado,
        status__in=status_list_apuracao
    ).count()

    if active_tab == 'apuracao':
        patds = PATD.objects.filter(
            oficial_responsavel=militar_logado,
            status__in=status_list_apuracao
        ).select_related('militar').order_by('-data_inicio')
    elif active_tab == 'todas':
        patds = PATD.objects.filter(
            oficial_responsavel=militar_logado
        ).select_related('militar').order_by('-data_inicio')
    else: # default is 'aprovar'
        patds = PATD.objects.filter(
            oficial_responsavel=militar_logado,
            status='aguardando_aprovacao_atribuicao'
        ).select_related('militar').order_by('-data_inicio')

    context = {
        'patds': patds,
        'active_tab': active_tab,
        'count_aprovar': count_aprovar,
        'count_apuracao': count_apuracao
    }

    return render(request, 'patd_atribuicoes_pendentes.html', context)

@login_required
@require_POST
def aceitar_atribuicao(request, pk):
    patd = get_object_or_404(PATD, pk=pk)

    if not (hasattr(request.user, 'profile') and request.user.profile.militar and request.user.profile.militar == patd.oficial_responsavel):
        messages.error(request, "Você não tem permissão para aceitar esta atribuição.")
        return redirect('Ouvidoria:patd_detail', pk=pk)

    form = AceitarAtribuicaoForm(request.POST)
    if form.is_valid():
        senha = form.cleaned_data['senha']
        user = authenticate(username=request.user.username, password=senha)
        if user is not None:
            # --- INÍCIO DA LÓGICA DE STATUS ---
            status_definido = False
            if patd.status_anterior:
                patd.status = patd.status_anterior
                patd.status_anterior = None
                status_definido = True
            else:
                patd.status = 'ciencia_militar'
                status_definido = True
            # --- FIM DA LÓGICA DE STATUS ---

            # Lógica para copiar assinatura padrão do oficial (existente)
            if patd.oficial_responsavel and patd.oficial_responsavel.assinatura:
                signature_data_base64 = patd.oficial_responsavel.assinatura
                try:
                    format, imgstr = signature_data_base64.split(';base64,')
                    ext = format.split('/')[-1]
                    file_content = ContentFile(base64.b64decode(imgstr), name=f'sig_oficial_{patd.pk}.{ext}')

                    if patd.assinatura_oficial:
                        patd.assinatura_oficial.delete(save=False)

                    patd.assinatura_oficial.save(file_content.name, file_content, save=False) # save=False aqui
                except Exception as e:
                    logger.error(f"Erro ao converter assinatura padrão do oficial para a PATD {pk}: {e}")
                    messages.error(request, "Erro ao processar a assinatura padrão do oficial.")
                    # Não salva a PATD e redireciona
                    return redirect('Ouvidoria:patd_atribuicoes_pendentes')

            # --- NOVA VERIFICAÇÃO DE ASSINATURAS DE CIÊNCIA ---
            if patd.status == 'ciencia_militar':
                try:
                    document_pages = get_document_pages(patd) # Gera o documento para contar placeholders
                    coringa_doc_text = document_pages[0] if document_pages else ""
                    required_initial_signatures = coringa_doc_text.count('{Assinatura Militar Arrolado}')
                    provided_signatures = sum(1 for s in (patd.assinaturas_militar or []) if s is not None)

                    logger.info(f"PATD {pk} aceita. Status: {patd.status}. Assinaturas requeridas: {required_initial_signatures}, Assinaturas providas: {provided_signatures}")

                    if provided_signatures >= required_initial_signatures:
                        if patd.data_ciencia is None: # Define a data da ciência se ainda não estiver definida
                            patd.data_ciencia = timezone.now()
                        patd.status = 'aguardando_justificativa' # Avança o status
                        logger.info(f"PATD {pk}: Assinaturas de ciência completas. Avançando status para 'aguardando_justificativa'.")
                        # Tenta avançar mais se a defesa já existir (caso raro, mas possível)
                        _try_advance_status_from_justificativa(patd)
                except Exception as e:
                    logger.error(f"Erro ao verificar assinaturas de ciência após aceite para PATD {pk}: {e}")
            # --- FIM DA NOVA VERIFICAÇÃO ---

            patd.save() # Salva a PATD com o status atualizado e a assinatura do oficial (se houver)
            messages.success(request, f'Atribuição da PATD Nº {patd.numero_patd} aceite com sucesso.')
            return redirect('Ouvidoria:patd_atribuicoes_pendentes')
        else:
            messages.error(request, "Senha incorreta. A atribuição não foi aceite.")
    else:
        messages.error(request, "Formulário inválido.")

    return redirect('Ouvidoria:patd_atribuicoes_pendentes')


@login_required
@require_GET
def patd_atribuicoes_pendentes_json(request):
    count = 0
    if hasattr(request.user, 'profile') and request.user.profile.militar:
        militar_logado = request.user.profile.militar

        count_aprovar = PATD.objects.filter(
            oficial_responsavel=militar_logado,
            status='aguardando_aprovacao_atribuicao'
        ).count()

        status_list_apuracao = ['em_apuracao', 'apuracao_preclusao', 'aguardando_punicao', 'aguardando_punicao_alterar']
        count_apuracao = PATD.objects.filter(
            oficial_responsavel=militar_logado,
            status__in=status_list_apuracao
        ).count()

        count = count_aprovar + count_apuracao

    return JsonResponse({'count': count})


@login_required
@comandante_redirect
@ouvidoria_required
def index(request):
    config = Configuracao.load()
    context = {
        'prazo_defesa_dias': config.prazo_defesa_dias,
        'prazo_defesa_minutos': config.prazo_defesa_minutos,
    }
    if request.method == 'POST':
        action = request.POST.get('action', 'analyze')

        # --- Modificação na ação 'create_militar_and_patd' (para consistência) ---
        if action == 'create_militar_and_patd':
            form = MilitarForm(request.POST)
            if form.is_valid():
                new_militar = form.save()
                transgressao = request.POST.get('transgressao')
                data_ocorrencia_str = request.POST.get('data_ocorrencia')
                protocolo_comaer = request.POST.get('protocolo_comaer', '')
                oficio_transgressao = request.POST.get('oficio_transgressao', '')
                data_oficio_str = request.POST.get('data_oficio', '') # String da data do ofício

                data_ocorrencia = None
                if data_ocorrencia_str:
                    try:
                        data_ocorrencia = datetime.strptime(data_ocorrencia_str, '%Y-%m-%d').date()
                    except (ValueError, TypeError):
                        pass

                # --- Lógica robusta para data_oficio ---
                data_oficio = None
                if data_oficio_str:
                    # Tenta remover prefixos comuns ANTES de converter
                    cleaned_data_oficio_str = re.sub(r"^[A-Za-z\s]+,\s*", "", data_oficio_str).strip() # Remove "Cidade, "
                    formats_to_try = ['%d/%m/%Y', '%d de %B de %Y', '%Y-%m-%d', '%d.%m.%Y']
                    for fmt in formats_to_try:
                        try:
                            if '%B' in fmt:
                                try:
                                    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
                                except locale.Error:
                                    logger.warning("Locale pt_BR.UTF-8 não encontrado ao criar PATD manualmente.")
                                    continue
                            # Usa a string limpa para conversão
                            data_oficio = datetime.strptime(cleaned_data_oficio_str, fmt).date()
                            break
                        except (ValueError, TypeError):
                            continue
                # --- Fim da lógica robusta ---

                patd = PATD.objects.create(
                    militar=new_militar,
                    transgressao=transgressao,
                    numero_patd=get_next_patd_number(),
                    data_ocorrencia=data_ocorrencia,
                    protocolo_comaer=protocolo_comaer,
                    oficio_transgressao=oficio_transgressao,
                    data_oficio=data_oficio # Salva a data convertida ou None
                )
                return JsonResponse({
                    'status': 'success',
                    'message': f'Militar "{new_militar.nome_guerra}" cadastrado e PATD Nº {patd.numero_patd} criada com sucesso!'
                })
            else:
                return JsonResponse({'status': 'error', 'errors': form.errors.as_json()}, status=400)


        # --- Modificação na ação 'analyze' ---
        elif action == 'analyze':
            pdf_file = request.FILES.get('pdf_file')
            if not pdf_file:
                return JsonResponse({'status': 'error', 'message': "Nenhum ficheiro foi enviado."}, status=400)

            try:
                # ... (código para ler PDF e chamar a IA - permanece igual) ...
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                    for chunk in pdf_file.chunks():
                        temp_file.write(chunk)
                    temp_file_path = temp_file.name

                loader = PyPDFLoader(temp_file_path)
                content = " ".join(page.page_content for page in loader.load_and_split())
                os.remove(temp_file_path)

                logger.info("Conteúdo do PDF extraído com sucesso. Chamando a IA para análise...")

                resultado_analise: AnaliseTransgressao = analisar_documento_pdf(content)

                logger.info(f"Resultado da análise da IA: {resultado_analise}")

                patds_criadas = []
                militares_nao_encontrados = []
                duplicatas_encontradas = []

                transgressao_comum = resultado_analise.transgressao
                data_ocorrencia = None
                if resultado_analise.data_ocorrencia:
                    try:
                        data_ocorrencia = datetime.strptime(resultado_analise.data_ocorrencia, '%Y-%m-%d').date()
                    except (ValueError, TypeError):
                        logger.warning(f"Formato inválido para data_ocorrencia: {resultado_analise.data_ocorrencia}")
                        pass

                # --- Lógica robusta REFINADA para data_oficio ---
                data_oficio = None
                data_oficio_str_from_ai = resultado_analise.data_oficio
                if data_oficio_str_from_ai:
                    # 1. Limpa prefixos comuns (Ex: "Rio de Janeiro, ")
                    cleaned_data_oficio_str = re.sub(r"^[A-Za-z\s]+,\s*", "", data_oficio_str_from_ai).strip()
                    logger.debug(f"String da data do ofício original: '{data_oficio_str_from_ai}', Limpa: '{cleaned_data_oficio_str}'")

                    # 2. Tenta os formatos na string limpa
                    formats_to_try = ['%d/%m/%Y', '%Y-%m-%d', '%d.%m.%Y', '%d de %B de %Y'] # Prioriza formatos numéricos
                    for fmt in formats_to_try:
                        try:
                            if '%B' in fmt:
                                try:
                                    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
                                except locale.Error:
                                    logger.warning(f"Locale pt_BR.UTF-8 não disponível para formato {fmt}.")
                                    continue
                            # Usa a string limpa para conversão
                            data_oficio = datetime.strptime(cleaned_data_oficio_str, fmt).date()
                            logger.info(f"Data do ofício '{cleaned_data_oficio_str}' convertida com sucesso usando o formato {fmt}.")
                            break # Para se um formato funcionar
                        except (ValueError, TypeError):
                            logger.debug(f"Falha ao converter data do ofício '{cleaned_data_oficio_str}' com formato {fmt}. Tentando próximo.")
                            continue # Tenta o próximo formato

                    if data_oficio is None:
                        logger.warning(f"Não foi possível converter a string da data do ofício '{data_oficio_str_from_ai}' (limpa: '{cleaned_data_oficio_str}') para data.")
                # --- Fim da lógica robusta REFINADA ---

                protocolo_comaer_comum = resultado_analise.protocolo_comaer
                oficio_transgressao_comum = resultado_analise.oficio_transgressao

                if not hasattr(resultado_analise, 'acusados') or not isinstance(resultado_analise.acusados, list):
                     logger.error(f"A resposta da IA não continha uma lista válida de 'acusados'. Resposta: {resultado_analise}")
                     raise ValueError("Formato de resposta inválido da IA: lista de acusados ausente ou malformada.")

                # Itera sobre os acusados (restante da lógica permanece igual)
                for acusado in resultado_analise.acusados:
                    nome_extraido = acusado.nome_militar
                    posto_graduacao_extraido = acusado.posto_graduacao

                    if not nome_extraido:
                        logger.warning(f"IA retornou um acusado sem nome: {acusado}")
                        continue

                    logger.info(f"Processando acusado: {posto_graduacao_extraido} {nome_extraido}")

                    militar = Militar.objects.filter(
                        Q(nome_completo__icontains=nome_extraido) |
                        Q(nome_guerra__icontains=nome_extraido)
                    ).first()

                    if militar:
                        logger.info(f"Militar encontrado no BD: {militar}")
                        existing_patds = PATD.objects.filter(militar=militar, data_ocorrencia=data_ocorrencia)

                        def similar(a, b):
                            return SequenceMatcher(None, a, b).ratio()

                        duplicata = False
                        for patd_existente in existing_patds:
                            if similar(transgressao_comum.strip().lower(), patd_existente.transgressao.strip().lower()) > 0.8:
                                patd_url = reverse('Ouvidoria:patd_detail', kwargs={'pk': patd_existente.pk})
                                duplicatas_encontradas.append({
                                    'nome_militar': str(militar),
                                    'numero_patd': patd_existente.numero_patd,
                                    'url': patd_url
                                })
                                duplicata = True
                                logger.info(f"PATD duplicada encontrada para {militar}: Nº {patd_existente.numero_patd}")
                                break

                        if not duplicata:
                            logger.info(f"Criando nova PATD para {militar}...")
                            patd = PATD.objects.create(
                                militar=militar,
                                transgressao=transgressao_comum,
                                numero_patd=get_next_patd_number(),
                                data_ocorrencia=data_ocorrencia,
                                protocolo_comaer=protocolo_comaer_comum,
                                oficio_transgressao=oficio_transgressao_comum,
                                data_oficio=data_oficio # Salva a data convertida ou None
                            )
                            patds_criadas.append({
                                'nome_militar': str(militar),
                                'numero_patd': patd.numero_patd
                            })
                            logger.info(f"PATD Nº {patd.numero_patd} criada para {militar}.")
                    else:
                        # ... (lógica para militar não encontrado) ...
                        logger.warning(f"Militar '{nome_extraido}' não encontrado no banco de dados.")
                        nome_para_cadastro = f"{posto_graduacao_extraido} {nome_extraido}".strip()
                        militares_nao_encontrados.append({
                            'nome_completo_sugerido': nome_para_cadastro,
                            'transgressao': transgressao_comum,
                            'data_ocorrencia': resultado_analise.data_ocorrencia,
                            'protocolo_comaer': protocolo_comaer_comum,
                            'oficio_transgressao': oficio_transgressao_comum,
                            # Passa a STRING original da IA para o formulário
                            'data_oficio': data_oficio_str_from_ai
                        })

                response_data = {
                    'status': 'processed',
                    'patds_criadas': patds_criadas,
                    'militares_nao_encontrados': militares_nao_encontrados,
                    'duplicatas_encontradas': duplicatas_encontradas
                }
                logger.info(f"Análise concluída. Resposta: {response_data}")
                return JsonResponse(response_data)

            except Exception as e:
                # ... (bloco de tratamento de exceção - permanece igual) ...
                 error_type = type(e).__name__
                 error_message = str(e)
                 error_traceback = traceback.format_exc()

                 logger.error(f"Erro detalhado na análise do PDF: {error_type} - {error_message}\nTraceback:\n{error_traceback}")

                 user_message = f"Ocorreu um erro inesperado durante a análise ({error_type}). Verifique os logs do servidor para mais detalhes."
                 # ... (mensagens de erro específicas) ...

                 return JsonResponse({
                     'status': 'error',
                     'message': user_message,
                     'detail': f"{error_type}: {error_message}"
                 }, status=500)

    # Lógica para GET request (permanece igual)
    return render(request, 'indexOuvidoria.html', context)


@login_required
@ouvidoria_required
def importar_excel(request):
    config = Configuracao.load()
    context = {
        'prazo_defesa_dias': config.prazo_defesa_dias,
        'prazo_defesa_minutos': config.prazo_defesa_minutos,
    }
    if request.method == 'POST':
        excel_file = request.FILES.get('excel_file')
        if not excel_file:
            messages.error(request, "Nenhum ficheiro foi enviado.")
            return redirect('Ouvidoria:importar_excel')

        try:
            df = pd.read_excel(excel_file).fillna('')
            column_mapping = {
                'pst.': 'posto', 'quad.': 'quad', 'esp.': 'especializacao', 'saram': 'saram',
                'nome completo': 'nome_completo', 'nome de guerra': 'nome_guerra', 'turma': 'turma',
                'situação': 'situacao', 'om': 'om', 'setor': 'setor', 'subsetor': 'subsetor'
            }
            df.columns = df.columns.str.lower().str.strip()
            militares_criados = 0
            militares_atualizados = 0
            linhas_com_erro = 0
            for index, row in df.iterrows():
                data_dict = {}
                for excel_col, model_field in column_mapping.items():
                    if excel_col in df.columns:
                        data_dict[model_field] = row[excel_col]
                postos_oficiais = ['TC', 'MJ', 'CP', '1T', '2T']
                is_recruta = str(data_dict.get('posto', '')).upper() == 'REC'
                if 'posto' in data_dict:
                    data_dict['oficial'] = str(data_dict.get('posto', '')).upper() in postos_oficiais
                identifier = {}
                if is_recruta:
                    if 'nome_completo' in data_dict and data_dict['nome_completo']:
                        identifier['nome_completo'] = data_dict['nome_completo']
                        data_dict.pop('saram', None)
                    else:
                        linhas_com_erro += 1
                        continue
                else:
                    if 'saram' in data_dict and str(data_dict['saram']).strip():
                        try:
                            identifier['saram'] = int(data_dict['saram'])
                        except (ValueError, TypeError):
                            linhas_com_erro += 1
                            continue
                    else:
                        linhas_com_erro += 1
                        continue
                try:
                    for field in Militar._meta.get_fields():
                        if not field.is_relation and field.name not in data_dict and not field.blank and field.name not in identifier:
                            if is_recruta and field.name == 'saram':
                                continue
                            data_dict[field.name] = ''
                    obj, created = Militar.objects.update_or_create(**identifier, defaults=data_dict)
                    if created:
                        militares_criados += 1
                    else:
                        militares_atualizados += 1
                except Exception as e:
                    logger.warning(f"Não foi possível processar a linha {index+2}: {e}. Dados: {row}")
                    linhas_com_erro += 1
                    continue
            msg = f"Importação concluída! {militares_criados} militares criados e {militares_atualizados} atualizados."
            if linhas_com_erro > 0:
                msg += f" {linhas_com_erro} linhas foram ignoradas por dados inválidos ou ausência de identificador."
            messages.success(request, msg)
        except Exception as e:
            logger.error(f"Erro na importação do Excel: {e}")
            messages.error(request, f"Ocorreu um erro ao processar o ficheiro: {e}")
        return redirect('Ouvidoria:importar_excel')
    return render(request, 'importar_excel.html', context)

@method_decorator([login_required, ouvidoria_required], name='dispatch')
class MilitarListView(ListView):
    model = Militar
    template_name = 'militar_list.html'
    context_object_name = 'militares'
    paginate_by = 25

    def get_queryset(self):
        query = self.request.GET.get('q')
        rank_order = Case(
            When(posto='TC', then=Value(0)), When(posto='MJ', then=Value(1)), When(posto='CP', then=Value(2)),
            When(posto='1T', then=Value(3)), When(posto='2T', then=Value(4)), When(posto='SO', then=Value(5)),
            When(posto='1S', then=Value(6)), When(posto='2S', then=Value(7)), When(posto='3S', then=Value(8)),
            When(posto='CB', then=Value(9)), When(posto='S1', then=Value(10)), When(posto='S2', then=Value(11)),
            default=Value(99), output_field=IntegerField(),
        )
        qs = super().get_queryset().annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_completo')
        if query:
            qs = qs.filter(
                Q(nome_completo__icontains=query) |
                Q(nome_guerra__icontains=query) |
                Q(saram__icontains=query)
            )
        # --- AJAX Handling ---
        # if self.request.headers.get('X-Requested-With') == 'XMLHttpRequest':
             # Se for AJAX, renderiza apenas a parte da lista e paginação
             # context = {'militares': qs, 'is_paginated': False} # Adapte a paginação se necessário
             # Retorna o HTML parcial
             # return render(self.request, 'partials/militar_list_partial.html', context) # Crie este template parcial
        return qs

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        config = Configuracao.load()
        context['prazo_defesa_dias'] = config.prazo_defesa_dias
        context['prazo_defesa_minutos'] = config.prazo_defesa_minutos
        return context

@method_decorator([login_required, ouvidoria_required], name='dispatch')
class MilitarCreateView(CreateView):
    model = Militar
    form_class = MilitarForm
    template_name = 'militar_form.html'
    success_url = reverse_lazy('Ouvidoria:militar_list')

@method_decorator([login_required, ouvidoria_required], name='dispatch')
class MilitarUpdateView(UpdateView):
    model = Militar
    form_class = MilitarForm
    template_name = 'militar_form.html'
    success_url = reverse_lazy('Ouvidoria:militar_list')

@method_decorator([login_required, ouvidoria_required], name='dispatch')
class MilitarDeleteView(DeleteView):
    model = Militar
    template_name = 'militar_confirm_delete.html'
    success_url = reverse_lazy('Ouvidoria:militar_list') # Alterado para militar_list

# Dicionário com os grupos de status
STATUS_GROUPS = {
    "Aguardando Oficial": {
        'definicao_oficial': 'Aguardando definição do Oficial',
        'aguardando_aprovacao_atribuicao': 'Aguardando aprovação de atribuição de oficial',
    },
    "Fase de Defesa": {
        'ciencia_militar': 'Aguardando ciência do militar',
        'aguardando_justificativa': 'Aguardando Justificativa',
        'prazo_expirado': 'Prazo expirado',
        'preclusao': 'Preclusão - Sem Defesa',
    },
    "Fase de Apuração": {
        'em_apuracao': 'Em Apuração',
        'apuracao_preclusao': 'Em Apuração (Preclusão)',
        'aguardando_punicao': 'Aguardando Aplicação da Punição',
        'aguardando_punicao_alterar': 'Aguardando Punição (alterar)',
    },
    "Decisão do Comandante": {
        'analise_comandante': 'Em Análise pelo Comandante',
        'aguardando_assinatura_npd': 'Aguardando Assinatura NPD',
    },
    "Fase de Reconsideração": {
        'periodo_reconsideracao': 'Período de Reconsideração',
        'em_reconsideracao': 'Em Reconsideração',
        'aguardando_comandante_base': 'Aguardando Comandante da Base',
        'aguardando_preenchimento_npd_reconsideracao': 'Aguardando preenchimento NPD Reconsideração',
    },
    "Aguardando Publicação": {
        'aguardando_publicacao': 'Aguardando publicação',
    }
}

@method_decorator([login_required, comandante_redirect, ouvidoria_required], name='dispatch')
class PATDListView(ListView):
    model = PATD
    template_name = 'patd_list.html'
    context_object_name = 'patds'
    paginate_by = 15

    def get_queryset(self):
        query = self.request.GET.get('q')
        status_filter = self.request.GET.get('status')

        qs = super().get_queryset().exclude(status='finalizado').select_related('militar', 'oficial_responsavel').order_by('-data_inicio')

        if query:
            qs = qs.filter(
                Q(numero_patd__icontains=query) |
                Q(militar__nome_completo__icontains=query) |
                Q(militar__nome_guerra__icontains=query)
            )

        if status_filter:
            if status_filter in STATUS_GROUPS:
                statuses_in_group = list(STATUS_GROUPS[status_filter].keys())
                qs = qs.filter(status__in=statuses_in_group)
            else:
                qs = qs.filter(status=status_filter)

        return qs

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        config = Configuracao.load()
        context['prazo_defesa_dias'] = config.prazo_defesa_dias
        context['prazo_defesa_minutos'] = config.prazo_defesa_minutos
        context['status_groups'] = STATUS_GROUPS
        context['current_status'] = self.request.GET.get('status', '')
        return context

@method_decorator([login_required, ouvidoria_required], name='dispatch')
class PatdFinalizadoListView(ListView):
    model = PATD
    template_name = 'patd_finalizado_list.html'
    context_object_name = 'patds'
    paginate_by = 15

    def get_queryset(self):
        return PATD.objects.filter(status='finalizado').select_related('militar').order_by('-data_inicio')

@method_decorator([login_required, ouvidoria_required], name='dispatch')
class PATDDetailView(DetailView):
    model = PATD
    template_name = 'patd_detail.html'
    context_object_name = 'patd'

    def get_queryset(self):
        return super().get_queryset().select_related(
            'militar', 'oficial_responsavel', 'testemunha1', 'testemunha2'
        ).prefetch_related('anexos')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        patd = self.get_object()
        config = Configuracao.load()

        document_pages = get_document_pages(patd)
        context['documento_texto_json'] = json.dumps(document_pages)

        context['assinaturas_militar_json'] = json.dumps(patd.assinaturas_militar or [])


        context['now_iso'] = timezone.now().isoformat()
        context['prazo_defesa_dias'] = config.prazo_defesa_dias
        context['prazo_defesa_minutos'] = config.prazo_defesa_minutos

        doc_context = _get_document_context(patd)
        context['comandante_assinatura'] = doc_context.get('assinatura_comandante_data')

        context['analise_data_json'] = json.dumps({
            'itens': patd.itens_enquadrados,
            'circunstancias': patd.circunstancias,
            'punicao': patd.punicao_sugerida
        }) if patd.punicao_sugerida else 'null'

        militar_acusado = patd.militar
        patds_anteriores = PATD.objects.filter(
            militar=militar_acusado
        ).exclude(pk=patd.pk).order_by('-data_inicio')

        historico_punicoes = []
        for p_antiga in patds_anteriores:
            if p_antiga.punicao:
                itens_str = ""
                if p_antiga.itens_enquadrados and isinstance(p_antiga.itens_enquadrados, list):
                    itens_str = ", ".join([str(item.get('numero', '')) for item in p_antiga.itens_enquadrados if 'numero' in item])

                historico_punicoes.append({
                    'numero_patd': p_antiga.numero_patd,
                    'punicao': f"{p_antiga.dias_punicao} de {p_antiga.punicao}" if p_antiga.dias_punicao else p_antiga.punicao,
                    'itens': itens_str,
                    'data': p_antiga.data_inicio.strftime('%d/%m/%Y')
                })

        context['historico_punicoes'] = historico_punicoes

        anexos_defesa = patd.anexos.filter(tipo='defesa')
        anexos_defesa_data = []
        for a in anexos_defesa:
            anexos_defesa_data.append({
                'id': a.id,
                'nome': os.path.basename(a.arquivo.name),
                'url': a.arquivo.url,
                'content_html': get_anexo_content_as_html(a)
            })
        context['anexos_defesa_json'] = json.dumps(anexos_defesa_data)

        anexos_reconsideracao = patd.anexos.filter(tipo='reconsideracao')
        anexos_reconsideracao_data = []
        for a in anexos_reconsideracao:
            anexos_reconsideracao_data.append({
                'id': a.id,
                'nome': os.path.basename(a.arquivo.name),
                'url': a.arquivo.url,
                'content_html': get_anexo_content_as_html(a)
            })
        context['anexos_reconsideracao_json'] = json.dumps(anexos_reconsideracao_data)

        anexos_reconsideracao_oficial = patd.anexos.filter(tipo='reconsideracao_oficial')
        anexos_reconsideracao_oficial_data = []
        for a in anexos_reconsideracao_oficial:
            anexos_reconsideracao_oficial_data.append({
                'id': a.id,
                'nome': os.path.basename(a.arquivo.name),
                'url': a.arquivo.url,
                'content_html': get_anexo_content_as_html(a)
            })
        context['anexos_reconsideracao_oficial_json'] = json.dumps(anexos_reconsideracao_oficial_data)

        return context

@method_decorator([login_required, ouvidoria_required], name='dispatch')
class PATDUpdateView(UserPassesTestMixin, UpdateView):
    model = PATD
    form_class = PATDForm
    template_name = 'patd_form.html'

    def test_func(self):
        return not has_comandante_access(self.request.user)

    def handle_no_permission(self):
        messages.error(self.request, "Acesso negado. Comandantes não podem editar o processo.")
        patd_pk = self.kwargs.get('pk')
        if patd_pk:
            return redirect('Ouvidoria:patd_detail', pk=patd_pk)
        return redirect('Ouvidoria:index')

    def get_success_url(self):
        return reverse_lazy('Ouvidoria:patd_detail', kwargs={'pk': self.object.pk})

@method_decorator([login_required, ouvidoria_required], name='dispatch')
class PATDDeleteView(DeleteView):
    model = PATD
    template_name = 'militar_confirm_delete.html'
    success_url = reverse_lazy('Ouvidoria:patd_list')

@method_decorator([login_required, ouvidoria_required], name='dispatch')
class MilitarPATDListView(ListView):
    model = PATD
    template_name = 'militar_patd_list.html'
    context_object_name = 'patds'
    paginate_by = 10

    def get_queryset(self):
        self.militar = get_object_or_404(Militar, pk=self.kwargs['pk'])
        return PATD.objects.filter(militar=self.militar).order_by('-data_inicio')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['militar'] = self.militar
        return context

@login_required
@oficial_responsavel_required
@require_POST
def salvar_assinatura(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)
        data = json.loads(request.body)
        signature_data_base64 = data.get('signature_data')

        if not signature_data_base64:
            return JsonResponse({'status': 'error', 'message': 'Nenhum dado de assinatura recebido.'}, status=400)

        try:
            format, imgstr = signature_data_base64.split(';base64,')
            ext = format.split('/')[-1]
            file_content = ContentFile(base64.b64decode(imgstr), name=f'sig_oficial_{pk}.{ext}')

            if patd.assinatura_oficial:
                patd.assinatura_oficial.delete(save=False)

            patd.assinatura_oficial.save(file_content.name, file_content, save=True)
        except Exception as e:
            logger.error(f"Erro ao converter Base64 para ficheiro para PATD {pk}: {e}")
            return JsonResponse({'status': 'error', 'message': 'Erro ao processar a imagem da assinatura.'}, status=500)

        return JsonResponse({'status': 'success', 'message': 'Assinatura salva com sucesso.'})
    except Exception as e:
        logger.error(f"Erro ao salvar assinatura do oficial para PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
@require_POST
def salvar_assinatura_ciencia(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)
        data = json.loads(request.body)
        signature_data_base64 = data.get('signature_data')
        assinatura_index = int(data.get('assinatura_index', -1))

        if not signature_data_base64 or assinatura_index < 0:
            return JsonResponse({'status': 'error', 'message': 'Dados de assinatura inválidos.'}, status=400)

        try:
            format, imgstr = signature_data_base64.split(';base64,')
            ext = format.split('/')[-1]
            file_name = f'sig_ciencia_{assinatura_index}_{pk}_{uuid4().hex[:6]}.{ext}'
            file_content = ContentFile(base64.b64decode(imgstr))

            anexo = Anexo.objects.create(patd=patd, tipo='assinatura_ciencia')
            anexo.arquivo.save(file_name, file_content, save=True)
            signature_url = anexo.arquivo.url

        except Exception as e:
            logger.error(f"Erro ao converter Base64 da assinatura de ciência para ficheiro (PATD {pk}): {e}")
            return JsonResponse({'status': 'error', 'message': 'Erro ao processar a imagem da assinatura.'}, status=500)

        if patd.assinaturas_militar is None:
            patd.assinaturas_militar = []

        while len(patd.assinaturas_militar) <= assinatura_index:
            patd.assinaturas_militar.append(None)

        patd.assinaturas_militar[assinatura_index] = signature_url

        if patd.status == 'ciencia_militar':
            document_pages = get_document_pages(patd)
            coringa_doc_text = document_pages[0] if document_pages else ""
            required_initial_signatures = coringa_doc_text.count('{Assinatura Militar Arrolado}')
            provided_signatures = sum(1 for s in patd.assinaturas_militar if s is not None)
            if provided_signatures >= required_initial_signatures:
                if patd.data_ciencia is None:
                    patd.data_ciencia = timezone.now()
                patd.status = 'aguardando_justificativa'

        _try_advance_status_from_justificativa(patd)
        _check_and_finalize_patd(patd)

        patd.save()

        return JsonResponse({'status': 'success', 'message': 'Assinatura registrada com sucesso.'})
    except Exception as e:
        logger.error(f"Erro ao salvar assinatura de ciência da PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
@require_POST
def salvar_alegacao_defesa(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)

        if patd.alegacao_defesa or patd.anexos.filter(tipo='defesa').exists():
            return JsonResponse({
                'status': 'error',
                'message': 'A alegação de defesa já foi enviada e não pode ser alterada.'
            }, status=403)

        alegacao_texto = request.POST.get('alegacao_defesa', '')
        arquivos = request.FILES.getlist('anexos_defesa')

        if not alegacao_texto and not arquivos:
            return JsonResponse({'status': 'error', 'message': 'É necessário fornecer um texto ou anexar pelo menos um ficheiro.'}, status=400)

        patd.alegacao_defesa = alegacao_texto
        patd.data_alegacao = timezone.now()

        for arquivo in arquivos:
            Anexo.objects.create(patd=patd, arquivo=arquivo, tipo='defesa')

        try:
            if not patd.alegacao_defesa_resumo:
                patd.alegacao_defesa_resumo = analisar_e_resumir_defesa(patd.alegacao_defesa)
            if not patd.ocorrencia_reescrita:
                ocorrencia_formatada = reescrever_ocorrencia(patd.transgressao)
                patd.ocorrencia_reescrita = ocorrencia_formatada
                patd.comprovante = ocorrencia_formatada
        except Exception as e:
            logger.error(f"Erro ao chamar a IA para processar textos da PATD {pk}: {e}")
            if not patd.alegacao_defesa_resumo:
                patd.alegacao_defesa_resumo = "Erro ao gerar resumo."
            if not patd.ocorrencia_reescrita:
                patd.ocorrencia_reescrita = patd.transgressao
                patd.comprovante = patd.transgressao

        patd.save()
        _try_advance_status_from_justificativa(patd)
        patd.save()

        return JsonResponse({'status': 'success', 'message': 'Alegação de defesa e anexos salvos com sucesso.'})
    except Exception as e:
        logger.error(f"Erro ao salvar alegação de defesa da PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
@require_POST
def salvar_assinatura_defesa(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)
        data = json.loads(request.body)
        signature_data_base64 = data.get('signature_data')

        if not signature_data_base64:
            return JsonResponse({'status': 'error', 'message': 'Nenhum dado de assinatura recebido.'}, status=400)

        try:
            format, imgstr = signature_data_base64.split(';base64,')
            ext = format.split('/')[-1]
            file_content = ContentFile(base64.b64decode(imgstr), name=f'sig_defesa_{pk}.{ext}')

            if patd.assinatura_alegacao_defesa:
                patd.assinatura_alegacao_defesa.delete(save=False)

            patd.assinatura_alegacao_defesa.save(file_content.name, file_content, save=True)
        except Exception as e:
            logger.error(f"Erro ao converter Base64 para ficheiro para PATD {pk}: {e}")
            return JsonResponse({'status': 'error', 'message': 'Erro ao processar a imagem da assinatura.'}, status=500)

        return JsonResponse({'status': 'success', 'message': 'Assinatura da defesa salva com sucesso.'})
    except Exception as e:
        logger.error(f"Erro ao salvar assinatura da defesa da PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
@require_POST
def salvar_assinatura_reconsideracao(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)

        if patd.status != 'em_reconsideracao':
            return JsonResponse({'status': 'error', 'message': 'A PATD não está na fase correta para assinar a reconsideração.'}, status=400)

        data = json.loads(request.body)
        signature_data_base64 = data.get('signature_data')

        if not signature_data_base64:
            return JsonResponse({'status': 'error', 'message': 'Nenhum dado de assinatura recebido.'}, status=400)

        try:
            format, imgstr = signature_data_base64.split(';base64,')
            ext = format.split('/')[-1]
            file_content = ContentFile(base64.b64decode(imgstr), name=f'sig_reconsideracao_{pk}.{ext}')

            if patd.assinatura_reconsideracao:
                patd.assinatura_reconsideracao.delete(save=False)

            patd.assinatura_reconsideracao.save(file_content.name, file_content, save=True)
            logger.info(f"Assinatura de reconsideração para PATD {pk} salva em {patd.assinatura_reconsideracao.path}")

        except Exception as e:
            logger.error(f"Erro ao converter Base64 para ficheiro para PATD {pk}: {e}")
            return JsonResponse({'status': 'error', 'message': 'Erro ao processar a imagem da assinatura.'}, status=500)

        _check_and_advance_reconsideracao_status(pk)

        return JsonResponse({'status': 'success', 'message': 'Assinatura salva com sucesso.'})
    except Exception as e:
        logger.error(f"Erro ao salvar assinatura da reconsideração da PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
@require_POST
def extender_prazo(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)
        if patd.status != 'prazo_expirado':
            return JsonResponse({'status': 'error', 'message': 'O prazo só pode ser estendido se estiver expirado.'}, status=400)

        data = json.loads(request.body)
        dias_extensao = int(data.get('dias', 0))
        minutos_extensao = int(data.get('minutos', 0))

        if not request.user.is_superuser and minutos_extensao != 0:
            return JsonResponse({'status': 'error', 'message': 'Apenas administradores podem alterar os minutos.'}, status=403)


        if dias_extensao < 0 or minutos_extensao < 0:
            return JsonResponse({'status': 'error', 'message': 'Valores de extensão de prazo inválidos.'}, status=400)

        config = Configuracao.load()

        delta_dias = config.prazo_defesa_dias - dias_extensao
        delta_minutos = config.prazo_defesa_minutos - minutos_extensao

        patd.data_ciencia = timezone.now() - timedelta(days=delta_dias, minutes=delta_minutos)
        patd.status = 'aguardando_justificativa'
        patd.save()

        return JsonResponse({'status': 'success', 'message': 'Prazo estendido com sucesso.'})

    except (ValueError, TypeError):
        return JsonResponse({'status': 'error', 'message': 'Dados de entrada inválidos.'}, status=400)
    except Exception as e:
        logger.error(f"Erro ao estender prazo da PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
@require_POST
def salvar_documento_patd(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)
        data = json.loads(request.body)
        texto_documento = data.get('texto_documento')

        if texto_documento is None:
            return JsonResponse({'status': 'error', 'message': 'Nenhum texto recebido.'}, status=400)

        patd.documento_texto = texto_documento
        patd.save()

        return JsonResponse({'status': 'success', 'message': 'Documento salvo com sucesso.'})
    except Exception as e:
        logger.error(f"Erro ao salvar documento da PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
@require_GET
def lista_oficiais(request):
    query = request.GET.get('q', '')
    oficiais = Militar.objects.filter(oficial=True)
    if query:
        oficiais = oficiais.filter(
            Q(nome_completo__icontains=query) |
            Q(nome_guerra__icontains=query)
        )
    oficiais = oficiais.order_by('posto', 'nome_guerra')
    data = list(oficiais.values('id', 'posto', 'nome_guerra', 'assinatura'))
    return JsonResponse(data, safe=False)

@login_required
@ouvidoria_required
@require_POST
def salvar_assinatura_padrao(request, pk):
    if not request.user.is_superuser:
        return JsonResponse({'status': 'error', 'message': 'Apenas administradores podem alterar assinaturas.'}, status=403)
    try:
        oficial = get_object_or_404(Militar, pk=pk, oficial=True)
        data = json.loads(request.body)
        signature_data = data.get('signature_data', '')
        oficial.assinatura = signature_data
        oficial.save()
        return JsonResponse({'status': 'success', 'message': 'Assinatura padrão salva com sucesso.'})
    except Militar.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Oficial não encontrado.'}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
def gerenciar_configuracoes_padrao(request):
    config = Configuracao.load()
    if request.method == 'POST':
        if not request.user.is_superuser:
            return JsonResponse({'status': 'error', 'message': 'Apenas administradores podem alterar as configurações.'}, status=403)
        try:
            data = json.loads(request.body)
            comandante_gsd_id = data.get('comandante_gsd_id')
            comandante_bagl_id = data.get('comandante_bagl_id')
            prazo_dias = data.get('prazo_defesa_dias')
            prazo_minutos = data.get('prazo_defesa_minutos')

            if comandante_gsd_id:
                comandante = get_object_or_404(Militar, pk=comandante_gsd_id, oficial=True)
                config.comandante_gsd = comandante
            else:
                config.comandante_gsd = None

            if comandante_bagl_id:
                comandante_bagl = get_object_or_404(Militar, pk=comandante_bagl_id, oficial=True)
                config.comandante_bagl = comandante_bagl
            else:
                config.comandante_bagl = None

            if prazo_dias is not None:
                config.prazo_defesa_dias = int(prazo_dias)
            if prazo_minutos is not None:
                config.prazo_defesa_minutos = int(prazo_minutos)

            config.save()
            return JsonResponse({'status': 'success', 'message': 'Configurações salvas com sucesso.'})
        except (ValueError, TypeError):
            return JsonResponse({'status': 'error', 'message': 'Prazo de defesa inválido.'}, status=400)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

    oficiais = Militar.objects.filter(oficial=True).order_by('posto', 'nome_guerra')
    oficiais_data = [{'id': o.id, 'texto': f"{o.posto} {o.nome_guerra}"} for o in oficiais]
    data = {
        'comandante_gsd_id': config.comandante_gsd.id if config.comandante_gsd else None,
        'comandante_bagl_id': config.comandante_bagl.id if config.comandante_bagl else None,
        'oficiais': oficiais_data,
        'prazo_defesa_dias': config.prazo_defesa_dias,
        'prazo_defesa_minutos': config.prazo_defesa_minutos
    }
    return JsonResponse(data)

@login_required
@ouvidoria_required
@require_GET
def patds_expirados_json(request):
    patds_expiradas = PATD.objects.filter(status='prazo_expirado').select_related('militar')
    data = [{'id': p.id, 'numero_patd': p.numero_patd, 'militar_nome': str(p.militar)} for p in patds_expiradas]
    return JsonResponse(data, safe=False)

@login_required
@ouvidoria_required
@require_POST
def extender_prazo_massa(request):
    try:
        data = json.loads(request.body)
        dias_extensao = int(data.get('dias', 5))
        minutos_extensao = int(data.get('minutos', 0))
        if dias_extensao < 0 or minutos_extensao < 0:
            return JsonResponse({'status': 'error', 'message': 'Valores de extensão inválidos.'}, status=400)
        patds_expiradas = PATD.objects.filter(status='prazo_expirado')
        if not patds_expiradas.exists():
            return JsonResponse({'status': 'no_action', 'message': 'Nenhuma PATD com prazo expirado para atualizar.'})
        count = 0
        for patd in patds_expiradas:
            config = Configuracao.load()
            delta_dias = config.prazo_defesa_dias - dias_extensao
            delta_minutos = config.prazo_defesa_minutos - minutos_extensao
            patd.data_ciencia = timezone.now() - timedelta(days=delta_dias, minutes=delta_minutos)
            patd.status = 'aguardando_justificativa'
            patd.save()
            count += 1
        return JsonResponse({'status': 'success', 'message': f'{count} prazos foram estendidos com sucesso.'})
    except (ValueError, TypeError):
        return JsonResponse({'status': 'error', 'message': 'Dados de entrada inválidos.'}, status=400)
    except Exception as e:
        logger.error(f"Erro ao estender prazos em massa: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
@require_POST
def verificar_e_atualizar_prazos(request):
    try:
        prazos_atualizados = 0
        patds_pendentes = PATD.objects.filter(status='aguardando_justificativa')
        config = Configuracao.load()

        for patd in patds_pendentes:
            if patd.data_ciencia:
                dias_uteis_a_adicionar = config.prazo_defesa_dias
                data_final = patd.data_ciencia
                dias_adicionados = 0
                while dias_adicionados < dias_uteis_a_adicionar:
                    data_final += timedelta(days=1)
                    if data_final.weekday() < 5: # 0-4 são dias úteis
                        dias_adicionados += 1

                deadline = (data_final + timedelta(days=1)).replace(hour=0, minute=0, second=0, microsecond=0)

                if timezone.now() > deadline:
                    patd.status = 'prazo_expirado'
                    patd.save(update_fields=['status'])
                    prazos_atualizados += 1

        patds_em_reconsideracao = PATD.objects.filter(status='periodo_reconsideracao')
        reconsideracoes_finalizadas = 0
        for patd in patds_em_reconsideracao:
            if patd.data_publicacao_punicao:
                deadline = patd.data_publicacao_punicao + timedelta(days=15)
                if timezone.now() > deadline:
                    patd.status = 'aguardando_publicacao'
                    patd.save(update_fields=['status'])
                    reconsideracoes_finalizadas += 1

        total_updated = prazos_atualizados + reconsideracoes_finalizadas
        return JsonResponse({'status': 'success', 'updated_count': total_updated})

    except Exception as e:
        logger.error(f"Erro ao verificar e atualizar prazos: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
@require_POST
def prosseguir_sem_alegacao(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)
        if patd.status != 'prazo_expirado':
            return JsonResponse({'status': 'error', 'message': 'Ação permitida apenas para PATDs com prazo expirado.'}, status=400)
        patd.status = 'apuracao_preclusao'
        patd.save(update_fields=['status'])
        return JsonResponse({'status': 'success', 'message': 'PATD atualizada para Apuração (Preclusão).'})
    except Exception as e:
        logger.error(f"Erro ao prosseguir sem alegação para PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
@require_POST
def salvar_assinatura_testemunha(request, pk, testemunha_num):
    try:
        patd = get_object_or_404(PATD, pk=pk)
        data = json.loads(request.body)
        signature_data_base64 = data.get('signature_data')

        if not signature_data_base64:
            return JsonResponse({'status': 'error', 'message': 'Nenhum dado de assinatura recebido.'}, status=400)

        try:
            format, imgstr = signature_data_base64.split(';base64,')
            ext = format.split('/')[-1]
            file_content = ContentFile(base64.b64decode(imgstr), name=f'sig_testemunha_{testemunha_num}_{pk}.{ext}')

            if testemunha_num == 1:
                if patd.assinatura_testemunha1:
                    patd.assinatura_testemunha1.delete(save=False)
                patd.assinatura_testemunha1.save(file_content.name, file_content, save=False)
            elif testemunha_num == 2:
                if patd.assinatura_testemunha2:
                    patd.assinatura_testemunha2.delete(save=False)
                patd.assinatura_testemunha2.save(file_content.name, file_content, save=False)
            else:
                return JsonResponse({'status': 'error', 'message': 'Número de testemunha inválido.'}, status=400)

            patd.save()

        except Exception as e:
            logger.error(f"Erro ao converter Base64 para ficheiro para PATD {pk}: {e}")
            return JsonResponse({'status': 'error', 'message': 'Erro ao processar a imagem da assinatura.'}, status=500)

        if _check_and_finalize_patd(patd):
             patd.save()

        return JsonResponse({'status': 'success', 'message': f'Assinatura da {testemunha_num}ª testemunha salva.'})
    except Exception as e:
        logger.error(f"Erro ao salvar assinatura da testemunha {testemunha_num} para PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@oficial_responsavel_required
@require_POST
def analisar_punicao(request, pk):
    patd = get_object_or_404(PATD, pk=pk)
    force_reanalyze = False

    if request.content_type == 'application/json':
        try:
            data = json.loads(request.body)
            force_reanalyze = data.get('force_reanalyze', False)
        except json.JSONDecodeError:
            pass

    if patd.punicao_sugerida and not force_reanalyze:
        return JsonResponse({
            'status': 'success',
            'analise_data': {
                'itens': patd.itens_enquadrados,
                'circunstancias': patd.circunstancias,
                'punicao': patd.punicao_sugerida
            }
        })

    try:
        itens_obj = enquadra_item(patd.transgressao)
        itens_list = [item for item in itens_obj.item]
        patd.itens_enquadrados = itens_list

        militar_acusado = patd.militar
        patds_anteriores = PATD.objects.filter(
            militar=militar_acusado
        ).exclude(pk=patd.pk).order_by('-data_inicio') # Adiciona order_by para pegar o mais recente

        historico_list = []
        
        # --- INÍCIO DA CORREÇÃO ---
        # Busca o comportamento anterior baseado na última PATD finalizada do militar
        patd_mais_recente = patds_anteriores.first()
        comportamento_anterior = "Permanece no \"Bom comportamento\"" # Valor padrão
        
        if patd_mais_recente and patd_mais_recente.comportamento:
            comportamento_anterior = patd_mais_recente.comportamento
        # --- FIM DA CORREÇÃO ---

        if patds_anteriores.exists():
            for p_antiga in patds_anteriores:
                if p_antiga.itens_enquadrados and isinstance(p_antiga.itens_enquadrados, list):
                    itens_str = ", ".join([f"Item {item.get('numero')}" for item in p_antiga.itens_enquadrados if 'numero' in item])
                    if itens_str:
                         historico_list.append(f"PATD anterior (Nº {p_antiga.numero_patd}) foi enquadrada em: {itens_str}.")

        historico_militar = "\n".join(historico_list) if historico_list else "Nenhuma punição anterior registrada."
        justificativa = patd.alegacao_defesa or "Nenhuma alegação de defesa foi apresentada."

        # --- CORREÇÃO NA CHAMADA DA FUNÇÃO ---
        circunstancias_obj = verifica_agravante_atenuante(
            historico_militar, 
            patd.transgressao, 
            justificativa, 
            patd.itens_enquadrados,
            comportamento_anterior # Argumento adicionado
        )
        # --- FIM DA CORREÇÃO NA CHAMADA ---

        circunstancias_dict = circunstancias_obj.item[0]
        patd.circunstancias = {
            'atenuantes': circunstancias_dict.get('atenuantes', []),
            'agravantes': circunstancias_dict.get('agravantes', [])
        }

        punicao_obj = sugere_punicao(
            transgressao=patd.transgressao,
            agravantes=patd.circunstancias.get('agravantes', []),
            atenuantes=patd.circunstancias.get('atenuantes', []),
            itens=patd.itens_enquadrados,
            observacao="Análise inicial"
        )
        patd.punicao_sugerida = punicao_obj.punicao.get('punicao', 'Erro na sugestão.')

        patd.save()

        final_response_data = {
            'status': 'success',
            'analise_data': {
                'itens': patd.itens_enquadrados,
                'circunstancias': patd.circunstancias,
                'punicao': patd.punicao_sugerida
            }
        }

        return JsonResponse(final_response_data)

    except Exception as e:
        logger.error(f"Erro na análise da IA para PATD {pk}: {e}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': f'Ocorreu um erro durante a análise da IA: {e}'
        }, status=500)

@login_required
@oficial_responsavel_required
@require_POST
def salvar_apuracao(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)
        data = json.loads(request.body)

        if patd.status == 'apuracao_preclusao':
            if not _check_preclusao_signatures(patd):
                return JsonResponse({
                    'status': 'error',
                    'message': 'A apuração não pode ser concluída. Faltam assinaturas das testemunhas no termo de preclusão.'
                }, status=400)

        patd.itens_enquadrados = data.get('itens_enquadrados')
        patd.circunstancias = data.get('circunstancias')
        punicao_sugerida_str = data.get('punicao_sugerida', '')
        patd.punicao_sugerida = punicao_sugerida_str

        match = re.search(r'(\d+)\s+dias\s+de\s+(.+)', punicao_sugerida_str, re.IGNORECASE)
        if match:
            dias_num = int(match.group(1))
            punicao_tipo = match.group(2).strip()
            dias_texto = num2words(dias_num, lang='pt_BR')
            patd.dias_punicao = f"{dias_texto} ({dias_num:02d}) dias"
            patd.punicao = punicao_tipo
        else:
            patd.dias_punicao = ""
            patd.punicao = punicao_sugerida_str

        patd.transgressao_afirmativa = f"foi verificado que o militar realmente cometeu a transgressão de '{patd.transgressao}'."

        if not patd.texto_relatorio:
            patd.texto_relatorio = texto_relatorio(patd.transgressao, patd.alegacao_defesa)

        patd.definir_natureza_transgressao()
        patd.calcular_e_atualizar_comportamento()

        patd.status = 'aguardando_punicao'

        patd.save()

        return JsonResponse({'status': 'success', 'message': 'Apuração salva com sucesso!'})

    except Exception as e:
        logger.error(f"Erro ao salvar apuração da PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@require_GET
def search_militares_json(request):
    """Retorna uma lista de militares para a pesquisa no modal."""
    query = request.GET.get('q', '')
    militares = Militar.objects.all()

    if query:
        militares = militares.filter(
            Q(nome_completo__icontains=query) |
            Q(nome_guerra__icontains=query) |
            Q(posto__icontains=query)
        )

    militares = militares.order_by('posto', 'nome_guerra')[:50]
    data = list(militares.values('id', 'posto', 'nome_guerra', 'nome_completo'))
    return JsonResponse(data, safe=False)

@method_decorator([login_required, comandante_required], name='dispatch')
class ComandanteDashboardView(ListView):
    model = PATD
    template_name = 'comandante_dashboard.html'
    context_object_name = 'patds'

    def get_queryset(self):
        return PATD.objects.filter(status='analise_comandante').select_related('militar').order_by('-data_inicio')

# --- ALTERAÇÃO: patd_aprovar modificado ---
@login_required
@comandante_required
@require_POST # Garante que só aceita POST
def patd_aprovar(request, pk):
    patd = get_object_or_404(PATD, pk=pk)
    form = ComandanteAprovarForm(request.POST)

    # Verifica se as testemunhas estão definidas (lógica existente)
    errors = []
    if not patd.testemunha1 or not patd.testemunha2:
        errors.append("É necessário definir as duas testemunhas no processo.")

    if errors:
        error_message = f"PATD Nº {patd.numero_patd}: Não foi possível aprovar. " + " ".join(errors)
        messages.error(request, error_message)
        return redirect(request.META.get('HTTP_REFERER', 'Ouvidoria:comandante_dashboard'))

    # Verifica o formulário e a senha
    if form.is_valid():
        senha = form.cleaned_data['senha_comandante']
        # Autentica o usuário logado (que deve ser o comandante) com a senha fornecida
        user = authenticate(username=request.user.username, password=senha)

        if user is not None:
            # Senha correta, prossegue com a aprovação
            patd.status = 'aguardando_assinatura_npd'
            patd.save()
            messages.success(request, f"PATD Nº {patd.numero_patd} aprovada com sucesso. Aguardando assinatura da NPD.")
            return redirect('Ouvidoria:comandante_dashboard')
        else:
            # Senha incorreta
            messages.error(request, "Senha do Comandante incorreta. Aprovação não realizada.")
    else:
        # Formulário inválido (deve ter faltado a senha)
        messages.error(request, "Erro no formulário. A senha é obrigatória.")

    # Redireciona de volta em caso de erro de senha ou formulário
    return redirect(request.META.get('HTTP_REFERER', 'Ouvidoria:comandante_dashboard'))


@login_required
@comandante_required
@require_POST
def patd_retornar(request, pk):
    patd = get_object_or_404(PATD, pk=pk)
    comentario = request.POST.get('comentario')

    if not comentario:
        messages.error(request, "O comentário é obrigatório para retornar a PATD.")
        return redirect(request.META.get('HTTP_REFERER', 'Ouvidoria:comandante_dashboard'))

    patd.status = 'aguardando_punicao_alterar'
    patd.comentario_comandante = comentario
    patd.save()
    messages.warning(request, f"PATD Nº {patd.numero_patd} retornada para alteração com observações.")
    # Não precisa de senha aqui
    return redirect(request.META.get('HTTP_REFERER', 'Ouvidoria:comandante_dashboard'))

@login_required
@oficial_responsavel_required
@require_POST
def avancar_para_comandante(request, pk):
    patd = get_object_or_404(PATD, pk=pk)

    if not patd.testemunha1 or not patd.testemunha2:
        detail_url = reverse('Ouvidoria:patd_detail', kwargs={'pk': pk})
        return redirect(f'{detail_url}?erro=testemunhas')

    patd.status = 'analise_comandante'
    patd.save()
    messages.success(request, f"PATD Nº {patd.numero_patd} enviada para análise do Comandante.")
    return redirect('Ouvidoria:patd_detail', pk=pk)

@login_required
@ouvidoria_required
@require_POST
def solicitar_reconsideracao(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)
        if patd.status != 'periodo_reconsideracao':
            return JsonResponse({'status': 'error', 'message': 'A PATD não está no período de reconsideração.'}, status=400)

        patd.status = 'em_reconsideracao'
        patd.save(update_fields=['status'])
        messages.success(request, f'PATD Nº {patd.numero_patd} movida para "Em Reconsideração".')
        return JsonResponse({'status': 'success', 'message': 'Status atualizado com sucesso.'})
    except Exception as e:
        logger.error(f"Erro ao solicitar reconsideração para PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
@require_POST
def salvar_reconsideracao(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)
        if patd.status != 'em_reconsideracao':
            return JsonResponse({'status': 'error', 'message': 'A PATD não está em fase de reconsideração.'}, status=400)

        texto = request.POST.get('texto_reconsideracao', '')
        arquivos = request.FILES.getlist('anexos_reconsideracao')

        if not texto and not arquivos:
            return JsonResponse({'status': 'error', 'message': 'É necessário fornecer um texto ou anexar pelo menos um ficheiro.'}, status=400)

        patd.texto_reconsideracao = texto
        if not patd.data_reconsideracao:
            patd.data_reconsideracao = timezone.now()
        patd.save(update_fields=['texto_reconsideracao', 'data_reconsideracao'])

        for arquivo in arquivos:
            Anexo.objects.create(patd=patd, arquivo=arquivo, tipo='reconsideracao')

        _check_and_advance_reconsideracao_status(pk)

        return JsonResponse({'status': 'success', 'message': 'Pedido de reconsideração e anexos salvos com sucesso.'})
    except Exception as e:
        logger.error(f"Erro ao salvar texto de reconsideração para PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@require_GET
def comandante_pendencias_json(request):
    if not has_comandante_access(request.user):
        return JsonResponse({'count': 0})

    count = PATD.objects.filter(status='analise_comandante').count()
    return JsonResponse({'count': count})

@login_required
@ouvidoria_required
@require_POST
def excluir_anexo(request, pk):
    try:
        anexo = get_object_or_404(Anexo, pk=pk)

        patd = anexo.patd
        user_militar = request.user.profile.militar if hasattr(request.user, 'profile') else None

        if not (request.user.is_superuser or user_militar == patd.oficial_responsavel):
             return JsonResponse({'status': 'error', 'message': 'Você não tem permissão para excluir este anexo.'}, status=403)

        if anexo.arquivo and os.path.isfile(anexo.arquivo.path):
            os.remove(anexo.arquivo.path)

        anexo.delete()

        return JsonResponse({'status': 'success', 'message': 'Anexo excluído com sucesso.'})
    except Exception as e:
        logger.error(f"Erro ao excluir anexo {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@oficial_responsavel_required # Alterado para oficial_responsavel_required
@require_POST
def finalizar_publicacao(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)

        # Apenas quem tem acesso à Ouvidoria pode finalizar
        if not has_ouvidoria_access(request.user):
             messages.error(request, "Você não tem permissão para finalizar este processo.")
             return redirect('Ouvidoria:patd_detail', pk=pk)

        boletim = request.POST.get('boletim_publicacao')

        if not boletim or not boletim.strip():
            messages.error(request, "O número do boletim é obrigatório para finalizar o processo.")
            return redirect('Ouvidoria:patd_detail', pk=pk)

        patd.boletim_publicacao = boletim
        patd.status = 'finalizado'
        patd.data_termino = timezone.now()
        patd.save()

        messages.success(request, f"PATD Nº {patd.numero_patd} finalizada com sucesso e publicada no boletim {boletim}.")
        return redirect('Ouvidoria:patd_detail', pk=pk)
    except Exception as e:
        logger.error(f"Erro ao finalizar PATD {pk}: {e}")
        messages.error(request, "Ocorreu um erro ao tentar finalizar o processo.")
        return redirect('Ouvidoria:patd_detail', pk=pk)

@login_required
@oficial_responsavel_required
@require_POST
def justificar_patd(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)

        if patd.status not in ['em_apuracao', 'apuracao_preclusao']:
             return JsonResponse({'status': 'error', 'message': 'A PATD não está na fase correta para ser justificada.'}, status=400)

        patd.justificado = True
        patd.status = 'aguardando_publicacao'

        patd.punicao_sugerida = "Transgressão Justificada"
        patd.punicao = ""
        patd.dias_punicao = ""

        patd.texto_relatorio = """Após análise dos fatos, alegações e circunstâncias, este Oficial Apurador conclui que a transgressão disciplinar imputada ao militar está JUSTIFICADA, nos termos do Art. 13, item 1 do RDAER."""

        patd.save()
        return JsonResponse({'status': 'success', 'message': 'A transgressão foi justificada com sucesso.'})
    except Exception as e:
        logger.error(f"Erro ao justificar a PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@oficial_responsavel_required
@require_POST
def anexar_documento_reconsideracao_oficial(request, pk):
    patd = get_object_or_404(PATD, pk=pk)
    if patd.status != 'aguardando_comandante_base':
        messages.error(request, "Ação não permitida no status atual.")
        return redirect('Ouvidoria:patd_detail', pk=pk)

    anexo_file = request.FILES.get('anexo_oficial')
    if not anexo_file:
        messages.error(request, "Nenhum ficheiro foi enviado.")
        return redirect('Ouvidoria:patd_detail', pk=pk)

    Anexo.objects.create(patd=patd, arquivo=anexo_file, tipo='reconsideracao_oficial')

    patd.status = 'aguardando_preenchimento_npd_reconsideracao'
    patd.save()

    messages.success(request, "Documento anexado com sucesso. O processo aguarda o preenchimento da NPD de Reconsideração.")
    return redirect('Ouvidoria:patd_detail', pk=pk)

@login_required
@oficial_responsavel_required
@require_POST
def salvar_nova_punicao(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)
        data = json.loads(request.body)
        dias = int(data.get('dias'))
        tipo = data.get('tipo')

        if dias < 0 or not tipo:
            return JsonResponse({'status': 'error', 'message': 'Dados inválidos.'}, status=400)

        dias_texto = num2words(dias, lang='pt_BR')

        patd.nova_punicao_dias = f"{dias_texto} ({dias:02d}) dias"
        patd.nova_punicao_tipo = tipo

        patd.dias_punicao = f"{dias_texto} ({dias:02d}) dias"
        patd.punicao = tipo

        patd.definir_natureza_transgressao()
        patd.calcular_e_atualizar_comportamento()

        patd.status = 'aguardando_publicacao'
        patd.save()

        return JsonResponse({'status': 'success', 'message': 'Nova punição salva com sucesso.'})
    except Exception as e:
        logger.error(f"Erro ao salvar nova punição para PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
@ouvidoria_required
def exportar_patd_docx(request, pk):
    """
    Gera e serve um ficheiro DOCX a partir do conteúdo HTML da PATD,
    incluindo imagens e formatação correta.
    """
    patd = get_object_or_404(PATD, pk=pk)
    context = _get_document_context(patd)
    config = Configuracao.load()
    comandante_gsd = config.comandante_gsd

    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.15)
    section.right_margin = Cm(2.5)
    section.gutter = Cm(0)

    full_html_content = "".join(get_document_pages(patd))
    soup = BeautifulSoup(full_html_content, 'html.parser')

    militar_sig_counter = 0

    placeholder_regex = re.compile(r'({[^}]+})')

    for element in soup.find_all(['p', 'img']):
        if element.name == 'p':
            p = document.add_paragraph()

            p_format = p.paragraph_format
            p_format.left_indent = Inches(0)
            p_format.first_line_indent = Inches(0)
            p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)

            if element.has_attr('style'):
                if 'text-align: center' in element['style']:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'text-align: right' in element['style']:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif 'text-align: justify' in element['style']:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            for content in element.contents:
                if isinstance(content, NavigableString):
                    text_content = str(content)
                    parts = placeholder_regex.split(text_content)

                    for part in parts:
                        if not part: continue

                        is_image_placeholder = False
                        if placeholder_regex.match(part):
                            placeholder = part.strip()
                            try:
                                if placeholder == '{Assinatura_Imagem_Comandante_GSD}' and comandante_gsd and comandante_gsd.assinatura:
                                    _, img_str = comandante_gsd.assinatura.split(';base64,')
                                    p.add_run().add_picture(io.BytesIO(base64.b64decode(img_str)), height=Cm(1.5))
                                    is_image_placeholder = True

                                elif placeholder == '{Assinatura_Imagem_Oficial_Apurador}' and patd.assinatura_oficial and patd.assinatura_oficial.path and os.path.exists(patd.assinatura_oficial.path):
                                    p.add_run().add_picture(patd.assinatura_oficial.path, height=Cm(1.5))
                                    is_image_placeholder = True

                                elif placeholder == '{Assinatura_Imagem_Testemunha_1}' and patd.assinatura_testemunha1 and patd.assinatura_testemunha1.path and os.path.exists(patd.assinatura_testemunha1.path):
                                    p.add_run().add_picture(patd.assinatura_testemunha1.path, height=Cm(1.5))
                                    is_image_placeholder = True

                                elif placeholder == '{Assinatura_Imagem_Testemunha_2}' and patd.assinatura_testemunha2 and patd.assinatura_testemunha2.path and os.path.exists(patd.assinatura_testemunha2.path):
                                    p.add_run().add_picture(patd.assinatura_testemunha2.path, height=Cm(1.5))
                                    is_image_placeholder = True

                                elif placeholder == '{Assinatura Militar Arrolado}':
                                    assinaturas_arrolado = patd.assinaturas_militar or []
                                    if militar_sig_counter < len(assinaturas_arrolado) and assinaturas_arrolado[militar_sig_counter]:
                                        anexo_url = assinaturas_arrolado[militar_sig_counter]
                                        # Assume que a URL é relativa à MEDIA_URL
                                        anexo_path = os.path.join(settings.MEDIA_ROOT, anexo_url.replace(settings.MEDIA_URL, '', 1))
                                        if os.path.exists(anexo_path):
                                            p.add_run().add_picture(anexo_path, height=Cm(1.5))
                                        militar_sig_counter += 1
                                        is_image_placeholder = True
                                    else:
                                         # Se não houver assinatura ou for None, incrementa o contador mesmo assim
                                         militar_sig_counter += 1


                                elif placeholder == '{Brasao da Republica}':
                                    img_path = os.path.join(settings.BASE_DIR, 'Static', 'img', 'brasao.png')
                                    if os.path.exists(img_path):
                                        p.add_run().add_picture(img_path, width=Cm(3))
                                        is_image_placeholder = True

                            except Exception as e:
                                logger.error(f"Error processing image placeholder {placeholder}: {e}")
                                is_image_placeholder = False # Não adiciona texto se era para ser imagem

                        if not is_image_placeholder:
                            sub_parts = re.split(r'(\*\*.*?\*\*)', part)
                            for sub_part in sub_parts:
                                if sub_part.startswith('**') and sub_part.endswith('**'):
                                    p.add_run(sub_part.strip('*')).bold = True
                                else:
                                    p.add_run(sub_part)

                elif content.name == 'strong':
                    p.add_run(content.get_text()).bold = True

        elif element.name == 'img':
             if 'brasao.png' in element.get('src', ''):
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()

                img_path = os.path.join(settings.BASE_DIR, 'Static', 'img', 'brasao.png')
                if os.path.exists(img_path):
                    run.add_picture(img_path, width=Cm(3))

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=PATD_{patd.numero_patd}.docx'
    document.save(response)

    return response