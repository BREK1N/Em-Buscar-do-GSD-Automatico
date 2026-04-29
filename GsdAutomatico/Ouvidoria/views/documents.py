import io, json, os, re, logging, base64, uuid, traceback
import tempfile, subprocess
from datetime import datetime, timedelta

from django.conf import settings
from django.utils import timezone

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_POST
from django.core.files.base import ContentFile
from django.core.files import File
from django.contrib.staticfiles.storage import staticfiles_storage
from django.contrib.staticfiles import finders
import docx
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.text import WD_BREAK
from bs4 import BeautifulSoup, NavigableString
import fitz

from ..models import PATD, Configuracao, Anexo
from Secao_pessoal.models import Efetivo
from .decorators import ouvidoria_required, oficial_responsavel_required, comandante_redirect
from .helpers import (
    _get_document_context, _render_document_from_template,
    get_document_pages, format_militar_string,
    _try_advance_status_from_justificativa,
)
from ..analise_transgressao import analisar_e_resumir_defesa, reescrever_ocorrencia

logger = logging.getLogger(__name__)

@login_required
@ouvidoria_required
@require_POST
def salvar_alegacao_defesa(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)

        if patd.alegacao_defesa or patd.anexos.filter(tipo='defesa').exists():
            return JsonResponse({
                'status': 'error',
                'message': 'A alegação de defesa já foi enviada e não pode ser alterada.'
            }, status=403)

        alegacao_texto = request.POST.get('alegacao_defesa', '')
        arquivos = request.FILES.getlist('anexos_defesa')

        if not alegacao_texto and not arquivos:
            return JsonResponse({'status': 'error', 'message': 'É necessário fornecer um texto ou anexar pelo menos um ficheiro.'}, status=400)

        patd.alegacao_defesa = alegacao_texto
        patd.data_alegacao = timezone.now()

        for arquivo in arquivos:
            Anexo.objects.create(patd=patd, arquivo=arquivo, tipo='defesa')

        try:
            if not patd.alegacao_defesa_resumo:
                patd.alegacao_defesa_resumo = analisar_e_resumir_defesa(patd.alegacao_defesa)
            if not patd.ocorrencia_reescrita:
                ocorrencia_formatada = reescrever_ocorrencia(patd.transgressao)
                patd.ocorrencia_reescrita = ocorrencia_formatada
                patd.comprovante = ocorrencia_formatada
        except Exception as e:
            logger.error(f"Erro ao chamar a IA para processar textos da PATD {pk}: {e}")
            if not patd.alegacao_defesa_resumo:
                patd.alegacao_defesa_resumo = "Erro ao gerar resumo."
            if not patd.ocorrencia_reescrita:
                patd.ocorrencia_reescrita = patd.transgressao
                patd.comprovante = patd.transgressao

        patd.save()
        _try_advance_status_from_justificativa(patd)
        patd.save()

        return JsonResponse({'status': 'success', 'message': 'Alegação de defesa e anexos salvos com sucesso.'})
    except Exception as e:
        logger.error(f"Erro ao salvar alegação de defesa da PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': 'Ocorreu um erro interno.'}, status=500)


@login_required
@ouvidoria_required
@require_POST
def extender_prazo(request, pk):
    try:
        patd = get_object_or_404(PATD, pk=pk)
        if patd.status != 'prazo_expirado':
            return JsonResponse({'status': 'error', 'message': 'O prazo só pode ser estendido se estiver expirado.'}, status=400)

        try:
            data = json.loads(request.body)
        except (json.JSONDecodeError, ValueError):
            return JsonResponse({'status': 'error', 'message': 'JSON inválido.'}, status=400)
        dias_extensao = int(data.get('dias', 0))
        minutos_extensao = int(data.get('minutos', 0))

        if not request.user.is_superuser and minutos_extensao != 0:
            return JsonResponse({'status': 'error', 'message': 'Apenas administradores podem alterar os minutos.'}, status=403)


        if dias_extensao < 0 or minutos_extensao < 0:
            return JsonResponse({'status': 'error', 'message': 'Valores de extensão de prazo inválidos.'}, status=400)

        config = Configuracao.load()

        delta_dias = config.prazo_defesa_dias - dias_extensao
        delta_minutos = config.prazo_defesa_minutos - minutos_extensao

        patd.data_ciencia = timezone.now() - timedelta(days=delta_dias, minutes=delta_minutos)
        patd.status = 'aguardando_justificativa'
        patd.save()

        return JsonResponse({'status': 'success', 'message': 'Prazo estendido com sucesso.'})

    except (ValueError, TypeError):
        return JsonResponse({'status': 'error', 'message': 'Dados de entrada inválidos.'}, status=400)
    except Exception as e:
        logger.error(f"Erro ao estender prazo da PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': 'Ocorreu um erro interno.'}, status=500)


@login_required
@ouvidoria_required
@require_POST
def salvar_documento_patd(request, pk):
    patd = get_object_or_404(PATD, pk=pk)

    # --- BLOQUEIO DE SEGURANÇA ---
    # Este bloco impede qualquer edição (Texto, Datas ou Localidade) se estiver finalizado.
    if patd.status == 'finalizado':
        return JsonResponse({
            'status': 'error', 
            'message': 'Este processo está finalizado e não permite mais edições no documento.'
        }, status=403)
        
    try:
        try:
            data = json.loads(request.body)
        except (json.JSONDecodeError, ValueError):
            return JsonResponse({'status': 'error', 'message': 'JSON inválido.'}, status=400)

        # Obtém os dados
        texto_documento = data.get('texto_documento')
        dates = data.get('dates', {})
        texts = data.get('texts', {})

        # Validação básica
        if texto_documento is None:
            return JsonResponse({'status': 'error', 'message': 'Nenhum texto recebido.'}, status=400)

        # 1. Atualiza o texto do documento
        patd.documento_texto = texto_documento
        
        # 2. Atualiza a Localidade (Manipulação segura de JSONField)
        if 'localidade' in texts:
            # Garante que é um dicionário antes de editar
            circunstancias = patd.circunstancias if isinstance(patd.circunstancias, dict) else {}
            circunstancias['localidade'] = texts['localidade']
            patd.circunstancias = circunstancias # Reatribuição força o Django a reconhecer a mudança

        # 3. Atualiza as Datas dinamicamente
        for field_name, date_str in dates.items():
            # Só atualiza se o campo existir no modelo PATD para evitar erros
            if hasattr(patd, field_name):
                if date_str:
                    try:
                        date_obj = datetime.strptime(date_str, '%Y-%m-%d').date()
                        setattr(patd, field_name, date_obj)
                    except (ValueError, TypeError):
                        logger.warning(f"Formato de data inválido para o campo {field_name}: {date_str}")
                else:
                    # Se vier vazio, define como None (null no banco)
                    setattr(patd, field_name, None)

        patd.save()

        return JsonResponse({'status': 'success', 'message': 'Documento, datas e localidade salvos com sucesso.'})
        
    except json.JSONDecodeError:
        return JsonResponse({'status': 'error', 'message': 'JSON inválido.'}, status=400)
    except Exception as e:
        logger.error(f"Erro ao salvar documento da PATD {pk}: {e}")
        return JsonResponse({'status': 'error', 'message': 'Erro interno ao salvar.'}, status=500)


@login_required
@comandante_redirect
@ouvidoria_required
@require_POST
def upload_ficha_individual(request, pk):
    patd = get_object_or_404(PATD, pk=pk)

    if patd.status == 'finalizado':
        messages.error(request, "Não é possível anexar arquivos em processos finalizados.")
        return redirect('Ouvidoria:patd_detail', pk=pk)
    
    if 'ficha_individual' not in request.FILES:
        messages.error(request, "Nenhum arquivo enviado.")
        return redirect('Ouvidoria:patd_detail', pk=pk)

    ficha_individual_file = request.FILES['ficha_individual']

    # Remove existing ficha_individual if it exists
    patd.anexos.filter(tipo='ficha_individual').delete()

    # Create a new Anexo for the ficha_individual
    Anexo.objects.create(
        patd=patd,
        arquivo=ficha_individual_file,
        tipo='ficha_individual'
    )

    messages.success(request, "Ficha individual atualizada com sucesso.")
    return redirect('Ouvidoria:patd_detail', pk=pk)


@login_required
@ouvidoria_required
def upload_formulario_resumo(request, pk):
    patd = get_object_or_404(PATD, pk=pk)
    if patd.status == 'finalizado':
        messages.error(request, "Não é possível anexar arquivos em processos finalizados.")
        return redirect('Ouvidoria:patd_detail', pk=pk)
    if 'formulario_resumo' not in request.FILES:
        messages.error(request, "Nenhum arquivo enviado.")
        return redirect('Ouvidoria:patd_detail', pk=pk)
    patd.anexos.filter(tipo='formulario_resumo').delete()
    Anexo.objects.create(patd=patd, arquivo=request.FILES['formulario_resumo'], tipo='formulario_resumo')
    messages.success(request, "Formulário de resumo anexado com sucesso.")
    return redirect('Ouvidoria:patd_detail', pk=pk)


def _append_anexo_content(document, anexo):
    """
    Função auxiliar para adicionar o conteúdo de um anexo (PDF, DOCX, Imagem)
    a um documento docx existente, sem adicionar uma quebra de página inicial.
    """
    file_path = anexo.arquivo.path
    file_name = os.path.basename(file_path)
    ext = os.path.splitext(file_name)[1].lower()
    
    try:
        if not os.path.exists(file_path):
            document.add_paragraph(f"[Erro: Ficheiro anexo '{file_name}' não encontrado no servidor.]")
            return

        if ext in ['.png', '.jpg', '.jpeg']:
            # Adiciona imagem, mantendo a proporção com largura máxima de 6 polegadas
            document.add_picture(file_path, width=Inches(6.5))
        
        elif ext == '.docx':
            # Anexa o conteúdo do DOCX na posição correta do body (antes do sectPr).
            # body.append() colocaria os elementos DEPOIS do <w:sectPr>, tornando o XML
            # inválido e fazendo o Word mover o conteúdo para o fim do documento.
            import copy
            from docx.oxml.ns import qn as _qn
            sub_doc = docx.Document(file_path)
            body = document.element.body
            # sectPr deve sempre ser o último filho do body
            sect_pr = body.find(_qn('w:sectPr'))
            for element in list(sub_doc.element.body):
                tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
                if tag == 'sectPr':
                    continue  # Não copia a seção do sub-doc
                node = copy.deepcopy(element)
                # Remove sectPr inline em parágrafos (quebras de seção embutidas)
                # que criariam novas seções sem rodapé
                for pPr in node.findall('.//' + _qn('w:pPr')):
                    inline_sect = pPr.find(_qn('w:sectPr'))
                    if inline_sect is not None:
                        pPr.remove(inline_sect)
                if sect_pr is not None:
                    body.insert(list(body).index(sect_pr), node)
                else:
                    body.append(node)
        
        elif ext == '.pdf':
            try:
                pdf_doc = fitz.open(file_path)
                # Área útil do documento exportado (A4 retrato com as margens definidas)
                # left=2.15cm, right=2.5cm → largura útil ≈ 21 - 2.15 - 2.5 = 16.35 cm
                # top=1.5cm, bottom=2.54cm → altura útil ≈ 29.7 - 1.5 - 2.54 = 25.66 cm
                max_w_cm = 16.3
                max_h_cm = 24.0

                # Pré-filtra páginas em branco do PDF antes de inserir no DOCX.
                # Critério combinado: brancura visual >= 99% E texto < 50 chars.
                # Isso remove páginas com artefato "tsteteetetlçask" (15 chars, ~100% branco)
                # sem remover páginas reais (que têm <99% brancura OU 116+ chars de texto).
                non_blank_pages = []
                for page in pdf_doc:
                    check_pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5), colorspace=fitz.csGRAY)
                    arr = bytearray(check_pix.samples)
                    total = len(arr)
                    white_ratio = sum(1 for b in arr if b >= 250) / total if total > 0 else 0
                    if white_ratio >= 0.999:
                        page_text = page.get_text().strip()
                        if len(page_text) < 50:
                            continue  # Página em branco ou com apenas artefato de texto → pula
                    rect = page.rect
                    is_landscape = rect.width > rect.height
                    if is_landscape:
                        mat = fitz.Matrix(250 / 72, 250 / 72).prerotate(90)
                    else:
                        mat = fitz.Matrix(250 / 72, 250 / 72)
                    pix = page.get_pixmap(matrix=mat)
                    non_blank_pages.append((pix.tobytes("png"), is_landscape))

                for i, (img_bytes, is_landscape) in enumerate(non_blank_pages):
                    if i > 0:
                        last_p = document.paragraphs[-1] if document.paragraphs else document.add_paragraph()
                        last_p.add_run().add_break(WD_BREAK.PAGE)
                    pic_p = document.add_paragraph()
                    pic_p.paragraph_format.space_before = Pt(0)
                    pic_p.paragraph_format.space_after = Pt(0)
                    if is_landscape:
                        pic_p.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(max_w_cm))
                    else:
                        pic_p.add_run().add_picture(io.BytesIO(img_bytes), height=Cm(max_h_cm))

                pdf_doc.close()
            except Exception as e:
                logger.error(f"Erro ao converter PDF para imagem {file_name}: {e}")
                document.add_paragraph(f"[Erro ao processar o anexo PDF '{file_name}'.]")

        else:
            # Tipo de ficheiro não suportado para embutir
            document.add_paragraph(f"[Conteúdo do anexo '{file_name}' (tipo: {ext}) não suportado para inclusão direta no DOCX.]")

    except Exception as e:
        logger.error(f"Erro geral ao processar anexo {file_name} para DOCX: {e}")
        document.add_paragraph(f"[Erro ao processar anexo '{file_name}': {e}]")


def add_page_number(paragraph):
    """
    Adiciona um campo de número de página a um parágrafo no rodapé.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Adiciona o texto "Página "
    run = paragraph.add_run()
    run.add_text('Página ')

    # --- INÍCIO DA CORREÇÃO ---
    # Cria o elemento fldChar para 'begin'
    fldChar_begin = docx.oxml.shared.OxmlElement('w:fldChar')
    fldChar_begin.set(docx.oxml.shared.qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    # Adiciona o código de instrução do campo (PAGE)
    instrText = docx.oxml.shared.OxmlElement('w:instrText')
    instrText.set(docx.oxml.shared.qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)

    # Cria o elemento fldChar para 'end'
    fldChar_end = docx.oxml.shared.OxmlElement('w:fldChar')
    fldChar_end.set(docx.oxml.shared.qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)


def _append_alegacao_docx(document, patd, context):
    """Lê o template da alegação, substitui os placeholders e o anexa ao documento principal."""
    # --- INÍCIO DA CORREÇÃO ---
    # Reutiliza o mesmo regex da função principal para consistência
    placeholder_regex = re.compile(r'({[^}]+})')
    # --- FIM DA CORREÇÃO ---
    try:
        pb_para = document.add_paragraph()
        pb_para.paragraph_format.space_before = Pt(0)
        pb_para.paragraph_format.space_after = Pt(0)
        pb_para.add_run().add_break(WD_BREAK.PAGE)

        doc_path = os.path.join(settings.BASE_DIR, 'pdf', 'PATD_Alegacao_DF.docx')
        alegacao_doc = Document(doc_path)
        for source_p in alegacao_doc.paragraphs:
            if source_p.text.strip():
                new_p = document.add_paragraph()
                new_p.paragraph_format.alignment = source_p.paragraph_format.alignment

                text_to_process = source_p.text
                
                for placeholder, value in context.items():
                    text_to_process = text_to_process.replace(str(placeholder), str(value))

                # --- INÍCIO DA CORREÇÃO: Reutilizar lógica de processamento de placeholders e formatação ---
                # Divide o texto em partes de texto normal e placeholders
                # O regex agora inclui a formatação de negrito
                parts = re.split(r'(\*\*.*?\*\*|{[^}]+})', text_to_process)

                for part in parts:
                    if not part: continue

                    # Lógica para placeholders de imagem (assinaturas)
                    if placeholder_regex.match(part):
                        placeholder = part.strip()
                        is_image_placeholder = False
                        try:
                            if placeholder == '{Assinatura Alegacao Defesa}' and patd.assinatura_alegacao_defesa and patd.assinatura_alegacao_defesa.path and os.path.exists(patd.assinatura_alegacao_defesa.path):
                                new_p.add_run().add_picture(patd.assinatura_alegacao_defesa.path, height=Cm(2.5))
                                is_image_placeholder = True
                        except Exception as e:
                            logger.error(f"Erro ao processar placeholder de imagem na alegação: {placeholder}: {e}")
                        
                        if is_image_placeholder:
                            continue # Pula para a próxima parte se a imagem foi adicionada

                    # Lógica para o texto da alegação
                    if '{Alegação de defesa}' in part:
                        alegacao_text = patd.alegacao_defesa or "[ALEGAÇÃO NÃO FORNECIDA]"
                        lines = alegacao_text.splitlines()
                        for i, line in enumerate(lines):
                            if i > 0: new_p.add_run().add_break() # Adiciona quebra de linha
                            new_p.add_run(line)
                    # Lógica para texto em negrito
                    elif part.startswith('**') and part.endswith('**'):
                        new_p.add_run(part.strip('*')).bold = True
                    # Ignora outros placeholders que não são imagens
                    elif placeholder_regex.match(part):
                        continue
                    # Adiciona texto normal
                    else:
                        new_p.add_run(part)
                # --- FIM DA CORREÇÃO ---
    except Exception as e:
        logger.error(f"Erro ao anexar documento de alegação: {e}")
        document.add_paragraph(f"[ERRO AO PROCESSAR DOCUMENTO DE ALEGAÇÃO: {e}]")


def _remove_blank_pages_from_docx(docx_bytes, filename_base):
    """
    Converte o DOCX para PDF via LibreOffice, remove páginas completamente em branco
    usando PyMuPDF e retorna os bytes do PDF limpo.
    Retorna (pdf_bytes, True) se bem-sucedido, ou (None, False) em caso de falha.
    """
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, f'{filename_base}.docx')
            with open(docx_path, 'wb') as f:
                f.write(docx_bytes)

            # Converte DOCX → PDF com LibreOffice headless
            env = os.environ.copy()
            env.setdefault('HOME', '/tmp')
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', tmpdir, docx_path],
                capture_output=True, text=True, timeout=120, env=env
            )
            if result.returncode != 0:
                logger.error(f"LibreOffice conversion failed: {result.stderr}")
                return None, False

            pdf_path = os.path.join(tmpdir, f'{filename_base}.pdf')
            if not os.path.exists(pdf_path):
                logger.error(f"PDF output not found after LibreOffice conversion")
                return None, False

            # Abre o PDF e detecta páginas em branco renderizando cada uma como imagem
            # e verificando se é visualmente toda branca (>99.5% pixels brancos)
            src_pdf = fitz.open(pdf_path)
            blank_page_indices = set()
            for i, page in enumerate(src_pdf):
                # Renderiza em baixa resolução só para detecção (rápido)
                pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5), colorspace=fitz.csGRAY)
                samples = pix.samples  # bytes em escala de cinza
                total_pixels = len(samples)
                if total_pixels == 0:
                    blank_page_indices.add(i)
                    continue
                # Conta pixels brancos (valor >= 250 em escala de cinza 0-255)
                # bytearray permite soma rápida sem loop Python
                arr = bytearray(samples)
                white_pixels = sum(1 for b in arr if b >= 250)
                white_ratio = white_pixels / total_pixels
                if white_ratio >= 0.980:
                    blank_page_indices.add(i)

            if not blank_page_indices:
                # Nenhuma página em branco — retorna o PDF como está
                pdf_bytes = open(pdf_path, 'rb').read()
                src_pdf.close()
                return pdf_bytes, True

            # Remove as páginas em branco criando um novo PDF apenas com as páginas válidas
            clean_pdf = fitz.open()
            for i in range(len(src_pdf)):
                if i not in blank_page_indices:
                    clean_pdf.insert_pdf(src_pdf, from_page=i, to_page=i)

            clean_bytes = clean_pdf.tobytes(deflate=True)
            clean_pdf.close()
            src_pdf.close()
            logger.info(f"Removidas {len(blank_page_indices)} página(s) em branco do PDF exportado.")
            return clean_bytes, True

    except subprocess.TimeoutExpired:
        logger.error("LibreOffice conversion timed out")
        return None, False
    except Exception as e:
        logger.error(f"Erro ao remover páginas em branco: {e}")
        return None, False


def _propagate_footer_to_all_sections(document):
    """
    Garante que todas as seções do documento (além da primeira) herdem o rodapé
    da seção principal. Isso evita que sub-documentos inseridos que possuam
    quebras de seção removam o número de página do rodapé.
    """
    from docx.oxml.ns import qn as _qn
    import copy as _copy

    body = document.element.body
    main_sect_pr = body.find(_qn('w:sectPr'))
    if main_sect_pr is None:
        return

    # Obtém os elementos de rodapé da seção principal para replicar
    main_footer_refs = main_sect_pr.findall(_qn('w:footerReference'))

    # Procura por sectPr inline em parágrafos (novas seções criadas por sub-docs)
    for pPr in body.findall('.//' + _qn('w:pPr')):
        inline_sect = pPr.find(_qn('w:sectPr'))
        if inline_sect is None:
            continue
        # Remove referências de rodapé existentes nessa seção
        for old_ref in inline_sect.findall(_qn('w:footerReference')):
            inline_sect.remove(old_ref)
        # Copia as referências de rodapé da seção principal
        for ref in main_footer_refs:
            inline_sect.insert(0, _copy.deepcopy(ref))


def preview_patd_pdf(request, pk):
    request.GET = request.GET.copy()
    request.GET['preview'] = '1'
    return exportar_patd_docx(request, pk)


@login_required
@ouvidoria_required
def exportar_patd_pdf(request, pk):
    """Gera PDF idêntico ao visualizador usando WeasyPrint com todas as imagens em base64."""
    import re as _re
    from django.conf import settings as _settings
    from django.contrib.staticfiles import finders as _finders

    patd = get_object_or_404(PATD, pk=pk)
    document_pages_raw = get_document_pages(patd, for_docx=True)

    PB = '<div class="manual-page-break"></div>'

    # ── Helpers base64 ────────────────────────────────────────────────────────

    def _file_to_b64(path):
        if not path or not os.path.exists(path):
            return None
        ext = os.path.splitext(path)[1].lower().lstrip('.')
        mime = {'png':'image/png','jpg':'image/jpeg','jpeg':'image/jpeg',
                'gif':'image/gif','webp':'image/webp','bmp':'image/bmp',
                'svg':'image/svg+xml'}.get(ext, 'image/png')
        with open(path, 'rb') as f:
            return f'data:{mime};base64,{base64.b64encode(f.read()).decode()}'

    def _field_to_b64(field):
        if not field or not field.name:
            return None
        try:
            return _file_to_b64(field.path)
        except Exception:
            return None

    def _media_path(media_url):
        rel = media_url.lstrip('/')
        media_prefix = _settings.MEDIA_URL.lstrip('/')
        if rel.startswith(media_prefix):
            rel = rel[len(media_prefix):]
        return os.path.join(_settings.MEDIA_ROOT, rel)

    def _static_path(src):
        """Resolve qualquer URL /Static/... ou /static/... para path no disco."""
        filename = src.rstrip('/').split('/')[-1].split('?')[0]
        path = _finders.find(f'img/{filename}') or _finders.find(filename)
        if not path:
            for base in [_settings.STATIC_ROOT, getattr(_settings, 'STATICFILES_DIRS', [None])[0] or '']:
                for candidate in [os.path.join(base, 'img', filename), os.path.join(base, filename)]:
                    if os.path.exists(candidate):
                        return candidate
        return path

    def _pdf_to_imgs(path):
        """Converte PDF em lista de <img> base64 via fitz ajustado à página de destino."""
        imgs = []
        try:
            # Dimensões de destino em pontos (1 cm = 28.346 pt)
            dest_w_pt = _meta['w'] * 28.346
            dest_h_pt = _meta['h'] * 28.346
            doc = fitz.open(path)
            for page in doc:
                pr = page.rect
                src_w, src_h = pr.width, pr.height
                # Calcula escala para caber inteiro na página de destino (sem cortar)
                scale = min(dest_w_pt / src_w, dest_h_pt / src_h) if src_w and src_h else 1.0
                mat = fitz.Matrix(scale, scale)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                b64 = base64.b64encode(pix.tobytes('png')).decode()
                imgs.append(f'data:image/png;base64,{b64}')
            doc.close()
        except Exception as e:
            logger.error("fitz falhou em %s: %s", path, e)
        return imgs

    # ── Assinaturas ───────────────────────────────────────────────────────────

    SIG_STYLE = 'height:90px;display:block;margin:4px auto;'  # centralizado + maior

    def _sig_img(field):
        b64 = _field_to_b64(field)
        return (f'<p style="text-align:center;margin:0;"><img src="{b64}" style="{SIG_STYLE}"/></p>'
                if b64 else '<span style="color:#888;font-style:italic;">[Sem assinatura]</span>')

    config = Configuracao.load()
    cmd_sig_html = '<span style="color:#888;font-style:italic;">[Sem assinatura]</span>'
    # Prioridade: assinatura salva no despacho → assinatura padrão do comandante GSD
    if patd.assinatura_cmd_gsd_despacho:
        b64 = _field_to_b64(patd.assinatura_cmd_gsd_despacho)
        if b64:
            cmd_sig_html = f'<p style="text-align:center;margin:0;"><img src="{b64}" style="{SIG_STYLE}"/></p>'
    elif config.comandante_gsd:
        cmd_assinatura = getattr(config.comandante_gsd, 'assinatura', None)
        if cmd_assinatura:
            if isinstance(cmd_assinatura, str) and cmd_assinatura.startswith('data:'):
                cmd_sig_html = f'<p style="text-align:center;margin:0;"><img src="{cmd_assinatura}" style="{SIG_STYLE}"/></p>'
            elif hasattr(cmd_assinatura, 'path'):
                b64 = _field_to_b64(cmd_assinatura)
                if b64:
                    cmd_sig_html = f'<p style="text-align:center;margin:0;"><img src="{b64}" style="{SIG_STYLE}"/></p>'

    assinaturas_militar = patd.assinaturas_militar or []
    mil_idx = [0]

    def _mil_sig(_m):
        i = mil_idx[0]; mil_idx[0] += 1
        if i < len(assinaturas_militar) and assinaturas_militar[i]:
            b64 = _file_to_b64(_media_path(assinaturas_militar[i]))
            if b64:
                return f'<p style="text-align:center;margin:0;"><img src="{b64}" style="{SIG_STYLE}"/></p>'
        return '<span style="color:#888;font-style:italic;">[Sem assinatura]</span>'

    _sig_cache = {
        'defesa':      _sig_img(patd.assinatura_alegacao_defesa),
        'recon':       _sig_img(patd.assinatura_reconsideracao),
        'oficial':     _sig_img(patd.assinatura_oficial),
        'test1':       _sig_img(patd.assinatura_testemunha1),
        'test2':       _sig_img(patd.assinatura_testemunha2),
    }

    # ── Substituição universal de src (aspas simples e duplas) ───────────────

    def _inline_all_srcs(html):
        """Converte todos os src= para base64 ou fitz, independente de aspas.
        Tags <embed>/<iframe> com PDF são substituídas inteiramente pelas imagens."""

        # 1. Substitui tags <embed> e <iframe> com PDF src pela sequência de imagens
        def _replace_embed_tag(m):
            tag_html = m.group(0)
            src_m = _re.search(r"""src=['"]([^'"]+)['"]""", tag_html)
            if not src_m:
                return ''
            src = src_m.group(1)
            if '/media' in src and src.lower().endswith('.pdf'):
                path = _media_path(src)
                if os.path.exists(path):
                    uris = _pdf_to_imgs(path)
                    return ''.join(
                        f'<img src="{u}" style="width:100%;height:auto;display:block;"/>'
                        for u in uris
                    )
            return ''

        html = _re.sub(r'<embed\b[^>]*/?>|<iframe\b[^>]*>.*?</iframe>', _replace_embed_tag, html, flags=_re.DOTALL)

        # 2. Substitui src= de imagens (aspas simples ou duplas) por base64
        def _sub(m):
            q   = m.group(1)
            src = m.group(2)

            if '/media' in src and not src.lower().endswith('.pdf'):
                b64 = _file_to_b64(_media_path(src))
                return f'src={q}{b64}{q}' if b64 else m.group(0)

            if '/static' in src.lower() or '/Static' in src:
                path = _static_path(src)
                b64 = _file_to_b64(path) if path else None
                return f'src={q}{b64}{q}' if b64 else m.group(0)

            return m.group(0)

        html = _re.sub(r"src=(['\"])([^'\"]+)\1", _sub, html)
        return html

    # ── Placeholders de texto ─────────────────────────────────────────────────

    def _process_placeholders(html):
        html = html.replace('{Assinatura Alegacao Defesa}',         _sig_cache['defesa'])
        html = html.replace('{Assinatura Reconsideracao}',          _sig_cache['recon'])
        html = html.replace('{Assinatura_Imagem_Oficial_Apurador}', _sig_cache['oficial'])
        html = html.replace('{Assinatura_Imagem_Testemunha_1}',     _sig_cache['test1'])
        html = html.replace('{Assinatura_Imagem_Testemunha_2}',     _sig_cache['test2'])
        html = html.replace('{Assinatura_Imagem_Comandante_GSD}',   cmd_sig_html)
        mil_idx[0] = 0
        html = _re.sub(r'\{Assinatura Militar Arrolado\}', _mil_sig, html)
        html = _re.sub(r'\{Botao [^}]+\}', '', html)
        html = html.replace('{Localidade}', '')
        html = _re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html)
        html = _re.sub(r'<div class="page-meta"[^>]*></div>', '', html)
        html = _re.sub(r'<p>(<br\s*\/?>|\s|&nbsp;)*</p>', '', html)
        html = _inline_all_srcs(html)
        return html

    # ── Dimensões de página ───────────────────────────────────────────────────

    _meta = {'w':21.0,'h':29.7,'top':2.5,'bot':2.5,'left':2.5,'right':2.5,
             'font':'Times New Roman','size':12.0}

    def _extract_meta(html):
        m = _re.search(
            r'<div class="page-meta"\s+data-width="([^"]+)"\s+data-height="([^"]+)"\s+'
            r'data-top="([^"]+)"\s+data-bottom="([^"]+)"\s+data-left="([^"]+)"\s+'
            r'data-right="([^"]+)"\s+data-font="([^"]+)"\s+data-fontsize="([^"]+)"', html)
        if m:
            _meta.update({'w':float(m.group(1)),'h':float(m.group(2)),
                          'top':float(m.group(3)),'bot':float(m.group(4)),
                          'left':float(m.group(5)),'right':float(m.group(6)),
                          'font':m.group(7),'size':float(m.group(8))})

    def _wrap(inner, padding=True):
        w,h = _meta['w'],_meta['h']
        if padding:
            t,b,l,r = _meta['top'],_meta['bot'],_meta['left'],_meta['right']
            pad = f'padding:{t}cm {r}cm {b}cm {l}cm;'
            font_css = f"font-family:'{_meta['font']}',serif;font-size:{_meta['size']}pt;"
        else:
            pad = 'padding:0;'
            font_css = ''
        return (f'<div style="width:{w}cm;height:{h}cm;{pad}{font_css}'
                f'box-sizing:border-box;overflow:hidden;page-break-after:always;">'
                f'{inner}</div>')

    # ── Placeholders de anexos ────────────────────────────────────────────────

    def _make_anexo_pages(anexo):
        ext = os.path.splitext(anexo.arquivo.name)[1].lower()
        path = anexo.arquivo.path
        if ext in ('.png','.jpg','.jpeg','.gif','.bmp','.webp'):
            b64 = _file_to_b64(path)
            # Imagem centralizada e contida na folha
            return [f'<img src="{b64}" style="width:100%;height:100%;object-fit:contain;display:block;"/>'] if b64 else []
        elif ext == '.pdf':
            # _pdf_to_imgs retorna lista de data URIs
            return [f'<img src="{uri}" style="width:100%;height:100%;object-fit:contain;display:block;"/>'
                    for uri in _pdf_to_imgs(path)]
        return [f'<p style="font-style:italic;">[Anexo: {os.path.basename(anexo.arquivo.name)}]</p>']

    placeholder_map = {
        '{ANEXOS_DEFESA_PLACEHOLDER}':                list(patd.anexos.filter(tipo='defesa')),
        '{ANEXOS_RECONSIDERACAO_PLACEHOLDER}':        list(patd.anexos.filter(tipo='reconsideracao')),
        '{ANEXO_OFICIAL_RECONSIDERACAO_PLACEHOLDER}': list(patd.anexos.filter(tipo='reconsideracao_oficial')),
        '{FORMULARIO_RESUMO_PLACEHOLDER}':            list(patd.anexos.filter(tipo='formulario_resumo')),
    }

    # ── Montar páginas ────────────────────────────────────────────────────────

    page_htmls = []

    for item in document_pages_raw:
        if item.strip() == PB:
            continue

        clean = item.strip().replace('<p>','').replace('</p>','').strip()
        matched_key = next((k for k in placeholder_map if clean == k), None)
        if matched_key:
            _extract_meta(item)
            for anexo in placeholder_map[matched_key]:
                for pg in _make_anexo_pages(anexo):
                    page_htmls.append(_wrap(pg, padding=False))
            continue

        _extract_meta(item)
        for sec in item.split(PB):
            if not sec.strip():
                continue
            page_htmls.append(_wrap(_process_placeholders(sec)))

    css_page = f'@page{{margin:0;size:{_meta["w"]}cm {_meta["h"]}cm;}}'
    full_html = (
        '<!DOCTYPE html><html lang="pt-BR"><head><meta charset="UTF-8"><style>'
        f'{css_page}'
        'body{margin:0;padding:0;background:white;}'
        'table{border-collapse:collapse;width:100%;}td,th{padding:2px 4px;}'
        'img{max-width:100%;}'
        'input,button,textarea,select{display:none!important;}'
        '.page-number{display:none;}'
        'a{color:inherit;text-decoration:none;}'
        '</style></head><body>'
        + ''.join(page_htmls)
        + '</body></html>'
    )

    filename_base = f"PATD_{patd.numero_patd or pk}"

    try:
        from weasyprint import HTML as WeasyprintHTML
        pdf_bytes = WeasyprintHTML(string=full_html).write_pdf()
    except Exception as e:
        logger.error("WeasyPrint falhou ao gerar PDF PATD %s: %s", pk, e)
        return HttpResponse('Erro ao gerar PDF.', status=500)

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename_base}.pdf"'
    response.write(pdf_bytes)
    _set_download_cookie(request, response)
    return response


def _set_download_cookie(request, response):
    """Seta cookie de token para o JS detectar fim do download."""
    token = request.GET.get('download_token')
    if token:
        response.set_cookie('download_token', token, max_age=60, samesite='Lax')


def exportar_patd_docx(request, pk):


    """

    Gera e serve um ficheiro DOCX a partir do conteúdo HTML da PATD,

    incluindo imagens, formatação correta E ANEXOS.

    """


    patd = get_object_or_404(PATD, pk=pk)


    context = _get_document_context(patd, for_docx=True)


    config = Configuracao.load()


    comandante_gsd = config.comandante_gsd




    document = Document()




    # --- INÍCIO DA MODIFICAÇÃO: Adicionar proteção contra edição ---


    from docx.oxml.ns import qn


    from docx.oxml import OxmlElement




    # Obtém o elemento <w:settings> do documento


    doc_settings = document.settings.element


    # Cria o elemento <w:documentProtection> para tornar o documento somente leitura


    protection = OxmlElement('w:documentProtection')


    protection.set(qn('w:edit'), 'readOnly')


    doc_settings.append(protection)


    # --- FIM DA MODIFICAÇÃO ---




    style = document.styles['Normal']


    font = style.font


    font.name = 'Times New Roman'


    font.size = Pt(12)




    section = document.sections[0]


    section.top_margin = Cm(1.5)


    section.bottom_margin = Cm(2.54)


    section.left_margin = Cm(2.15)


    section.right_margin = Cm(2.5)


    section.gutter = Cm(0)




    # Adiciona o número da página no rodapé


    add_page_number(section.footer.paragraphs[0])


    

    full_html_content = "".join(get_document_pages(patd, for_docx=True))




    anexos_defesa = patd.anexos.filter(tipo='defesa')


    anexos_reconsideracao = patd.anexos.filter(tipo='reconsideracao')


    anexos_reconsideracao_oficial = patd.anexos.filter(tipo='reconsideracao_oficial')




    soup = BeautifulSoup(full_html_content, 'html.parser')

    militar_sig_counter = 0
    placeholder_regex = re.compile(r'({[^}]+})')

    # Controle de estado para evitar páginas em branco duplicadas
    last_action_was_page_break = True  # Começa True para não adicionar quebra antes do 1º elemento
    was_last_p_empty = False
    last_para_was_pdf_image = False  # True quando o último parágrafo é uma imagem de página de PDF

    def _add_run_with_text(paragraph, text, bold=False, italic=False, underline=False, font_size_pt=None):
        """Adiciona um run de texto ao parágrafo com a formatação indicada."""
        if not text:
            return
        run = paragraph.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if underline:
            run.underline = True
        if font_size_pt:
            run.font.size = Pt(font_size_pt)

    def _process_node(node, paragraph, bold=False, italic=False, underline=False, font_size_pt=None):
        """
        Percorre recursivamente os nós HTML de um parágrafo e adiciona
        runs ao parágrafo DOCX com a formatação correta.
        Suporta: NavigableString, <strong>, <em>, <u>, <span>, <img>, <br>, <button>, <input>
        """
        nonlocal militar_sig_counter

        if isinstance(node, NavigableString):
            text = str(node)
            if not text:
                return
            # Divide o texto em partes: placeholders e texto normal
            parts = placeholder_regex.split(text)
            for part in parts:
                if not part:
                    continue
                if placeholder_regex.match(part):
                    _resolve_placeholder(part.strip(), paragraph, bold, italic, underline, font_size_pt)
                else:
                    # Trata markdown **texto** que pode aparecer em campos gerados pelo LLM
                    md_parts = re.split(r'(\*\*.*?\*\*)', part)
                    for md in md_parts:
                        if md.startswith('**') and md.endswith('**'):
                            _add_run_with_text(paragraph, md[2:-2], bold=True, italic=italic, underline=underline, font_size_pt=font_size_pt)
                        else:
                            _add_run_with_text(paragraph, md, bold, italic, underline, font_size_pt)
            return

        tag = getattr(node, 'name', None)
        if tag is None:
            return

        if tag == 'strong':
            for child in node.children:
                _process_node(child, paragraph, bold=True, italic=italic, underline=underline, font_size_pt=font_size_pt)
        elif tag == 'em':
            for child in node.children:
                _process_node(child, paragraph, bold=bold, italic=True, underline=underline, font_size_pt=font_size_pt)
        elif tag == 'u':
            for child in node.children:
                _process_node(child, paragraph, bold=bold, italic=italic, underline=True, font_size_pt=font_size_pt)
        elif tag == 'span':
            # Extrai font-size do style se existir
            fs = font_size_pt
            style = node.get('style', '')
            m = re.search(r'font-size:\s*([\d.]+)pt', style)
            if m:
                fs = float(m.group(1))
            for child in node.children:
                _process_node(child, paragraph, bold=bold, italic=italic, underline=underline, font_size_pt=fs)
        elif tag == 'img':
            src = node.get('src', '')
            if 'brasao.png' in src:
                img_path = finders.find('img/brasao.png')
                if img_path and os.path.exists(img_path):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.add_run().add_picture(img_path, width=Cm(3))
            elif src.startswith('data:image'):
                # Imagem base64 (assinatura inline)
                try:
                    _, b64data = src.split(';base64,')
                    paragraph.add_run().add_picture(io.BytesIO(base64.b64decode(b64data)), height=Cm(2.5))
                except Exception as e:
                    logger.error(f"Erro ao processar imagem base64 inline: {e}")
        elif tag == 'br':
            paragraph.add_run().add_break()
        elif tag in ('button', 'input', 'a'):
            # Botões/links do visualizador não vão para o DOCX
            pass
        else:
            # Tag desconhecida — desce nos filhos
            for child in node.children:
                _process_node(child, paragraph, bold=bold, italic=italic, underline=underline, font_size_pt=font_size_pt)

    def _resolve_placeholder(ph, paragraph, bold=False, italic=False, underline=False, font_size_pt=None):
        """Resolve um placeholder de assinatura/imagem no DOCX."""
        nonlocal militar_sig_counter
        try:
            if ph == '{Assinatura_Imagem_Comandante_GSD}':
                if patd.assinatura_cmd_gsd_despacho and patd.assinatura_cmd_gsd_despacho.path and os.path.exists(patd.assinatura_cmd_gsd_despacho.path):
                    paragraph.add_run().add_picture(patd.assinatura_cmd_gsd_despacho.path, height=Cm(2.5))
                elif comandante_gsd and comandante_gsd.assinatura:
                    _, img_str = comandante_gsd.assinatura.split(';base64,')
                    paragraph.add_run().add_picture(io.BytesIO(base64.b64decode(img_str)), height=Cm(2.5))

            elif ph == '{Assinatura_Imagem_Oficial_Apurador}':
                sig = patd.assinatura_oficial
                if sig and sig.path and os.path.exists(sig.path):
                    paragraph.add_run().add_picture(sig.path, height=Cm(2.5))

            elif ph == '{Assinatura_Imagem_Testemunha_1}':
                sig = patd.assinatura_testemunha1
                if sig and sig.path and os.path.exists(sig.path):
                    paragraph.add_run().add_picture(sig.path, height=Cm(2.5))

            elif ph == '{Assinatura_Imagem_Testemunha_2}':
                sig = patd.assinatura_testemunha2
                if sig and sig.path and os.path.exists(sig.path):
                    paragraph.add_run().add_picture(sig.path, height=Cm(2.5))

            elif ph == '{Assinatura Militar Arrolado}':
                assinaturas = patd.assinaturas_militar or []
                if militar_sig_counter < len(assinaturas) and assinaturas[militar_sig_counter]:
                    url = assinaturas[militar_sig_counter]
                    path = os.path.join(settings.MEDIA_ROOT, url.replace(settings.MEDIA_URL, '', 1))
                    if os.path.exists(path):
                        paragraph.add_run().add_picture(path, height=Cm(2.5))
                militar_sig_counter += 1

            elif ph == '{Assinatura Alegacao Defesa}':
                sig = patd.assinatura_alegacao_defesa
                if sig and sig.path and os.path.exists(sig.path):
                    paragraph.add_run().add_picture(sig.path, height=Cm(2.5))

            elif ph == '{Assinatura Reconsideracao}':
                sig = patd.assinatura_reconsideracao
                if sig and sig.path and os.path.exists(sig.path):
                    paragraph.add_run().add_picture(sig.path, height=Cm(2.5))

            elif ph == '{Botao Definir Nova Punicao}':
                # Substitui pelo texto da nova punição se existir, senão omite
                if patd.nova_punicao_tipo:
                    nova = f"{patd.nova_punicao_dias} de {patd.nova_punicao_tipo}" if patd.nova_punicao_dias else patd.nova_punicao_tipo
                    _add_run_with_text(paragraph, nova, bold, italic, underline, font_size_pt)

            elif ph in ('{Botao Assinar Oficial}', '{Botao Assinar Testemunha 1}',
                        '{Botao Assinar Testemunha 2}', '{Botao Adicionar Alegacao}',
                        '{Botao Adicionar Reconsideracao}',
                        '{Botao Assinar Defesa}', '{Botao Assinar Reconsideracao}',
                        '{Botao Assinar Ciencia}', '{Botao Adicionar Texto Reconsideracao}'):
                # Botões de ação não vão para o DOCX
                pass

            else:
                # Placeholder de texto desconhecido — escreve como texto
                _add_run_with_text(paragraph, ph, bold, italic, underline, font_size_pt)

        except Exception as e:
            logger.error(f"Erro ao resolver placeholder {ph} no DOCX: {e}")
            _add_run_with_text(paragraph, f'[{ph}]')

    # Itera em ordem correta: top-level <p> e <div class="manual-page-break">.
    # <div> wrappers (ex: data-document-id) são transparentes — seus filhos são emitidos inline.
    def _iter_doc_elements(node):
        for child in node.children:
            tag = getattr(child, 'name', None)
            if tag == 'p':
                yield child
            elif tag == 'div':
                classes = child.get('class', [])
                if 'manual-page-break' in classes:
                    yield child
                else:
                    # Div wrapper — desce transparentemente
                    yield from _iter_doc_elements(child)

    for element in _iter_doc_elements(soup):

        if element.name == 'div' and 'manual-page-break' in element.get('class', []):
            if not last_action_was_page_break:
                last_p = document.paragraphs[-1] if document.paragraphs else document.add_paragraph()
                last_p.add_run().add_break(WD_BREAK.PAGE)
                last_action_was_page_break = True
            last_para_was_pdf_image = False
            was_last_p_empty = False
            continue




        if element.name == 'p':


            text_content_for_check = element.get_text().replace('\xa0', '').strip()




            # Lida com placeholders de anexo


            if "{ANEXOS_DEFESA_PLACEHOLDER}" in text_content_for_check and anexos_defesa.exists():
                for i, anexo in enumerate(anexos_defesa):
                    if i > 0:
                        last_p = document.paragraphs[-1] if document.paragraphs else document.add_paragraph()
                        last_p.add_run().add_break(WD_BREAK.PAGE)
                    _append_anexo_content(document, anexo)
                last_action_was_page_break = False
                was_last_p_empty = True
                last_para_was_pdf_image = False
                continue

            if "{ANEXOS_RECONSIDERACAO_PLACEHOLDER}" in text_content_for_check and anexos_reconsideracao.exists():
                for i, anexo in enumerate(anexos_reconsideracao):
                    if i > 0:
                        last_p = document.paragraphs[-1] if document.paragraphs else document.add_paragraph()
                        last_p.add_run().add_break(WD_BREAK.PAGE)
                    _append_anexo_content(document, anexo)
                last_action_was_page_break = False
                was_last_p_empty = True
                last_para_was_pdf_image = False
                continue

            if "{ANEXO_OFICIAL_RECONSIDERACAO_PLACEHOLDER}" in text_content_for_check and anexos_reconsideracao_oficial.exists():
                for i, anexo in enumerate(anexos_reconsideracao_oficial):
                    if i > 0:
                        last_p = document.paragraphs[-1] if document.paragraphs else document.add_paragraph()
                        last_p.add_run().add_break(WD_BREAK.PAGE)
                    _append_anexo_content(document, anexo)
                last_action_was_page_break = False
                was_last_p_empty = True
                last_para_was_pdf_image = False
                continue

            formulario_resumo_qs = patd.anexos.filter(tipo='formulario_resumo')
            if "{FORMULARIO_RESUMO_PLACEHOLDER}" in text_content_for_check and formulario_resumo_qs.exists():
                for i, anexo in enumerate(formulario_resumo_qs):
                    if i > 0:
                        last_p = document.paragraphs[-1] if document.paragraphs else document.add_paragraph()
                        last_p.add_run().add_break(WD_BREAK.PAGE)
                    _append_anexo_content(document, anexo)
                last_action_was_page_break = False
                was_last_p_empty = True
                last_para_was_pdf_image = False
                continue

            # Trata <embed> e <img> de anexos inline (oficio, ficha, formulario).
            # <embed> sem correspondência é descartado (PDF não renderizável no DOCX).
            # <img> sem correspondência (ex: brasão) cai no processamento normal abaixo.
            embed_tag = element.find('embed')
            img_tag = element.find('img', attrs={'src': True})
            if embed_tag or img_tag:
                src = (embed_tag or img_tag).get('src', '')
                found_anexo = False
                if src:
                    for anexo in patd.anexos.filter(tipo__in=['oficio_lancamento', 'ficha_individual', 'formulario_resumo', 'relatorio_delta_base', 'npd_base']):
                        if anexo.arquivo.url == src:
                            _append_anexo_content(document, anexo)
                            last_action_was_page_break = False
                            was_last_p_empty = True
                            last_para_was_pdf_image = False
                            found_anexo = True
                            break
                if found_anexo or embed_tag:
                    continue  # embed sempre pula; img só pula se era um anexo




            is_empty_paragraph = not text_content_for_check and not element.find('img')


            

            if is_empty_paragraph:
                # Ignora parágrafos vazios logo após uma quebra de página
                if not last_action_was_page_break and not was_last_p_empty:
                    p = document.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    was_last_p_empty = True
                    last_action_was_page_break = False
                continue
            else:
                was_last_p_empty = False
                last_action_was_page_break = False
                last_para_was_pdf_image = False




            p = document.add_paragraph()


            p_format = p.paragraph_format


            p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


            p_format.space_before = Pt(0)


            p_format.space_after = Pt(0)


            

            if element.has_attr('style'):
                style_str = element['style']
                if 'text-align: center' in style_str:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'text-align: right' in style_str:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif 'text-align: justify' in style_str:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                elif 'text-align: left' in style_str:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                indent_m = re.search(r'padding-left:\s*([\d.]+)pt', style_str)
                if indent_m:
                    p_format.left_indent = Pt(float(indent_m.group(1)))
                indent_fi = re.search(r'text-indent:\s*([\d.]+)pt', style_str)
                if indent_fi:
                    p_format.first_line_indent = Pt(float(indent_fi.group(1)))




            for content in element.contents:
                _process_node(content, p)

    
    # Garante que o rodapé com número de página apareça em TODAS as seções do documento.
    # Seções adicionais (criadas por sub-documentos) devem herdar o rodapé da seção 1.
    _propagate_footer_to_all_sections(document)

    # Salva o DOCX em memória
    docx_buffer = io.BytesIO()
    document.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    # Tenta converter para PDF, remover páginas em branco e exportar como PDF
    filename_base = f'PATD_{patd.numero_patd}'
    pdf_bytes, converted = _remove_blank_pages_from_docx(docx_bytes, filename_base)

    is_preview = request.GET.get('preview') == '1'
    disposition = 'inline' if is_preview else 'attachment'

    if converted and pdf_bytes:
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'{disposition}; filename={filename_base}.pdf'
        response.write(pdf_bytes)
    else:
        # Fallback: exporta o DOCX sem conversão
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={filename_base}.docx'
        response.write(docx_bytes)

    _set_download_cookie(request, response)
    return response


