import logging
import pandas as pd
import openpyxl
import os
import tempfile
import httpx, json
import random
import unicodedata
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse_lazy
from django.core.paginator import Paginator
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.contrib.auth.decorators import login_required, user_passes_test
from django.utils.decorators import method_decorator
from django.utils import timezone
from django.views.decorators.clickjacking import xframe_options_sameorigin
from django.views.decorators.http import require_POST
from django.views.decorators.http import require_GET
from django.contrib.auth.models import Group
from Secao_pessoal.models import (
    Efetivo, Posto, Quad, Especializacao, OM, Setor, Subsetor, SolicitacaoTrocaSetor,
    HistoricoInspsau, MovimentacaoEfetivo, LotacaoPessoal,
)
from django.contrib.auth import get_user_model as _get_user_model
from caixa_entrada.models import Notificacao, Mensagem as _Mensagem
from auditoria.utils import registrar, resolver_label

_PESSOAL_PERMISSAO_MAP = {
    'Seção de Pessoal (S1)': 'S1- Efetivo',
}

_User = _get_user_model()


def _enviar_mensagem_sistema(remetente_militar, destinatario_militar, assunto, corpo):
    """Envia uma Mensagem da nova caixa de entrada entre dois militares."""
    try:
        rem  = _User.objects.filter(profile__militar=remetente_militar).first()
        dest = _User.objects.filter(profile__militar=destinatario_militar).first()
        if rem and dest:
            msg = _Mensagem.objects.create(remetente=rem, assunto=assunto, corpo=corpo)
            msg.destinatarios.add(dest)
    except Exception:
        pass
from .forms import MilitarForm, LotacaoPessoalForm
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.contrib import messages
from django.db.models import Q, Max, Case, When, Value, IntegerField, Count, Sum
from difflib import SequenceMatcher
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.utils import get_column_letter
from datetime import datetime, date, timedelta
try:
    import pytesseract
except ImportError:
    pytesseract = None
from PIL import Image
import fitz 
import io
import re
# Se estiver no Windows e der erro, descomente e ajuste o caminho do seu tesseract:
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
from .analise_inspsau import analisar_inspsau_pdf
from .analise_fq import analisar_fq_documento
from chamada.models import RegistroChamada as ChamadaRegistro

logger = logging.getLogger(__name__)

def is_s1_member(user):
    return user.groups.filter(name='Seção de Pessoal (S1)').exists()

s1_required = user_passes_test(is_s1_member)

def get_chefe_por_grupo(nome_setor):
    if not nome_setor:
        return None
    nome_grupo = f"Chefe - {nome_setor}"
    try:
        grupo = Group.objects.get(name__iexact=nome_grupo)
        user_chefe = grupo.user_set.filter(is_active=True, profile__militar__isnull=False).first()
        if user_chefe:
            return user_chefe.profile.militar
    except Group.DoesNotExist:
        pass
    
    setor_obj = Setor.objects.filter(nome=nome_setor).first()
    return setor_obj.chefe if setor_obj and hasattr(setor_obj, 'chefe') else None


def obter_situacao_inspsau(letra):
    """Mapeia a finalidade da INSPSAU para o campo Situação (limite de 50 caracteres)."""
    if not letra:
        return 'Inspeção de Saúde'
    letra = letra.upper().strip()
    if letra.startswith('G'): return 'De Junta'
    elif letra.startswith('L'): return 'Licença Pessoal Navegação Aérea - LPNA'
    elif letra.startswith('A'): return 'Incorporação ou Desincorporação'
    elif letra.startswith('B'): return 'Matrícula em Escolas de Formação'
    elif letra.startswith('C'): return 'Concurso Cargos Civis no COMAER'
    elif letra.startswith('D'): return 'Verificação Periódica - Temporários'
    elif letra.startswith('E'): return 'Manutenção de Tratamento de Saúde'
    elif letra.startswith('F1'): return 'Missão no Exterior'
    elif letra.startswith('F2'): return 'Localidade Especial'
    elif letra.startswith('H'): return 'Verificação Periódica de Carreira'
    elif letra.startswith('I'): return 'Cursos Operacionais ou Ativ. Aérea'
    elif letra.startswith('J'): return 'Designação PTTC'
    elif letra.startswith('N'): return 'Ordem Judicial/Reversão/DSA'
    elif letra.startswith('O'): return 'Benefícios e Licenças'
    elif letra.startswith('P'): return 'Acidentes/Incidentes Aeronáuticos'
    elif letra.startswith('R1'): return 'Estado de Saúde Desertor/Insubmisso'
    elif letra.startswith('R2'): return 'Verificação de Capacidade Cognitiva'
    else: return f'Inspeção Finalidade {letra}'[:50]

@s1_required
def index(request):
    _POSTOS      = ['TC', 'MJ', 'CP', '1T', '2T', 'ASP', 'SO', '1S', '2S', '3S', 'CB', 'S1', 'S2', 'REC']
    _POSTOS_SET  = ['TC', 'MJ', 'CP', '1T', '2T', 'ASP', 'SO', '1S', '2S', '3S', 'CB', 'S1', 'S2']
    _DISPLAY     = {'ASP': 'AP'}
    _SETORES     = ['ESI', 'EPA', 'EFSD', 'SAP', 'SOP', 'CMD']
    hoje         = date.today()

    def _por_posto(qs):
        raw = {}
        for item in qs.values('posto').annotate(c=Count('id')):
            p = (item['posto'] or '').strip().upper()
            raw[p] = raw.get(p, 0) + item['c']
        rows = [(_DISPLAY.get(p, p), raw.get(p, 0)) for p in _POSTOS]
        total_srec = sum(v for p, v in rows if p != 'REC')
        total = sum(v for _, v in rows)
        return {'rows': rows, 'total_srec': total_srec, 'total': total}

    def _setor_key(s):
        s = (s or '').upper().strip()
        if 'ESI' in s: return 'ESI'
        if 'PCG' in s or 'CARCERAGEM' in s or 'CANIL' in s or 'EPA' in s: return 'EPA'
        if 'EFSD' in s: return 'EFSD'
        if 'SAP' in s: return 'SAP'
        if 'SOP' in s: return 'SOP'
        if 'CMD' in s: return 'CMD'
        return None

    qs_geral      = Efetivo.objects.exclude(situacao__iexact='Baixado')
    qs_operac     = Efetivo.objects.filter(
        Q(situacao__iexact='Ativo') | Q(situacao__iexact='Ativa') |
        Q(situacao__iexact='PSV GSD-GL') |
        Q(situacao__iexact='PSV', data_vencimento_prestacao__lt=hoje)
    )
    qs_psv_ext    = Efetivo.objects.filter(situacao__iexact='PSV')
    qs_psv_gsd    = Efetivo.objects.filter(situacao__iexact='PSV GSD-GL')
    qs_baixados   = Efetivo.all_objects.filter(situacao__iexact='Baixado')
    qs_desertores = Efetivo.all_objects.filter(situacao__iexact='Desertor')

    # --- Tabelas por posto ---
    tabelas = [
        {'titulo': 'Quantitativo de Efetivo Total',
         'subtitulo': 'Efetivo do GSD-GL',
         'data': _por_posto(qs_geral)},
        {'titulo': 'Quantitativo de Efetivo Operacional Total',
         'subtitulo': 'Efetivo do GSD-GL',
         'data': _por_posto(qs_operac)},
        {'titulo': 'Quantitativo de Efetivo Prestando Serviço Externo',
         'subtitulo': 'Militares Prestando Serviço',
         'data': _por_posto(qs_psv_ext)},
        {'titulo': 'Quantitativo de Efetivo Prestando Serviço para o GSD-GL',
         'subtitulo': 'Efetivo do GSD-GL',
         'data': _por_posto(qs_psv_gsd)},
        {'titulo': 'Quantitativo de Desligados',
         'subtitulo': 'Desligados',
         'data': _por_posto(qs_baixados)},
    ]

    # --- Tabela cruzada: postos × setores ---
    _disp_set = [_DISPLAY.get(p, p) for p in _POSTOS_SET]
    setor_grid = {s: {p: 0 for p in _disp_set} for s in _SETORES}
    setor_totals_d = {s: 0 for s in _SETORES}

    for item in qs_operac.values('posto', 'setor').annotate(c=Count('id')):
        sk = _setor_key(item['setor'])
        if sk is None:
            continue
        p_disp = _DISPLAY.get((item['posto'] or '').strip().upper(), (item['posto'] or '').strip().upper())
        if p_disp in setor_grid[sk]:
            setor_grid[sk][p_disp] += item['c']
            setor_totals_d[sk] += item['c']

    setor_rows = []
    for p in _disp_set:
        valores = [setor_grid[s][p] for s in _SETORES]
        setor_rows.append({'posto': p, 'valores': valores, 'total': sum(valores)})
    setor_totals = [setor_totals_d[s] for s in _SETORES]
    total_setor  = sum(setor_totals)
    total_operac = qs_operac.count()
    setor_pie_pct = [
        round(v / total_operac * 100, 1) if total_operac else 0 for v in setor_totals
    ]

    # --- Por Especialidade (agrupado só pela especialidade, sem posto) ---
    esp_raw = (
        qs_geral.values('especializacao')
        .annotate(c=Count('id'))
        .order_by('especializacao')
    )
    esp_rows = []
    for item in esp_raw:
        e = (item['especializacao'] or '').strip().upper()
        esp_rows.append((e if e else '(sem especialidade)', item['c']))

    tlp_raw = (
        LotacaoPessoal.objects
        .values('posto', 'especializacao')
        .annotate(total=Sum('vagas_previstas'))
        .order_by('posto', 'especializacao')
    )
    tlp_rows = []
    for item in tlp_raw:
        p = _DISPLAY.get((item['posto'] or '').strip().upper(), (item['posto'] or '').strip().upper())
        e = (item['especializacao'] or '').strip().upper()
        tlp_rows.append((f"{p} {e}".strip() if e else p, item['total'] or 0))

    # --- Cards de quantitativos rápidos ---
    total_geral       = qs_geral.count()
    total_agd         = Efetivo.objects.filter(
        Q(situacao__icontains='desligamento') | Q(situacao__iexact='AGD. DESLIGAMENTO')
    ).count()
    total_psv_servico = qs_psv_ext.count()
    total_desertores_n= qs_desertores.count()
    total_junta       = Efetivo.objects.filter(
        Q(situacao__icontains='junta') | Q(situacao__iexact='De Junta')
    ).count()
    total_justica     = Efetivo.objects.filter(situacao__icontains='justi').count()
    total_psv_vencida = Efetivo.objects.filter(
        situacao__iexact='PSV', data_vencimento_prestacao__lt=hoje
    ).count()

    # --- Tabelas resumo por categoria (Efetivo Total e Operacional) ---
    _dados_geral  = _por_posto(qs_geral)
    _dados_operac_cat = _por_posto(qs_operac)

    def _categorias(dados):
        m = {label: count for label, count in dados['rows']}
        return [
            ('OF',   sum(m.get(p, 0) for p in ['TC', 'MJ', 'CP', '1T', '2T', 'AP'])),
            ('SO',   m.get('SO', 0)),
            ('SGT',  sum(m.get(p, 0) for p in ['1S', '2S', '3S'])),
            ('CABO', m.get('CB', 0)),
            ('S1',   m.get('S1', 0)),
            ('S2',   m.get('S2', 0)),
            ('REC',  m.get('REC', 0)),
        ]

    context = {
        'tabelas': tabelas,
        'desertores': _por_posto(qs_desertores),
        'setores_dash': _SETORES,
        'setor_rows': setor_rows,
        'setor_totals': setor_totals,
        'total_setor': total_setor,
        'total_operac': total_operac,
        'setor_pie_labels': json.dumps(_SETORES),
        'setor_pie_data': json.dumps(setor_totals),
        'setor_pie_pct': json.dumps(setor_pie_pct),
        'esp_rows': esp_rows,
        'tlp_rows': tlp_rows,
        # Cards rápidos
        'total_geral':       total_geral,
        'total_agd':         total_agd,
        'total_psv_servico': total_psv_servico,
        'total_desertores_n':total_desertores_n,
        'total_junta':       total_junta,
        'total_justica':     total_justica,
        'total_psv_vencida': total_psv_vencida,
        # Tabelas por categoria
        'cat_geral':  _categorias(_dados_geral),
        'cat_operac': _categorias(_dados_operac_cat),
    }
    return render(request, 'Secao_pessoal/index.html', context)

@s1_required
@require_POST
def tornar_recrutas_soldados(request):
    updated = Efetivo.objects.filter(posto='REC').update(posto='S2', especializacao='NE')
    if updated:
        messages.success(request, f'{updated} recruta(s) promovido(s) para S2 (especialidade → NE) com sucesso.')
    else:
        messages.info(request, 'Nenhum recruta (REC) encontrado no efetivo ativo.')
    return redirect('Secao_pessoal:militar_list')


@s1_required
def painel_chefe(request):
    efetivo = Efetivo.objects.exclude(situacao__iexact='Baixado')
    
    # EFETIVO GERAL
    total_efetivo = efetivo.count()
    total_indisponiveis = efetivo.exclude(Q(situacao__iexact='Ativo') | Q(situacao__exact='') | Q(situacao__isnull=True)).count()
    
    # SAÚDE (INSPSAU)
    hoje = timezone.now().date()
    daqui_30_dias = hoje + timedelta(days=30)
    inspsau_vencidas = efetivo.filter(inspsau_validade__lt=hoje).count()
    inspsau_a_vencer = efetivo.filter(inspsau_validade__gte=hoje, inspsau_validade__lte=daqui_30_dias).count()
    total_junta = efetivo.filter(situacao__iexact='De Junta').count()

    # CONTROLE
    trocas_pendentes = SolicitacaoTrocaSetor.objects.filter(status__in=['pendente_atual', 'pendente_destino']).count()
    total_baixados = Efetivo.objects.filter(situacao__iexact='Baixado').count()

    context = {
        'total_efetivo': total_efetivo,
        'total_indisponiveis': total_indisponiveis,
        'inspsau_vencidas': inspsau_vencidas,
        'inspsau_a_vencer': inspsau_a_vencer,
        'total_junta': total_junta,
        'trocas_pendentes': trocas_pendentes,
        'total_baixados': total_baixados,
    }

    return render(request, 'Secao_pessoal/painel_chefe.html', context)

@s1_required
def inspsau(request):
    if request.method == 'POST':
        # --- Lógica para lidar com a confirmação do usuário ---
        if 'militar_id_confirmado' in request.POST:
            militar_id = request.POST.get('militar_id_confirmado')
            finalidade_ia = request.POST.get('finalidade')
            validade_ia_str = request.POST.get('validade')
            parecer_ia = request.POST.get('parecer')
            pdf_file = request.FILES.get('pdf_file')

            if not all([militar_id, finalidade_ia, pdf_file]):
                return JsonResponse({'status': 'error', 'message': 'Dados de confirmação incompletos.'}, status=400)

            try:
                militar = Efetivo.objects.get(id=militar_id)

                significado = obter_situacao_inspsau(finalidade_ia)
                validade_obj = None
                if validade_ia_str and str(validade_ia_str).lower() != 'none':
                    try:
                        validade_obj = datetime.strptime(validade_ia_str, '%d/%m/%Y').date()
                    except (ValueError, TypeError):
                        pass

                # VERIFICA DUPLICIDADE NO HISTÓRICO
                if HistoricoInspsau.objects.filter(militar=militar, finalidade=finalidade_ia, validade=validade_obj).exists():
                    message = f"Já existe um registro de inspeção com finalidade '{finalidade_ia}' e validade '{validade_ia_str or 'N/A'}' para o militar {militar.posto} {militar.nome_guerra} no histórico."
                    return JsonResponse({'status': 'error', 'message': message}, status=409)

                militar.documento_inspsau = pdf_file
                if finalidade_ia and finalidade_ia.upper().startswith('G'):
                    militar.situacao = 'De Junta'
                elif militar.situacao == significado or militar.situacao == 'De Junta':
                    militar.situacao = 'Ativo'

                observacao_final = f"INSPSAU Finalidade: {finalidade_ia}."
                if validade_ia_str and str(validade_ia_str).lower() != 'none':
                    observacao_final += f" Validade: {validade_ia_str}"
                militar.observacao = observacao_final
                militar.inspsau_finalidade = finalidade_ia
                militar.inspsau_validade = validade_obj
                militar.inspsau_parecer = parecer_ia
                militar.save(update_fields=['observacao', 'situacao', 'documento_inspsau', 'inspsau_finalidade', 'inspsau_validade', 'inspsau_parecer'])
                messages.success(request, f"Inspeção do militar {militar.posto} {militar.nome_guerra} atualizada com sucesso. Finalidade: {finalidade_ia}.")
                return JsonResponse({'status': 'success'})
            except Efetivo.DoesNotExist:
                return JsonResponse({'status': 'error', 'message': 'Militar confirmado não foi encontrado.'}, status=404)
            except Exception as e:
                return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

    if request.method == 'POST' and request.FILES.get('pdf_file'):
        pdf_file = request.FILES['pdf_file']
        
        try:
            # Salva o PDF num ficheiro temporário
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                for chunk in pdf_file.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name

            # Extração Híbrida: Tenta texto nativo, se falhar ou tiver pouco texto, faz OCR
            content = ""
            doc = fitz.open(temp_file_path)
            for page in doc:
                page_text = page.get_text()
                if len(page_text.strip()) > 50:
                    content += page_text + "\n\n"
                else:
                    try:
                        pix = page.get_pixmap(dpi=300)
                        img = Image.open(io.BytesIO(pix.tobytes("png")))
                        if pytesseract:
                            content += pytesseract.image_to_string(img, lang='por') + "\n\n"
                        else:
                            logger.warning("pytesseract não instalado — OCR ignorado.")
                    except Exception as e:
                        logger.warning("Erro ao realizar OCR na página: %s", e)
            doc.close()
            os.remove(temp_file_path) # Remove o ficheiro temporário

            # Analisa o conteúdo com a IA
            resultado_analise = analisar_inspsau_pdf(content)

            # Busca o militar no banco de dados
            nome_completo_ia = resultado_analise.nome_completo
            posto_ia = resultado_analise.posto
            finalidade_ia = resultado_analise.finalidade
            validade_ia_str = getattr(resultado_analise, 'validade', None)
            parecer_ia = getattr(resultado_analise, 'parecer', '')

            significado = obter_situacao_inspsau(finalidade_ia)
            
            observacao_final = f"INSPSAU Finalidade: {finalidade_ia}."
            validade_obj = None
            if validade_ia_str and str(validade_ia_str).lower() != 'none':
                observacao_final += f" Validade: {validade_ia_str}"
                try:
                    validade_obj = datetime.strptime(validade_ia_str, '%d/%m/%Y').date()
                except (ValueError, TypeError):
                    pass
            
            militar = None
            if nome_completo_ia:
                # Tenta encontrar pelo nome completo e posto para maior precisão
                candidatos = Efetivo.all_objects.filter(
                    nome_completo__icontains=nome_completo_ia,
                    posto__iexact=posto_ia
                )
                if candidatos.count() == 1:
                    militar = candidatos.first()
                else:
                    # Se não encontrar ou houver ambiguidade, tenta só pelo nome
                    candidatos = Efetivo.all_objects.filter(nome_completo__icontains=nome_completo_ia)
                    if candidatos.count() == 1:
                        militar = candidatos.first()

            if militar:
                # VERIFICA DUPLICIDADE NO HISTÓRICO
                if HistoricoInspsau.objects.filter(militar=militar, finalidade=finalidade_ia, validade=validade_obj).exists():
                    message = f"Já existe um registro de inspeção com finalidade '{finalidade_ia}' e validade '{validade_ia_str or 'N/A'}' para o militar {militar.posto} {militar.nome_guerra} no histórico."
                    return JsonResponse({'status': 'error', 'message': message}, status=409)

                # Atualiza a observação do militar com a finalidade
                militar.documento_inspsau = pdf_file # Salva o arquivo PDF
                if finalidade_ia and finalidade_ia.upper().startswith('G'):
                    militar.situacao = 'De Junta'
                elif militar.situacao == significado or militar.situacao == 'De Junta':
                    militar.situacao = 'Ativo'
                militar.observacao = observacao_final
                militar.inspsau_finalidade = finalidade_ia
                militar.inspsau_validade = validade_obj
                militar.inspsau_parecer = parecer_ia
                militar.save(update_fields=['observacao', 'situacao', 'documento_inspsau', 'inspsau_finalidade', 'inspsau_validade', 'inspsau_parecer'])

                messages.success(request, f"Inspeção do militar {militar.posto} {militar.nome_guerra} atualizada com sucesso. Finalidade: {finalidade_ia}.")
                return JsonResponse({'status': 'success'})
            else:
                # --- INÍCIO DA LÓGICA DE BUSCA POR SIMILARIDADE ---
                best_match = None
                highest_ratio = 0.7  # Limiar de similaridade de 70%

                if nome_completo_ia:
                    nome_normalizado_ia = normalize_name(nome_completo_ia)
                    todos_militares = Efetivo.all_objects.all()

                    for m in todos_militares:
                        nome_normalizado_db = normalize_name(m.nome_completo)
                        ratio = SequenceMatcher(None, nome_normalizado_ia, nome_normalizado_db).ratio()
                        
                        if ratio > highest_ratio:
                            highest_ratio = ratio
                            best_match = m
                
                if best_match:
                    # Encontrou um militar parecido. Pede confirmação ao usuário.
                    return JsonResponse({
                        'status': 'confirm',
                        'message': f"O militar '{nome_completo_ia}' não foi encontrado. Você quis dizer '{best_match.posto} {best_match.nome_completo}'?",
                        'militar_encontrado': {
                            'id': best_match.id,
                        },
                        'dados_inspsau': {
                            'finalidade': finalidade_ia,
                            'validade': validade_ia_str,
                            'parecer': parecer_ia,
                        }
                    })
                else:
                    # Não encontrou nenhum militar, nem parecido. Retorna um erro.
                    # --- ALTERAÇÃO: Retorna status 'not_found' para acionar a busca manual no frontend ---
                    return JsonResponse({
                        'status': 'not_found',
                        'message': f"O militar '{nome_completo_ia}' não foi encontrado. Deseja procurar manualmente?",
                        'dados_inspsau': {
                            'finalidade': finalidade_ia,
                            'validade': validade_ia_str,
                            'parecer': parecer_ia,
                        }
                    })


        except Exception as e:
            messages.error(request, f"Erro ao processar o PDF: {str(e)}")
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

    # Dashboard básico: Lista militares que estão "De Junta" ou possuem um resultado de INSPSAU a exibir
    query = request.GET.get('q')
    militares_baixados = Efetivo.all_objects.filter(
        Q(situacao__iexact='De Junta') | Q(observacao__icontains='INSPSAU Finalidade')
    )

    if query:
        militares_baixados = militares_baixados.filter(
            Q(posto__icontains=query) |
            Q(nome_guerra__icontains=query) |
            Q(nome_completo__icontains=query) |
            Q(observacao__icontains=query)
        )

    militares_baixados = militares_baixados.order_by('nome_guerra')
    context = {'militares_baixados': militares_baixados, 'current_query': query or ''}
    return render(request, 'Secao_pessoal/inspsau.html', context)

#Efetivo

@method_decorator(s1_required, name='dispatch')
class MilitarListView(ListView):
    model = Efetivo
    template_name = 'Secao_pessoal/militar_list.html'
    context_object_name = 'militares'
    ordering = ['nome_guerra']
    paginate_by = 20

    def get_template_names(self):
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return ['Secao_pessoal/militar_list_partial.html']
        return ['Secao_pessoal/militar_list.html']

    def get_queryset(self):
        query = self.request.GET.get('q')
        rank_order = Case(
            When(posto='CL', then=Value(0)), When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
            When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)),When(posto='ASP', then=Value (6)), When(posto='SO', then=Value(7)),
            When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)), When(posto='3S', then=Value(10)),
            When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)), When(posto='S2', then=Value(13)),When(posto='REC', then=Value(14)),
            default=Value(99), output_field=IntegerField(),
        )
        qs = super().get_queryset().annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_completo')
        if query:
            q_objects = Q(nome_completo__icontains=query) | \
                        Q(nome_guerra__icontains=query) | \
                        Q(posto__icontains=query)

            if query.isdigit():
                q_objects |= Q(saram__icontains=query)

            qs = qs.filter(q_objects)

        def _split(param):
            return [v.strip() for v in param.split(',') if v.strip()]

        posto_f  = _split(self.request.GET.get('posto_f', ''))
        esp_f    = _split(self.request.GET.get('esp_f', ''))
        sit_f    = _split(self.request.GET.get('sit_f', ''))
        obs_f    = _split(self.request.GET.get('obs_f', ''))
        setor_f  = _split(self.request.GET.get('setor_f', ''))

        if posto_f:
            qs = qs.filter(posto__in=posto_f)
        if esp_f:
            esp_q = Q()
            for v in esp_f:
                if v == '__vazio__':
                    esp_q |= Q(especializacao='') | Q(especializacao__isnull=True)
                else:
                    esp_q |= Q(especializacao=v)
            qs = qs.filter(esp_q)
        if sit_f:
            sit_q = Q()
            for v in sit_f:
                if v == '__vazio__':
                    sit_q |= Q(situacao='') | Q(situacao__isnull=True)
                else:
                    sit_q |= Q(situacao=v)
            qs = qs.filter(sit_q)
        if obs_f:
            obs_q = Q()
            for v in obs_f:
                if v == '__vazio__':
                    obs_q |= Q(observacao='') | Q(observacao__isnull=True)
                else:
                    obs_q |= Q(observacao=v)
            qs = qs.filter(obs_q)
        if setor_f:
            set_q = Q()
            for v in setor_f:
                if v == '__vazio__':
                    set_q |= Q(setor='') | Q(setor__isnull=True)
                else:
                    set_q |= Q(setor=v)
            qs = qs.filter(set_q)

        return qs

@method_decorator(s1_required, name='dispatch')
class MilitarCreateView(CreateView):
    model = Efetivo
    form_class = MilitarForm
    template_name = 'Secao_pessoal/militar_form.html'
    success_url = reverse_lazy('Secao_pessoal:militar_list')

    def get_initial(self):
        initial = super().get_initial()
        initial['nome_completo'] = self.request.GET.get('nome_completo', '')
        initial['posto'] = self.request.GET.get('posto', '')
        initial['situacao'] = self.request.GET.get('situacao', 'Ativo')
        initial['observacao'] = self.request.GET.get('observacao', '')
        return initial

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Adicionar Militar'
        return context

@method_decorator(s1_required, name='dispatch')
class MilitarUpdateView(UpdateView):
    model = Efetivo
    form_class = MilitarForm
    template_name = 'Secao_pessoal/militar_form.html'
    success_url = reverse_lazy('Secao_pessoal:militar_list')

    def get_success_url(self):
        next_url = self.request.POST.get('next') or self.request.GET.get('next')
        if next_url and next_url.startswith('/'):
            return next_url
        return str(self.success_url)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar Militar'
        context['next_url'] = self.request.POST.get('next') or self.request.GET.get('next', '')
        return context

@method_decorator(s1_required, name='dispatch')
class MilitarDeleteView(DeleteView):
    model = Efetivo
    template_name = 'Secao_pessoal/militar_confirm_delete.html'
    success_url = reverse_lazy('Secao_pessoal:militar_list')

    def get_queryset(self):
        return Efetivo.all_objects.all()

    def post(self, request, *args, **kwargs):
        from django.db.models.deletion import ProtectedError
        self.object = self.get_object()
        nome_guerra = self.object.nome_guerra
        try:
            self.object.delete()
            messages.success(request, f"Militar {nome_guerra} excluído permanentemente.")
        except ProtectedError:
            messages.error(
                request,
                f"Não é possível excluir {nome_guerra} pois ele possui vínculos ativos na seção "
                f"de Informática (Cautela). Remova os registros vinculados antes de excluir o militar."
            )
        return redirect(self.get_success_url())

@method_decorator(s1_required, name='dispatch')
class MilitarTrashListView(ListView):
    model = Efetivo
    template_name = 'Secao_pessoal/militar_trash_list.html'
    context_object_name = 'militares'
    paginate_by = 20

    def get_queryset(self):
        return Efetivo.all_objects.filter(deleted=True).order_by('-deleted_at')

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx['total_lixeira'] = Efetivo.all_objects.filter(deleted=True).count()
        return ctx

@s1_required
@require_POST
def militar_restore(request, pk):
    militar = get_object_or_404(Efetivo.all_objects, pk=pk, deleted=True)
    nome = militar.nome_guerra or militar.nome_completo
    militar.deleted = False
    militar.deleted_at = None
    militar.save(update_fields=['deleted', 'deleted_at'])
    # deleted/deleted_at não estão em campos_monitorados — log explícito necessário
    registrar(
        request.user, secao='pessoal',
        permissao=resolver_label(request.user, _PESSOAL_PERMISSAO_MAP),
        acao='restaurou', descricao=f"restaurou o militar '{nome}' da lixeira",
        objeto_tipo='Efetivo', objeto_id=nome,
    )
    messages.success(request, f"Militar {nome} foi restaurado com sucesso.")
    return redirect('Secao_pessoal:militar_trash_list')

@s1_required
@require_POST
def militar_permanently_delete(request, pk):
    militar = get_object_or_404(Efetivo.all_objects, pk=pk, deleted=True)
    nome = militar.nome_guerra or militar.nome_completo
    militar.delete()
    messages.success(request, f"Militar {nome} foi excluído permanentemente.")
    return redirect('Secao_pessoal:militar_trash_list')

@s1_required
@require_POST
def militar_lixeira_esvaziar(request):
    count = Efetivo.all_objects.filter(deleted=True).count()
    if count:
        # log antes do bulk delete — QuerySet.delete() não dispara post_delete por registro
        registrar(
            request.user, secao='pessoal',
            permissao=resolver_label(request.user, _PESSOAL_PERMISSAO_MAP),
            acao='excluiu', descricao=f"esvaziou a lixeira de militares ({count} registro(s) excluídos permanentemente)",
            objeto_tipo='Efetivo', objeto_id='',
        )
    Efetivo.all_objects.filter(deleted=True).delete()
    messages.success(request, f'{count} militar(es) excluído(s) permanentemente da lixeira.')
    return redirect('Secao_pessoal:militar_trash_list')

@method_decorator(s1_required, name='dispatch')
class EfetivoOperacionalListView(ListView):
    """Efetivo Operacional — Ativo + PSV GSD-GL."""
    model = Efetivo
    context_object_name = 'militares'
    paginate_by = 20

    def get_template_names(self):
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return ['Secao_pessoal/efetivo_operacional_partial.html']
        return ['Secao_pessoal/efetivo_operacional.html']

    def _split(self, param):
        return [v.strip() for v in param.split(',') if v.strip()]

    def get_queryset(self):
        query = self.request.GET.get('q')
        rank_order = Case(
            When(posto='CL', then=Value(0)), When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
            When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)), When(posto='ASP', then=Value(6)), When(posto='SO', then=Value(7)),
            When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)), When(posto='3S', then=Value(10)),
            When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)), When(posto='S2', then=Value(13)), When(posto='REC', then=Value(14)),
            default=Value(99), output_field=IntegerField(),
        )
        qs = super().get_queryset().filter(
            Q(situacao__iexact='Ativo') | Q(situacao__iexact='Ativa') | Q(situacao__iexact='PSV GSD-GL')
        ).annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_completo')
        if query:
            q_objects = Q(nome_completo__icontains=query) | Q(nome_guerra__icontains=query) | Q(posto__icontains=query)
            if query.isdigit():
                q_objects |= Q(saram__icontains=query)
            qs = qs.filter(q_objects)
        # Filtros de coluna
        posto_f = self._split(self.request.GET.get('posto_f', ''))
        esp_f   = self._split(self.request.GET.get('esp_f', ''))
        sit_f   = self._split(self.request.GET.get('sit_f', ''))
        obs_f   = self._split(self.request.GET.get('obs_f', ''))
        setor_f = self._split(self.request.GET.get('setor_f', ''))
        if posto_f:
            qs = qs.filter(posto__in=posto_f)
        if esp_f:
            esp_q = Q()
            for v in esp_f:
                if v == '__vazio__': esp_q |= Q(especializacao='') | Q(especializacao__isnull=True)
                else: esp_q |= Q(especializacao=v)
            qs = qs.filter(esp_q)
        if sit_f:
            sit_q = Q()
            for v in sit_f:
                if v == '__vazio__': sit_q |= Q(situacao='') | Q(situacao__isnull=True)
                else: sit_q |= Q(situacao=v)
            qs = qs.filter(sit_q)
        if obs_f:
            obs_q = Q()
            for v in obs_f:
                if v == '__vazio__': obs_q |= Q(observacao='') | Q(observacao__isnull=True)
                else: obs_q |= Q(observacao=v)
            qs = qs.filter(obs_q)
        if setor_f:
            set_q = Q()
            for v in setor_f:
                if v == '__vazio__': set_q |= Q(setor='') | Q(setor__isnull=True)
                else: set_q |= Q(setor=v)
            qs = qs.filter(set_q)
        return qs


@method_decorator(s1_required, name='dispatch')
class MilitarBaixadoListView(ListView):
    """Junta Médica — exibe militares com situação 'De Junta'."""
    model = Efetivo
    context_object_name = 'militares'
    paginate_by = 20

    def get_template_names(self):
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return ['Secao_pessoal/militar_list_partial.html']
        return ['Secao_pessoal/junta_medica_list.html']

    def _split(self, param):
        return [v.strip() for v in param.split(',') if v.strip()]

    def get_queryset(self):
        query = self.request.GET.get('q')
        rank_order = Case(
            When(posto='CL', then=Value(0)), When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
            When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)), When(posto='ASP', then=Value(6)), When(posto='SO', then=Value(7)),
            When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)), When(posto='3S', then=Value(10)),
            When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)), When(posto='S2', then=Value(13)), When(posto='REC', then=Value(14)),
            default=Value(99), output_field=IntegerField(),
        )
        qs = super().get_queryset().filter(situacao__iexact='De Junta').annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_completo')
        if query:
            q_objects = Q(nome_completo__icontains=query) | Q(nome_guerra__icontains=query) | Q(posto__icontains=query)
            if query.isdigit():
                q_objects |= Q(saram__icontains=query)
            qs = qs.filter(q_objects)
        # Filtros de coluna
        posto_f = self._split(self.request.GET.get('posto_f', ''))
        esp_f   = self._split(self.request.GET.get('esp_f', ''))
        sit_f   = self._split(self.request.GET.get('sit_f', ''))
        obs_f   = self._split(self.request.GET.get('obs_f', ''))
        if posto_f:
            qs = qs.filter(posto__in=posto_f)
        if esp_f:
            esp_q = Q()
            for v in esp_f:
                if v == '__vazio__': esp_q |= Q(especializacao='') | Q(especializacao__isnull=True)
                else: esp_q |= Q(especializacao=v)
            qs = qs.filter(esp_q)
        if sit_f:
            sit_q = Q()
            for v in sit_f:
                if v == '__vazio__': sit_q |= Q(situacao='') | Q(situacao__isnull=True)
                else: sit_q |= Q(situacao=v)
            qs = qs.filter(sit_q)
        if obs_f:
            obs_q = Q()
            for v in obs_f:
                if v == '__vazio__': obs_q |= Q(observacao='') | Q(observacao__isnull=True)
                else: obs_q |= Q(observacao=v)
            qs = qs.filter(obs_q)
        return qs


@method_decorator(s1_required, name='dispatch')
class MilitarDesligadoListView(ListView):
    """Desligados — exibe militares com situação 'Baixado'."""
    model = Efetivo
    context_object_name = 'militares'
    paginate_by = 20

    @staticmethod
    def _split(val):
        return [v for v in (val or '').split(',') if v] if val else []

    def get_template_names(self):
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return ['Secao_pessoal/militar_list_partial.html']
        return ['Secao_pessoal/desligados_list.html']

    def get_queryset(self):
        query = self.request.GET.get('q')
        rank_order = Case(
            When(posto='CL', then=Value(0)), When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
            When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)), When(posto='ASP', then=Value(6)), When(posto='SO', then=Value(7)),
            When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)), When(posto='3S', then=Value(10)),
            When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)), When(posto='S2', then=Value(13)), When(posto='REC', then=Value(14)),
            default=Value(99), output_field=IntegerField(),
        )
        qs = super().get_queryset().filter(situacao__iexact='Baixado').annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_completo')
        if query:
            q_objects = Q(nome_completo__icontains=query) | Q(nome_guerra__icontains=query) | Q(posto__icontains=query)
            if query.isdigit():
                q_objects |= Q(saram__icontains=query)
            qs = qs.filter(q_objects)
        posto_f = self._split(self.request.GET.get('posto_f', ''))
        esp_f   = self._split(self.request.GET.get('esp_f', ''))
        sit_f   = self._split(self.request.GET.get('sit_f', ''))
        obs_f   = self._split(self.request.GET.get('obs_f', ''))
        if posto_f:
            qs = qs.filter(posto__in=posto_f)
        if esp_f:
            esp_q = Q()
            for v in esp_f:
                if v == '__vazio__': esp_q |= Q(especializacao='') | Q(especializacao__isnull=True)
                else: esp_q |= Q(especializacao=v)
            qs = qs.filter(esp_q)
        if sit_f:
            sit_q = Q()
            for v in sit_f:
                if v == '__vazio__': sit_q |= Q(situacao='') | Q(situacao__isnull=True)
                else: sit_q |= Q(situacao=v)
            qs = qs.filter(sit_q)
        if obs_f:
            obs_q = Q()
            for v in obs_f:
                if v == '__vazio__': obs_q |= Q(observacao='') | Q(observacao__isnull=True)
                else: obs_q |= Q(observacao=v)
            qs = qs.filter(obs_q)
        return qs

@s1_required
def movimentar_militar(request):
    rank_order = Case(
        When(posto='CL', then=Value(0)), When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
        When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)),When(posto='ASP', then=Value (6)), When(posto='SO', then=Value(7)),
        When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)), When(posto='3S', then=Value(10)),
        When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)), When(posto='S2', then=Value(13)),When(posto='REC', then=Value(14)),
        default=Value(99), output_field=IntegerField(),
    )
    militares = Efetivo.objects.exclude(situacao__iexact='Baixado').annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_completo')

    if request.method == 'POST':
        militar_id = request.POST.get('militar_movimentacao')
        data_movimentacao = request.POST.get('data_movimentacao')
        om_destino = request.POST.get('om_destino')
        sigad_movimentacao = request.POST.get('sigad_movimentacao')
        boletim_movimentacao = request.POST.get('boletim_movimentacao')
        observacao = request.POST.get('observacao')

        if militar_id and data_movimentacao:
            try:
                militar = Efetivo.objects.get(id=militar_id)
                try:
                    data_mov = datetime.strptime(data_movimentacao, '%Y-%m-%d').date()
                except ValueError:
                    messages.error(request, 'Data da movimentação inválida.')
                    return redirect('Secao_pessoal:movimentar_militar')

                MovimentacaoEfetivo.objects.create(
                    militar=militar,
                    data_movimentacao=data_mov,
                    om_destino=om_destino or '',
                    sigad_movimentacao=sigad_movimentacao or '',
                    boletim_movimentacao=boletim_movimentacao or '',
                    observacao=observacao or '',
                )
                militar.situacao = 'Movimentado'
                militar.save()
                messages.success(request, f'Militar {militar.posto} {militar.nome_guerra} movimentado com sucesso.')
            except Efetivo.DoesNotExist:
                messages.error(request, 'Erro: Militar não encontrado.')
        else:
            messages.error(request, 'Selecione o militar e informe a data da movimentação.')

        return redirect('Secao_pessoal:movimentar_militar')

    context = {
        'militares': militares
    }
    return render(request, 'Secao_pessoal/movimentar_militar.html', context)

@method_decorator(s1_required, name='dispatch')
class MilitarMovimentadoListView(ListView):
    model = Efetivo
    template_name = 'Secao_pessoal/militar_movimentado_list.html'
    context_object_name = 'militares'
    paginate_by = 20

    def get_template_names(self):
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return ['Secao_pessoal/militar_movimentado_list_partial.html']
        return ['Secao_pessoal/militar_movimentado_list.html']

    def get_queryset(self):
        query = self.request.GET.get('q')
        rank_order = Case(
            When(posto='CL', then=Value(0)), When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
            When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)),When(posto='ASP', then=Value (6)), When(posto='SO', then=Value(7)),
            When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)), When(posto='3S', then=Value(10)),
            When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)), When(posto='S2', then=Value(13)),When(posto='REC', then=Value(14)),
            default=Value(99), output_field=IntegerField(),
        )
        qs = super().get_queryset().filter(situacao__iexact='Movimentado').annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_completo')
        if query:
            q_objects = Q(nome_completo__icontains=query) | \
                        Q(nome_guerra__icontains=query) | \
                        Q(posto__icontains=query)
            if query.isdigit():
                q_objects |= Q(saram__icontains=query)
            qs = qs.filter(q_objects)
        return qs

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        militares = ctx.get('militares') or []
        militares_ids = [m.pk for m in militares]
        ultimas = {}
        for mov in MovimentacaoEfetivo.objects.filter(militar_id__in=militares_ids).order_by('militar_id', '-data_movimentacao'):
            ultimas.setdefault(mov.militar_id, mov)
        for militar in militares:
            militar.ultima_movimentacao = ultimas.get(militar.pk)
        return ctx

@s1_required
def gerar_ficha_desimpedimento(request, pk):
    militar = get_object_or_404(Efetivo.all_objects, pk=pk)

    document = docx.Document()
    section = document.sections[0]

    def add_centered(text, bold=False, size=12):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    add_centered('MINISTÉRIO DA DEFESA', bold=True)
    add_centered('COMANDO DA AERONÁUTICA', bold=True)
    add_centered(militar.om or '', bold=True)
    add_centered('ESQUADRÃO DE PESSOAL', bold=True)
    add_centered('FICHA DE DESIMPEDIMENTO', bold=True, size=14)
    document.add_paragraph('')

    def add_field(label, valor):
        p = document.add_paragraph()
        run_label = p.add_run(f'{label}: ')
        run_label.bold = True
        p.add_run(str(valor) if valor else '—')

    add_field('Nome do Militar', militar.nome_completo)
    add_field('Posto/Grad', militar.posto)
    add_field('Documento de Publicação', militar.documento_desligamento)
    add_field('OM', militar.om)
    add_field('Motivo do Desligamento', militar.motivo_desligamento)
    add_field('SARAM', militar.saram)
    add_field('Função (conforme publicado)', militar.funcao_desligamento)
    add_field('Setor', militar.setor)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    nome_arquivo = f"ficha_desimpedimento_{militar.nome_guerra or militar.pk}.docx".replace(' ', '_')
    response['Content-Disposition'] = f'attachment; filename="{nome_arquivo}"'
    document.save(response)
    return response

#EFETIVO IMPORT EXCEL
@s1_required
def importar_excel(request):
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        sincronizar = request.POST.get('sincronizar') == '1'
        postos_filtro_str = request.POST.get('postos_filtro', 'todos').strip()
        if postos_filtro_str and postos_filtro_str != 'todos':
            postos_filtro = set(postos_filtro_str.split(','))
        else:
            postos_filtro = None

        try:
            # Lê o arquivo Excel e converte tudo para string para evitar erros de tipo.
            # header=2 → linhas 1-2 são título, a linha 3 é o cabeçalho real (novo modelo).
            # Tenta header=2 primeiro; se não encontrar colunas conhecidas, tenta header=0
            # para compatibilidade com planilhas antigas.
            df = pd.read_excel(excel_file, dtype=str, header=2)
            df.columns = df.columns.str.strip()
            # Detecta se o cabeçalho correto foi lido: deve ter ao menos SARAM ou NOME COMPLETO
            _cols_norm = {re.sub(r'\s+', ' ', normalize_name(str(c))).strip() for c in df.columns}
            if 'SARAM' not in _cols_norm and 'NOME COMPLETO' not in _cols_norm:
                # Fallback: planilha antiga sem linhas de título
                df = pd.read_excel(excel_file, dtype=str, header=0)
                df.columns = df.columns.str.strip()

            # Renomeia as colunas para os nomes canônicos esperados abaixo, aceitando
            # variações de acentuação/maiúsculas/abreviação/pontuação no cabeçalho da
            # planilha. 'ESPC' é a abreviação usada no novo modelo; 'OBSERVAÇÕES' é nova.
            ALIASES_COLUNAS_EFETIVO = {
                'SARAM': ['SARAM'],
                'NOME COMPLETO': ['NOME COMPLETO', 'NOME'],
                'PST.': ['PST', 'POSTO', 'POSTOGRAD', 'POSTO GRAD'],
                'QUAD.': ['QUAD', 'QUADRO'],
                'ESP.': ['ESP', 'ESPC', 'ESPECIALIZACAO', 'ESPECIALIDADE'],
                'NOME DE GUERRA': ['NOME DE GUERRA', 'NOME GUERRA', 'GUERRA'],
                'TURMA': ['TURMA'],
                'SITUAÇÃO': ['SITUACAO'],
                'OM': ['OM'],
                'SETOR': ['SETOR'],
                'SUBSETOR': ['SUBSETOR'],
                'OBSERVAÇÃO': ['OBSERVACOES', 'OBSERVACAO', 'OBSERVACOES'],
            }
            colunas_normalizadas = {
                re.sub(r'\s+', ' ', re.sub(r'[.\/]', '', normalize_name(str(c)))).strip(): c
                for c in df.columns
            }
            renomear = {}
            for canonico, aliases in ALIASES_COLUNAS_EFETIVO.items():
                for alias in aliases:
                    col_real = colunas_normalizadas.get(alias)
                    if col_real:
                        renomear[col_real] = canonico
                        break
            df.rename(columns=renomear, inplace=True)

            # Substitui valores 'nan' (vazios) do pandas por string vazia
            df.fillna('', inplace=True)

            criados = 0
            atualizados = 0

            # Coleta valores únicos para criar no controle geral ao final
            postos_excel = set()
            quads_excel = set()
            especs_excel = set()
            oms_excel = set()
            setores_excel = set()
            subsetores_excel = set()

            # Mapa de militares já existentes sem SARAM confiável, indexado pelo nome
            # normalizado (sem acentos, maiúsculas, espaços colapsados), para evitar
            # criar registros duplicados quando o nome da planilha não é idêntico
            # byte-a-byte ao nome já salvo no banco.
            existentes_por_nome = {}
            for efetivo_existente in Efetivo.all_objects.all():
                chave = normalize_name(efetivo_existente.nome_completo).strip()
                if chave:
                    existentes_por_nome[chave] = efetivo_existente

            # IDs dos militares encontrados/criados nesta planilha, usados pelo modo
            # "importar e sincronizar" para saber quem NÃO está na planilha.
            pks_na_planilha = set()

            # Linhas com posto TEN/TENENTE que precisam de clarificação
            ten_rows_efetivo = []

            for index, row in df.iterrows():
                # Pega valores essenciais
                saram_valor = str(row.get('SARAM', '')).strip()
                nome_completo_valor = row.get('NOME COMPLETO', '').strip()

                # Se não tem SARAM E não tem nome, aí sim a linha é inútil e pulamos
                if not saram_valor and not nome_completo_valor:
                    continue

                # Tratamento do SARAM para salvar no banco (IntegerField exige número ou None)
                saram_db = None
                if saram_valor:
                    try:
                        # Converte de "12345.0" para 12345 caso o pandas tenha lido como float
                        saram_db = int(float(saram_valor))
                    except ValueError:
                        saram_db = None

                # Extrai nome_guerra e strip de posto prefixo
                nome_guerra_raw = row.get('NOME DE GUERRA', '').strip()
                posto_raw = row.get('PST.', '').strip()

                nome_guerra_limpo, posto_de_nome = _strip_posto_de_nome(nome_guerra_raw)

                # Determina posto final
                if posto_raw:
                    posto_norm, is_ten = _normalizar_posto(posto_raw)
                    if posto_norm is None:
                        posto_norm = posto_raw  # fallback (sem normalização)
                else:
                    # Usa posto detectado no nome_guerra se posto_raw estiver vazio
                    posto_norm = posto_de_nome or ''
                    is_ten = False

                # Pula postos não selecionados pelo usuário
                if postos_filtro and posto_norm not in postos_filtro:
                    continue

                # Se posto é TEN/TENENTE, guarda para clarificação e passa para a próxima
                if is_ten:
                    if postos_filtro and posto_norm not in postos_filtro:
                        continue
                    ten_rows_efetivo.append({
                        'saram': saram_valor,
                        'nome_completo': nome_completo_valor,
                        'nome_guerra': nome_guerra_limpo,
                        'posto_raw': posto_raw,
                        'quad': row.get('QUAD.', '').strip(),
                        'especializacao': row.get('ESP.', '').strip(),
                        'turma': _extrair_ano_turma(row.get('TURMA', '')),
                        'situacao': row.get('SITUAÇÃO', '').strip(),
                        'om': row.get('OM', '').strip(),
                        'setor': row.get('SETOR', '').strip(),
                        'subsetor': row.get('SUBSETOR', '').strip(),
                        'observacao': row.get('OBSERVAÇÃO', '').strip(),
                        'saram_db': saram_db,
                        'is_ten': True,
                    })
                    continue

                # Dicionário com os dados a serem salvos/atualizados
                dados_militar = {
                    'posto': posto_norm,
                    'quad': row.get('QUAD.', '').strip(),
                    'especializacao': row.get('ESP.', '').strip(),
                    'saram': saram_db,
                    'nome_completo': nome_completo_valor,
                    'nome_guerra': nome_guerra_limpo,
                    'turma': _extrair_ano_turma(row.get('TURMA', '')),
                    'situacao': row.get('SITUAÇÃO', '').strip(),
                    'om': row.get('OM', '').strip(),
                    'setor': row.get('SETOR', '').strip(),
                    'subsetor': row.get('SUBSETOR', '').strip(),
                    'observacao': row.get('OBSERVAÇÃO', '').strip(),
                    # Se o militar estava na lixeira, reimportá-lo via planilha deve
                    # restaurá-lo para o efetivo ativo — caso contrário ele continua
                    # "atualizado" mas invisível na lista geral.
                    'deleted': False,
                    'deleted_at': None,
                }

                # Acumula valores não-vazios para criar no controle geral
                if dados_militar['posto']:
                    postos_excel.add(dados_militar['posto'])
                if dados_militar['quad']:
                    quads_excel.add(dados_militar['quad'])
                if dados_militar['especializacao']:
                    especs_excel.add(dados_militar['especializacao'])
                if dados_militar['om']:
                    oms_excel.add(dados_militar['om'])
                if dados_militar['setor']:
                    setores_excel.add(dados_militar['setor'])
                if dados_militar['subsetor']:
                    subsetores_excel.add(dados_militar['subsetor'])

                # LÓGICA DE SALVAMENTO INTELIGENTE:
                if saram_db:
                    # Se tem SARAM, a chave principal é o SARAM
                    obj, created = Efetivo.all_objects.update_or_create(
                        saram=saram_db,
                        defaults=dados_militar
                    )
                else:
                    # Se é um recruta sem SARAM, casa pelo nome normalizado (ignora
                    # acentos/maiúsculas/espaços) para não criar duplicado quando o
                    # nome da planilha não é idêntico ao já salvo no banco.
                    chave_nome = normalize_name(nome_completo_valor).strip()
                    existente = existentes_por_nome.get(chave_nome)
                    if existente:
                        for campo, valor in dados_militar.items():
                            setattr(existente, campo, valor)
                        existente.save()
                        obj, created = existente, False
                    else:
                        obj = Efetivo.all_objects.create(**dados_militar)
                        created = True
                        existentes_por_nome[chave_nome] = obj

                if created:
                    criados += 1
                else:
                    atualizados += 1

                pks_na_planilha.add(obj.pk)

            removidos = 0
            sincronizar_abortada = False
            if sincronizar:
                if criados == 0 and atualizados == 0:
                    # Nenhuma linha da planilha foi reconhecida (provável incompatibilidade
                    # de colunas) — não sincroniza para não apagar todo o efetivo existente.
                    sincronizar_abortada = True
                else:
                    agora = timezone.now()
                    nao_encontrados = Efetivo.objects.exclude(pk__in=pks_na_planilha)
                    removidos = nao_encontrados.update(deleted=True, deleted_at=agora)
                    if removidos:
                        # bulk update não dispara signal — log explícito
                        registrar(
                            request.user, secao='pessoal',
                            permissao=resolver_label(request.user, _PESSOAL_PERMISSAO_MAP),
                            acao='excluiu',
                            descricao=f"importação Excel (sincronizar): {removidos} militar(es) movidos para lixeira por não constar na planilha",
                            objeto_tipo='Efetivo', objeto_id='',
                        )

            # Cria opções do controle geral sem duplicar (get_or_create é idempotente)
            for v in postos_excel:
                Posto.objects.get_or_create(nome=v)
            for v in quads_excel:
                Quad.objects.get_or_create(nome=v)
            for v in especs_excel:
                Especializacao.objects.get_or_create(nome=v)
            for v in oms_excel:
                OM.objects.get_or_create(nome=v)
            for v in setores_excel:
                Setor.objects.get_or_create(nome=v)
            for v in subsetores_excel:
                Subsetor.objects.get_or_create(nome=v)

            # Se há linhas com TEN/TENENTE ambíguo, redireciona para clarificação
            if ten_rows_efetivo:
                request.session['ten_rows_efetivo'] = ten_rows_efetivo
                request.session['efetivo_criados'] = criados
                request.session['efetivo_atualizados'] = atualizados
                request.session['efetivo_pks_na_planilha'] = list(pks_na_planilha)
                request.session['efetivo_sincronizar'] = sincronizar
                request.session['efetivo_postos'] = list(postos_excel)
                request.session['efetivo_quads'] = list(quads_excel)
                request.session['efetivo_especs'] = list(especs_excel)
                request.session['efetivo_oms'] = list(oms_excel)
                request.session['efetivo_setores'] = list(setores_excel)
                request.session['efetivo_subsetores'] = list(subsetores_excel)
                messages.info(request, f'{criados + atualizados} militar(es) processado(s). Esclareça o posto dos {len(ten_rows_efetivo)} militar(es) com TEN/TENENTE.')
                return redirect('Secao_pessoal:esclarecer_importacao_efetivo')

            if sincronizar_abortada:
                messages.error(request, 'Nenhum militar foi reconhecido na planilha (verifique se as colunas "SARAM" e "NOME COMPLETO" estão corretas). A sincronização foi cancelada para evitar apagar o efetivo existente.')
            elif sincronizar:
                messages.success(request, f'Sucesso! {criados} militares criados, {atualizados} atualizados e {removidos} removidos (enviados para a lixeira) por não constarem na planilha.')
            else:
                messages.success(request, f'Sucesso! {criados} militares criados e {atualizados} atualizados.')
            return redirect('Secao_pessoal:militar_list')

        except Exception as e:
            messages.error(request, f'Erro na importação: {str(e)}')
            return redirect('Secao_pessoal:importar_excel')

    return render(request, 'Secao_pessoal/importar_excel.html')


@s1_required
@require_POST
def preview_postos_excel(request):
    """Lê o arquivo Excel e retorna os postos encontrados (sem importar)."""
    excel_file = request.FILES.get('excel_file')
    if not excel_file:
        return JsonResponse({'error': 'Nenhum arquivo enviado'}, status=400)
    try:
        df = pd.read_excel(excel_file, dtype=str, header=2)
        df.columns = df.columns.str.strip()
        _cols = {re.sub(r'\s+', ' ', normalize_name(str(c))).strip() for c in df.columns}
        if 'SARAM' not in _cols and 'NOME COMPLETO' not in _cols:
            df = pd.read_excel(excel_file, dtype=str, header=0)
            df.columns = df.columns.str.strip()
        colunas_norm = {
            re.sub(r'\s+', ' ', re.sub(r'[.\/]', '', normalize_name(str(c)))).strip(): c
            for c in df.columns
        }
        posto_col = None
        for alias in ['PST', 'POSTO', 'POSTOGRAD', 'POSTO GRAD']:
            if alias in colunas_norm:
                posto_col = colunas_norm[alias]
                break
        if not posto_col:
            return JsonResponse({'postos': []})
        df.fillna('', inplace=True)
        hierarquia = ['TC', 'MJ', 'CP', '1T', '2T', 'ASP', 'SO', '1S', '2S', '3S', 'CB', 'S1', 'S2', 'REC']
        postos_norm = set()
        for p_raw in df[posto_col].unique():
            p_raw = str(p_raw).strip()
            if p_raw:
                p_norm, _ = _normalizar_posto(p_raw)
                if p_norm:
                    postos_norm.add(p_norm)
        postos_ord = sorted(postos_norm, key=lambda x: hierarquia.index(x) if x in hierarquia else 99)
        return JsonResponse({'postos': postos_ord})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@s1_required
def esclarecer_importacao_efetivo(request):
    ten_rows = request.session.get('ten_rows_efetivo', [])
    if not ten_rows:
        return redirect('Secao_pessoal:militar_list')

    if request.method == 'POST':
        criados    = request.session.pop('efetivo_criados', 0)
        atualizados= request.session.pop('efetivo_atualizados', 0)
        pks_na_planilha = set(request.session.pop('efetivo_pks_na_planilha', []))
        sincronizar = request.session.pop('efetivo_sincronizar', False)
        postos_excel   = set(request.session.pop('efetivo_postos', []))
        quads_excel    = set(request.session.pop('efetivo_quads', []))
        especs_excel   = set(request.session.pop('efetivo_especs', []))
        oms_excel      = set(request.session.pop('efetivo_oms', []))
        setores_excel  = set(request.session.pop('efetivo_setores', []))
        subsetores_excel = set(request.session.pop('efetivo_subsetores', []))
        request.session.pop('ten_rows_efetivo', None)

        existentes_por_nome = {}
        for efetivo_existente in Efetivo.all_objects.all():
            chave = normalize_name(efetivo_existente.nome_completo).strip()
            if chave:
                existentes_por_nome[chave] = efetivo_existente

        for i, r in enumerate(ten_rows):
            posto_escolhido = request.POST.get(f'posto_{i}', '')
            saram_db = r.get('saram_db')
            dados_militar = {
                'posto': posto_escolhido,
                'quad': r.get('quad', ''),
                'especializacao': r.get('especializacao', ''),
                'saram': saram_db,
                'nome_completo': r.get('nome_completo', ''),
                'nome_guerra': r.get('nome_guerra', ''),
                'turma': r.get('turma', ''),
                'situacao': r.get('situacao', ''),
                'om': r.get('om', ''),
                'setor': r.get('setor', ''),
                'subsetor': r.get('subsetor', ''),
                'observacao': r.get('observacao', ''),
                'deleted': False,
                'deleted_at': None,
            }
            if posto_escolhido:
                postos_excel.add(posto_escolhido)

            if saram_db:
                obj, created = Efetivo.all_objects.update_or_create(saram=saram_db, defaults=dados_militar)
            else:
                chave_nome = normalize_name(r.get('nome_completo', '')).strip()
                existente = existentes_por_nome.get(chave_nome)
                if existente:
                    for campo, valor in dados_militar.items():
                        setattr(existente, campo, valor)
                    existente.save()
                    obj, created = existente, False
                else:
                    obj = Efetivo.all_objects.create(**dados_militar)
                    created = True
                    existentes_por_nome[chave_nome] = obj
            if created:
                criados += 1
            else:
                atualizados += 1
            pks_na_planilha.add(obj.pk)

        # Sincronizar se necessário
        removidos = 0
        if sincronizar and pks_na_planilha:
            agora = timezone.now()
            nao_enc = Efetivo.objects.exclude(pk__in=pks_na_planilha)
            removidos = nao_enc.update(deleted=True, deleted_at=agora)

        for v in postos_excel: Posto.objects.get_or_create(nome=v)
        for v in quads_excel: Quad.objects.get_or_create(nome=v)
        for v in especs_excel: Especializacao.objects.get_or_create(nome=v)
        for v in oms_excel: OM.objects.get_or_create(nome=v)
        for v in setores_excel: Setor.objects.get_or_create(nome=v)
        for v in subsetores_excel: Subsetor.objects.get_or_create(nome=v)

        if sincronizar and removidos:
            messages.success(request, f'Sucesso! {criados} militares criados, {atualizados} atualizados e {removidos} removidos por não constarem na planilha.')
        else:
            messages.success(request, f'Sucesso! {criados} militares criados e {atualizados} atualizados.')
        return redirect('Secao_pessoal:militar_list')

    return render(request, 'Secao_pessoal/esclarecer_importacao.html', {'ambiguous': ten_rows, 'fonte': 'efetivo'})


# --- Lógica para Nome de Guerra ---

def normalize_name(name):
    """Remove acentos e caracteres especiais para comparação (ex: 'Corrêa' -> 'CORREA')."""
    if not name:
        return ""
    return ''.join(c for c in unicodedata.normalize('NFD', name) if unicodedata.category(c) != 'Mn').upper()

def gerar_sugestoes_guerra(nome_completo):
    """Gera uma lista de possíveis nomes de guerra baseados no nome completo, em ordem de prioridade."""
    ignore = ['de', 'da', 'do', 'dos', 'das', 'e']
    partes = [p for p in nome_completo.split() if p.lower() not in ignore]
    
    if not partes:
        return []
    
    sugestoes = []
    
    def add_unique(nome):
        nome = nome.upper()
        if nome not in sugestoes:
            sugestoes.append(nome)
    
    # 1. Último nome (Padrão)
    add_unique(partes[-1])
    
    # 2. Penúltimo (se houver e não for o primeiro)
    if len(partes) > 2:
        add_unique(partes[-2])

    # 3. Nomes Compostos (ex: Del Puerto, Villas Boas)
    if len(partes) >= 2:
        add_unique(f"{partes[-2]} {partes[-1]}")

    # 4. Primeiro + Último (Prioriza abreviação se o nome for longo)
    primeiro_ultimo = f"{partes[0]} {partes[-1]}"
    inicial_ultimo = f"{partes[0][0]}. {partes[-1]}"
    
    if len(partes[0]) > 5: # Ex: Leonardo (8) -> L. MARTINS
        add_unique(inicial_ultimo)
        add_unique(primeiro_ultimo)
    else:
        add_unique(primeiro_ultimo)
        add_unique(inicial_ultimo)
    
    # 5. Outras combinações
    if len(partes) > 2:
        primeiro_penultimo = f"{partes[0]} {partes[-2]}"
        inicial_penultimo = f"{partes[0][0]}. {partes[-2]}"
        
        if len(partes[0]) > 5:
            add_unique(inicial_penultimo)
            add_unique(primeiro_penultimo)
        else:
            add_unique(primeiro_penultimo)
            add_unique(inicial_penultimo)

    for i in range(1, len(partes)-1):
        add_unique(partes[i])
        # Adiciona combinações com nomes do meio, também verificando tamanho
        comb_nome = f"{partes[0]} {partes[i]}"
        comb_inicial = f"{partes[0][0]}. {partes[i]}"
        
        if len(partes[0]) > 5:
            add_unique(comb_inicial)
            add_unique(comb_nome)
        else:
            add_unique(comb_nome)
            add_unique(comb_inicial)

    return sugestoes

@s1_required
def nome_de_guerra(request):
    if request.method == 'POST':
        try:
            # 1. Buscar Recrutas no Banco de Dados
            recrutas = Efetivo.objects.filter(posto='REC').exclude(situacao__iexact='Baixado')
            
            if not recrutas.exists():
                messages.warning(request, "Nenhum militar com posto 'REC' encontrado no efetivo.")
                return redirect('Secao_pessoal:nome_de_guerra')

            # Carrega todos os nomes existentes e normaliza para comparação rápida e sem acentos
            existing_names_map = {}
            all_efetivo = Efetivo.objects.exclude(nome_guerra__isnull=True).exclude(nome_guerra='').values('id', 'nome_guerra')
            
            for item in all_efetivo:
                norm = normalize_name(item['nome_guerra'])
                if norm not in existing_names_map:
                    existing_names_map[norm] = set()
                existing_names_map[norm].add(item['id'])

            nomes_usados_sessao_norm = set() # Rastreia nomes gerados nesta sessão (normalizados)
            count_atualizados = 0

            for recruta in recrutas:
                sugestoes = gerar_sugestoes_guerra(recruta.nome_completo)
                nome_final = None
                
                for sugestao in sugestoes:
                    sugestao_norm = normalize_name(sugestao)
                    
                    # Verifica conflito no DB (se existe e não é o próprio recruta)
                    ids_com_esse_nome = existing_names_map.get(sugestao_norm, set())
                    conflito_db = bool(ids_com_esse_nome - {recruta.pk})
                    
                    # Verifica conflito na sessão atual
                    conflito_sessao = sugestao_norm in nomes_usados_sessao_norm

                    if not conflito_db and not conflito_sessao:
                        nome_final = sugestao
                        break
                
                if not nome_final:
                    base = sugestoes[0] if sugestoes else "MILITAR"
                    contador = 2
                    while True:
                        teste = f"{base} {contador}"
                        teste_norm = normalize_name(teste)
                        
                        ids_com_esse_nome = existing_names_map.get(teste_norm, set())
                        conflito_db = bool(ids_com_esse_nome - {recruta.pk})
                        conflito_sessao = teste_norm in nomes_usados_sessao_norm

                        if not conflito_db and not conflito_sessao:
                            nome_final = teste
                            break
                        contador += 1

                nomes_usados_sessao_norm.add(normalize_name(nome_final))

                # Atualiza diretamente no banco de dados
                recruta.nome_guerra = nome_final
                recruta.save()
                count_atualizados += 1
            
            messages.success(request, f"Sucesso! Nomes de guerra atualizados para {count_atualizados} recrutas.")
            return redirect('Secao_pessoal:nome_de_guerra')
            
        except Exception as e:
            messages.error(request, f"Erro ao gerar nomes: {str(e)}")
            return redirect('Secao_pessoal:nome_de_guerra')

    return render(request, 'Secao_pessoal/nome_de_guerra.html')

def comunicacoes(request):
    """Redirecionamento de compatibilidade — caixa de entrada movida para caixa_entrada app."""
    from django.shortcuts import redirect
    return redirect('caixa_entrada:comunicacoes')

@s1_required
def troca_de_setor(request):
    militares = Efetivo.objects.exclude(situacao__iexact='Baixado').order_by('nome_guerra')
    setores = Setor.objects.all().order_by('nome')

    if request.method == 'POST':
        militar_id = request.POST.get('militar_id')
        setor_destino = request.POST.get('setor_destino')

        if militar_id and setor_destino:
            militar = get_object_or_404(Efetivo, id=militar_id)
            setor_atual_str = militar.setor or ''
            
            def get_chefe_por_grupo(nome_setor):
                if not nome_setor:
                    return None
                nome_grupo = f"Chefe - {nome_setor}"
                try:
                    grupo = Group.objects.get(name__iexact=nome_grupo)
                    user_chefe = grupo.user_set.filter(is_active=True, profile__militar__isnull=False).first()
                    if user_chefe:
                        return user_chefe.profile.militar
                except Group.DoesNotExist:
                    pass
                return None
            
            chefe_atual = get_chefe_por_grupo(setor_atual_str)
            if not chefe_atual:
                setor_atual_obj = Setor.objects.filter(nome=setor_atual_str).first()
                chefe_atual = setor_atual_obj.chefe if setor_atual_obj else None
            
            chefe_destino = get_chefe_por_grupo(setor_destino)
            if not chefe_destino:
                setor_destino_obj = Setor.objects.filter(nome=setor_destino).first()
                chefe_destino = setor_destino_obj.chefe if setor_destino_obj else None
            
            if not chefe_destino:
                messages.error(request, f'O setor de destino ({setor_destino}) não possui um chefe configurado no sistema (Grupo "Chefe - {setor_destino}"). Por favor, defina um chefe para este setor na página de Administração.')
                return redirect('Secao_pessoal:troca_de_setor')
                
            status_inicial = 'pendente_atual' if chefe_atual else 'pendente_destino'
            
            solicitacao = SolicitacaoTrocaSetor.objects.create(
                militar=militar,
                setor_atual=setor_atual_str or 'Não definido',
                setor_destino=setor_destino,
                chefe_atual=chefe_atual,
                chefe_destino=chefe_destino,
                status=status_inicial
            )
            
            try:
                S1_militar = request.user.profile.militar
                if not S1_militar:
                    S1_militar = chefe_destino
            except:
                S1_militar = chefe_destino

            if chefe_atual:
                messages.success(request, f'Solicitação criada. Aguardando autorização do chefe atual ({chefe_atual.posto} {chefe_atual.nome_guerra}).')
            else:
                messages.success(request, f'Militar sem chefe atual. Solicitação enviada diretamente para autorização do chefe de destino ({chefe_destino.posto} {chefe_destino.nome_guerra}).')
                
            return redirect('Secao_pessoal:troca_de_setor')
        else:
            messages.error(request, 'Preencha todos os campos corretamente.')

    solicitacoes = SolicitacaoTrocaSetor.objects.all()

    context = {
        'militares': militares,
        'setores': setores,
        'solicitacoes': solicitacoes
    }
    return render(request, 'Secao_pessoal/troca_de_setor.html', context)

@login_required
@xframe_options_sameorigin
def responder_troca_setor(request, solicitacao_id, acao):
    solicitacao = get_object_or_404(SolicitacaoTrocaSetor, id=solicitacao_id)
    militar_logado = None
    try:
        if hasattr(request.user, 'profile'):
            militar_logado = request.user.profile.militar
    except:
        pass
        
    if not militar_logado and not request.user.is_superuser:
        messages.error(request, "Seu usuário não está vinculado a um militar.")
        return redirect('caixa_entrada:comunicacoes')
    
    if solicitacao.status == 'pendente_atual':
        if solicitacao.chefe_atual != militar_logado and not is_s1_member(request.user) and not request.user.is_superuser:
            messages.error(request, 'Você não tem permissão para responder a esta solicitação.')
            return redirect('caixa_entrada:comunicacoes')
            
        if acao == 'aprovar':
            solicitacao.status = 'pendente_destino'
            solicitacao.save()
            if solicitacao.chefe_destino:
                _enviar_mensagem_sistema(
                    militar_logado, solicitacao.chefe_destino,
                    f"Autorização pendente: troca de setor de {solicitacao.militar.nome_guerra}",
                    f"A saída do {solicitacao.militar.posto} {solicitacao.militar.nome_guerra} do setor {solicitacao.setor_atual} foi autorizada. Aguarda sua autorização de entrada no setor {solicitacao.setor_destino}.",
                )
            messages.success(request, 'Saída autorizada com sucesso. Aguardando autorização do setor de destino.')
        elif acao == 'rejeitar':
            solicitacao.status = 'rejeitado'
            solicitacao.save()
            _enviar_mensagem_sistema(
                militar_logado, solicitacao.militar,
                f"Solicitação de troca de setor rejeitada",
                f"Sua solicitação de troca do setor {solicitacao.setor_atual} para {solicitacao.setor_destino} foi rejeitada pelo {militar_logado.posto} {militar_logado.nome_guerra}.",
            )
            messages.success(request, 'Solicitação de troca rejeitada.')

    elif solicitacao.status == 'pendente_destino':
        if solicitacao.chefe_destino != militar_logado and not is_s1_member(request.user) and not request.user.is_superuser:
            messages.error(request, 'Você não tem permissão para responder a esta solicitação.')
            return redirect('caixa_entrada:comunicacoes')

        if acao == 'aprovar':
            solicitacao.status = 'aprovado'
            solicitacao.save()

            militar = solicitacao.militar
            militar.setor = solicitacao.setor_destino
            militar.save()
            _enviar_mensagem_sistema(
                militar_logado, solicitacao.militar,
                f"Troca de setor aprovada",
                f"Sua transferência para o setor {solicitacao.setor_destino} foi aprovada pelo {militar_logado.posto} {militar_logado.nome_guerra}. Você já foi movido para o novo setor.",
            )
            messages.success(request, 'Entrada autorizada com sucesso. O militar foi transferido de setor.')
        elif acao == 'rejeitar':
            solicitacao.status = 'rejeitado'
            solicitacao.save()
            _enviar_mensagem_sistema(
                militar_logado, solicitacao.militar,
                f"Solicitação de troca de setor rejeitada",
                f"Sua solicitação de entrada no setor {solicitacao.setor_destino} foi rejeitada pelo {militar_logado.posto} {militar_logado.nome_guerra}.",
            )
            messages.success(request, 'Solicitação de troca rejeitada.')
            
    return redirect(request.META.get('HTTP_REFERER', 'caixa_entrada:comunicacoes'))

POSTOS_DESIMPEDIMENTO_PRACA = ['CB', 'S1', 'S2']
POSTOS_DESIMPEDIMENTO_GRADUADO_OF = ['TC', 'MJ', 'CP', '1T', '2T', 'ASP', 'SO', '1S', '2S', '3S']


@s1_required
def desimpedimento_busca(request):
    rank_order = Case(
        When(posto='CL', then=Value(0)), When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
        When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)),When(posto='ASP', then=Value (6)), When(posto='SO', then=Value(7)),
        When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)), When(posto='3S', then=Value(10)),
        When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)), When(posto='S2', then=Value(13)),When(posto='REC', then=Value(14)),
        default=Value(99), output_field=IntegerField(),
    )
    query = request.GET.get('q')
    tab = request.GET.get('tab') if request.GET.get('tab') in ('praca', 'graduado') else 'praca'
    postos = POSTOS_DESIMPEDIMENTO_GRADUADO_OF if tab == 'graduado' else POSTOS_DESIMPEDIMENTO_PRACA

    militares = Efetivo.objects.filter(posto__in=postos).annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_completo')
    if query:
        q_objects = Q(nome_completo__icontains=query) | Q(nome_guerra__icontains=query) | Q(posto__icontains=query)
        if query.isdigit():
            q_objects |= Q(saram__icontains=query)
        militares = militares.filter(q_objects)

    paginator = Paginator(militares, 20)
    page_obj = paginator.get_page(request.GET.get('page'))

    template_name = 'Secao_pessoal/desimpedimento_busca_partial.html' if request.headers.get('x-requested-with') == 'XMLHttpRequest' else 'Secao_pessoal/desimpedimento_busca.html'
    return render(request, template_name, {
        'militares': page_obj,
        'page_obj': page_obj,
        'is_paginated': page_obj.has_other_pages(),
        'tab': tab,
    })


@s1_required
def desimpedimento_form(request, pk):
    militar = get_object_or_404(Efetivo, pk=pk)
    return render(request, 'Secao_pessoal/desimpedimento_form.html', {'militar': militar})


@s1_required
@require_POST
def desimpedimento_gerar_pdf(request, pk):
    from django.contrib.staticfiles import finders
    import shutil
    import subprocess

    militar = get_object_or_404(Efetivo, pk=pk)

    documento_desligamento = request.POST.get('documento_desligamento', '')
    motivo_desligamento = request.POST.get('motivo_desligamento', '')
    data_desligamento_str = request.POST.get('data_desligamento', '')
    salvar_no_cadastro = request.POST.get('salvar_no_cadastro') == '1'

    data_desligamento = None
    if data_desligamento_str:
        try:
            data_desligamento = datetime.strptime(data_desligamento_str, '%Y-%m-%d').date()
        except ValueError:
            pass

    if salvar_no_cadastro:
        militar.documento_desligamento = documento_desligamento
        militar.motivo_desligamento = motivo_desligamento
        militar.data_desligamento = data_desligamento
        militar.save()

    if militar.posto in POSTOS_DESIMPEDIMENTO_GRADUADO_OF:
        template_filename = 'ficha_desimpedimento_graduado_of_template.xlsx'
        print_area_last_row = 35
        image_offset_emu = 26543
    else:
        template_filename = 'ficha_desimpedimento_template.xlsx'
        print_area_last_row = 31
        image_offset_emu = 156210

    template_path = finders.find(f'Secao_pessoal/templates_pdf/{template_filename}')
    if not template_path:
        messages.error(request, 'Modelo de Ficha de Desimpedimento não encontrado.')
        return redirect('Secao_pessoal:desimpedimento_form', pk=pk)

    tmp_dir = tempfile.mkdtemp()
    xlsx_path = os.path.join(tmp_dir, 'ficha_preenchida.xlsx')
    pdf_path = os.path.join(tmp_dir, 'ficha_preenchida.pdf')
    shutil.copyfile(template_path, xlsx_path)

    wb = openpyxl.load_workbook(xlsx_path)
    ws = wb.worksheets[0]
    ws['C11'] = militar.nome_completo or ''
    ws['F11'] = f"{militar.posto or ''} {militar.quad or ''}".strip()
    ws['C12'] = documento_desligamento
    ws['F12'] = militar.om or ''
    ws['C13'] = motivo_desligamento
    ws['F13'] = militar.saram or ''
    ws['C14'] = militar.setor or ''
    ws['F14'] = militar.setor or ''
    ws['F15'] = data_desligamento.strftime('%d/%m/%Y') if data_desligamento else ''

    # Limita a área de impressão à tabela real (até a coluna F, alinhada com a logo
    # BINFAE-GL) e desliga a impressão das linhas de grade, já que a planilha tem
    # preenchimento branco em células vazias até a coluna R que expandia a página
    # impressa muito além da tabela.
    ws.print_area = f'A1:F{print_area_last_row}'
    ws.print_options.gridLines = False
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1

    # Sobe o documento na folha: reduz a margem superior (a logo fica rente à borda
    # de impressão) e soma a diferença na margem inferior, mantendo a altura
    # impressa igual e deixando a sobra de espaço no fim da página.
    ws.page_margins.top = 0.15
    ws.page_margins.bottom = 1.35

    # Move a logo do brasão um pouco para a direita para centralizá-la com a
    # frase "MINISTÉRIO DA DEFESA".
    if ws._images:
        ws._images[0].anchor._from.colOff += image_offset_emu

    wb.save(xlsx_path)

    try:
        subprocess.run(
            ['soffice', '--headless', '--norestore', '--convert-to', 'pdf', '--outdir', tmp_dir, xlsx_path],
            check=True, timeout=60, capture_output=True,
        )
        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    nome_arquivo = f"ficha_desimpedimento_{militar.nome_guerra or militar.pk}.pdf".replace(' ', '_')
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename="{nome_arquivo}"'
    return response

@s1_required
def baixa(request):
    rank_order = Case(
        When(posto='CL', then=Value(0)), When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
        When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)),When(posto='ASP', then=Value (6)), When(posto='SO', then=Value(7)),
        When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)), When(posto='3S', then=Value(10)),
        When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)), When(posto='S2', then=Value(13)),When(posto='REC', then=Value(14)),
        default=Value(99), output_field=IntegerField(),
    )
    militares = Efetivo.objects.exclude(situacao__iexact='Baixado').annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_completo')
    
    if request.method == 'POST':
        militar_id = request.POST.get('militar_baixa')
        data_baixa = request.POST.get('data_baixa')
        motivo_baixa = request.POST.get('motivo_baixa')
        
        if militar_id:
            try:
                militar = Efetivo.objects.get(id=militar_id) # type: ignore
                militar.situacao = 'Baixado' # Ao invés de deletar, apenas altera a situação
                militar.observacao = motivo_baixa # Salva o motivo da baixa
                militar.motivo_desligamento = motivo_baixa
                if data_baixa:
                    try:
                        militar.data_desligamento = datetime.strptime(data_baixa, '%Y-%m-%d').date()
                    except ValueError:
                        pass
                militar.save()
                messages.success(request, f'Militar {militar.posto} {militar.nome_guerra} desligado do efetivo com sucesso.')
            except Efetivo.DoesNotExist:
                messages.error(request, 'Erro: Militar não encontrado.')
                
        return redirect('Secao_pessoal:baixa')

    context = {
        'militares': militares
    }
    return render(request, 'Secao_pessoal/baixa.html', context)

@s1_required
def indisponiveis(request):
    rank_order = Case(
        When(posto='CL', then=Value(0)), When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
        When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)),When(posto='ASP', then=Value (6)), When(posto='SO', then=Value(7)),
        When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)), When(posto='3S', then=Value(10)),
        When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)), When(posto='S2', then=Value(13)),When(posto='REC', then=Value(14)),
        default=Value(99), output_field=IntegerField(),
    )
    qs = Efetivo.objects.exclude(
        Q(situacao__iexact='Ativo') | Q(situacao__iexact='Ativa') | Q(situacao__exact='') | Q(situacao__isnull=True)
    ).annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_completo')
    paginator = Paginator(qs, 20)
    page_obj = paginator.get_page(request.GET.get('page'))
    context = {
        'militares': page_obj,
        'page_obj': page_obj,
        'is_paginated': page_obj.has_other_pages(),
    }
    return render(request, 'Secao_pessoal/indisponiveis.html', context)

@s1_required
def reintegrar_militar(request, pk):
    if request.method == 'POST':
        try:
            militar = get_object_or_404(Efetivo.all_objects, id=pk)
            
            # Se o militar tem dados de inspeção, arquiva-os no histórico e limpa os campos
            if militar.inspsau_finalidade or militar.documento_inspsau:
                HistoricoInspsau.objects.create(
                    militar=militar,
                    finalidade=militar.inspsau_finalidade,
                    validade=militar.inspsau_validade,
                    documento=militar.documento_inspsau,
                    parecer=militar.inspsau_parecer
                )
                militar.documento_inspsau = None
                militar.inspsau_finalidade = None
                militar.inspsau_validade = None
                militar.inspsau_parecer = None

            militar.situacao = 'Ativo' # Retorna a situação para Ativo
            militar.observacao = '' # Limpa o histórico de período/motivo
            militar.save()
            
            messages.success(request, f'Militar {militar.posto} {militar.nome_guerra} reintegrado ao efetivo ativo com sucesso.')
        except Efetivo.DoesNotExist:
            messages.error(request, 'Erro: Militar não encontrado.')
            
    return redirect(request.META.get('HTTP_REFERER', 'Secao_pessoal:militar_list'))

@method_decorator(s1_required, name='dispatch')
class LotacaoPessoalListView(ListView):
    model = LotacaoPessoal
    template_name = 'Secao_pessoal/lotacao_list.html'
    context_object_name = 'lotacoes'
    paginate_by = 50

@method_decorator(s1_required, name='dispatch')
class LotacaoPessoalCreateView(CreateView):
    model = LotacaoPessoal
    form_class = LotacaoPessoalForm
    template_name = 'Secao_pessoal/lotacao_form.html'
    success_url = reverse_lazy('Secao_pessoal:lotacao_list')

@method_decorator(s1_required, name='dispatch')
class LotacaoPessoalUpdateView(UpdateView):
    model = LotacaoPessoal
    form_class = LotacaoPessoalForm
    template_name = 'Secao_pessoal/lotacao_form.html'
    success_url = reverse_lazy('Secao_pessoal:lotacao_list')

@method_decorator(s1_required, name='dispatch')
class LotacaoPessoalDeleteView(DeleteView):
    model = LotacaoPessoal
    template_name = 'Secao_pessoal/lotacao_confirm_delete.html'
    success_url = reverse_lazy('Secao_pessoal:lotacao_list')

@s1_required
def relatorio_tlp(request):
    existente_qs = Efetivo.objects.exclude(situacao__iexact='Baixado') \
        .values('posto', 'quad', 'especializacao', 'om').annotate(existente=Count('id'))
    existente_map = {
        (r['posto'], r['quad'], r['especializacao'], r['om']): r['existente']
        for r in existente_qs
    }

    linhas = []
    vistos = set()
    for lot in LotacaoPessoal.objects.all():
        chave = (lot.posto, lot.quad, lot.especializacao, lot.om)
        existente = existente_map.get(chave, 0)
        linhas.append({
            'posto': lot.posto, 'quad': lot.quad, 'especializacao': lot.especializacao,
            'om': lot.om, 'vagas_previstas': lot.vagas_previstas,
            'existente': existente, 'saldo': existente - lot.vagas_previstas,
        })
        vistos.add(chave)

    for chave, existente in existente_map.items():
        if chave not in vistos and any(chave):
            linhas.append({
                'posto': chave[0], 'quad': chave[1], 'especializacao': chave[2],
                'om': chave[3], 'vagas_previstas': 0, 'existente': existente, 'saldo': existente,
            })

    return render(request, 'Secao_pessoal/relatorio_tlp.html', {'linhas': linhas})

@s1_required
def gerenciar_opcoes(request):
    # Dicionário que mapeia o 'tipo' do formulário para a Model correspondente
    MAPA_OPCOES = {
        'posto': Posto,
        'quad': Quad,
        'especializacao': Especializacao,
        'om': OM,
        'setor': Setor,
        'subsetor': Subsetor,
    }

    if request.method == 'POST':
        action = request.POST.get('action')
        tipo_opcao = request.POST.get('tipo_opcao') # Pega o tipo de opção para o redirect
        
        # Ação para ADICIONAR um novo item
        if action == 'add':
            nome = request.POST.get('nome', '').strip()
            
            if tipo_opcao in MAPA_OPCOES and nome:
                model = MAPA_OPCOES[tipo_opcao]
                # 'get_or_create' evita duplicatas
                obj, created = model.objects.get_or_create(nome=nome)
                if created:
                    messages.success(request, f'Opção "{nome}" adicionada com sucesso.')
                else:
                    messages.warning(request, f'Opção "{nome}" já existia.')
            else:
                messages.error(request, 'Erro ao adicionar a opção. Verifique os dados.')

        # Ação para DELETAR um item
        elif action == 'delete':
            item_id = request.POST.get('item_id')

            if tipo_opcao in MAPA_OPCOES and item_id:
                model = MAPA_OPCOES[tipo_opcao]
                try:
                    item = model.objects.get(id=item_id)
                    item_nome = item.nome
                    item.delete()
                    messages.success(request, f'Opção "{item_nome}" removida com sucesso.')
                except model.DoesNotExist:
                    messages.error(request, 'Erro: A opção que você tentou remover não existe.')

        # Redireciona para a mesma página, mantendo o tipo selecionado na URL
        redirect_url = reverse_lazy('Secao_pessoal:gerenciar_opcoes')
        if tipo_opcao:
            return redirect(f'{redirect_url}?tipo={tipo_opcao}')
        return redirect(redirect_url)

    # Contexto para o método GET (carregamento da página)
    context = {
        'postos': Posto.objects.all(),
        'quads': Quad.objects.all(),
        'especializacoes': Especializacao.objects.all(),
        'oms': OM.objects.all(),
        'setores': Setor.objects.all(),
        'subsetores': Subsetor.objects.all(),
        'selected_tipo': request.GET.get('tipo', 'posto'), # Pega da URL ou define um padrão
    }
    return render(request, 'Secao_pessoal/gerenciar_opcoes.html', context)

@s1_required
def exportar_efetivo(request):
    if request.method == 'POST':
        # Filtro avançado (vem do botão Exportar na lista)
        q        = request.POST.get('q', '').strip()
        posto_f  = [v.strip() for v in request.POST.get('posto_f', '').split(',') if v.strip()]
        esp_f    = [v.strip() for v in request.POST.get('esp_f', '').split(',') if v.strip()]
        sit_f    = [v.strip() for v in request.POST.get('sit_f', '').split(',') if v.strip()]
        obs_f    = [v.strip() for v in request.POST.get('obs_f', '').split(',') if v.strip()]
        setor_f  = [v.strip() for v in request.POST.get('setor_f', '').split(',') if v.strip()]
        usa_filtro_avancado = any([q, posto_f, esp_f, sit_f, obs_f, setor_f])

        # Cria o Workbook e a planilha ativa
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Efetivo Exportado"

        # Estilos
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
        alignment_center = Alignment(horizontal="center", vertical="center")
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

        # Define os cabeçalhos solicitados
        headers = ["POSTO", "NOME DE GUERRA", "NOME COMPLETO", "SARAM"]
        ws.append(headers)

        # Aplica estilo ao cabeçalho
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = alignment_center

        # Base da query com ordenação hierárquica (Mesma lógica da MilitarListView)
        rank_order = Case(
            When(posto='CL', then=Value(0)), When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
            When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)),When(posto='ASP', then=Value (6)), When(posto='SO', then=Value(7)),
            When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)), When(posto='3S', then=Value(10)),
            When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)), When(posto='S2', then=Value(13)),When(posto='REC', then=Value(14)),
            default=Value(99), output_field=IntegerField(),
        )
        queryset = Efetivo.objects.exclude(situacao__iexact='Baixado').annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_guerra')

        # Aplica os filtros
        if usa_filtro_avancado:
            if q:
                queryset = queryset.filter(
                    Q(nome_completo__icontains=q) | Q(nome_guerra__icontains=q) | Q(posto__icontains=q)
                )
            if posto_f:
                queryset = queryset.filter(posto__in=posto_f)
            if esp_f:
                eq = Q()
                for v in esp_f:
                    if v == '__vazio__': eq |= Q(especializacao='') | Q(especializacao__isnull=True)
                    else: eq |= Q(especializacao=v)
                queryset = queryset.filter(eq)
            if sit_f:
                sq = Q()
                for v in sit_f:
                    if v == '__vazio__': sq |= Q(situacao='') | Q(situacao__isnull=True)
                    else: sq |= Q(situacao=v)
                queryset = queryset.filter(sq)
            if obs_f:
                oq = Q()
                for v in obs_f:
                    if v == '__vazio__': oq |= Q(observacao='') | Q(observacao__isnull=True)
                    else: oq |= Q(observacao=v)
                queryset = queryset.filter(oq)
            if setor_f:
                stq = Q()
                for v in setor_f:
                    if v == '__vazio__': stq |= Q(setor='') | Q(setor__isnull=True)
                    else: stq |= Q(setor=v)
                queryset = queryset.filter(stq)
        else:
            filtro = request.POST.get('filtro')
            if filtro == 'todos':
                pass
            elif filtro == 'oficiais':
                queryset = queryset.filter(oficial=True)
            elif filtro == 'pracas':
                queryset = queryset.filter(oficial=False)
            elif filtro:
                queryset = queryset.filter(posto=filtro)

        # Preenche os dados
        for militar in queryset:
            ws.append([
                str(militar.posto),
                militar.nome_guerra,
                militar.nome_completo,
                militar.saram if militar.saram else ""
            ])

        # Aplica bordas e alinhamento em todas as células de dados
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=4):
            for cell in row:
                cell.border = thin_border
                cell.alignment = Alignment(vertical="center")
        
        # Ajuste automático de largura das colunas
        for column_cells in ws.columns:
            length = max(len(str(cell.value)) if cell.value else 0 for cell in column_cells)
            ws.column_dimensions[get_column_letter(column_cells[0].column)].width = length + 2

        # Configura a resposta HTTP para download do arquivo
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=efetivo_exportado.xlsx'
        wb.save(response)
        return response

    # GET: Renderiza a página de seleção
    postos_db = Efetivo.objects.values_list('posto', flat=True).distinct()
    
    hierarquia = {
        'CL': 0, 'TC': 1, 'MJ': 2, 'CP': 3,
        '1T': 4, '2T': 5, 'ASP': 6, 'SO': 7,
        '1S': 8, '2S': 9, '3S': 10,
        'CB': 11, 'S1': 12, 'S2': 13, 'REC': 14
    }
    
    postos_existentes = sorted(postos_db, key=lambda x: hierarquia.get(x, 99))
    
    return render(request, 'Secao_pessoal/exportar_excel.html', {
        'postos': postos_existentes
    })

@method_decorator(s1_required, name='dispatch')
class HistoricoInspsauListView(ListView):
    model = HistoricoInspsau
    template_name = 'Secao_pessoal/historico_inspsau.html'
    context_object_name = 'historico_list'
    paginate_by = 20

    def get_queryset(self):
        query = self.request.GET.get('q')
        qs = super().get_queryset().select_related('militar').order_by('-data_registro')

        if query:
            qs = qs.filter(
                Q(militar__nome_completo__icontains=query) |
                Q(militar__nome_guerra__icontains=query) |
                Q(militar__posto__icontains=query) |
                Q(militar__saram__icontains=query) |
                Q(finalidade__icontains=query) |
                Q(parecer__icontains=query)
            )
        return qs

@method_decorator(s1_required, name='dispatch')
class PrestacaoServicoListView(ListView):
    model = Efetivo
    template_name = 'Secao_pessoal/prestacao_servico_list.html'
    context_object_name = 'militares'
    paginate_by = 20

    def get_queryset(self):
        query = self.request.GET.get('q')
        rank_order = Case(
            When(posto='CL', then=Value(0)), When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
            When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)),When(posto='ASP', then=Value (6)), When(posto='SO', then=Value(7)),
            When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)), When(posto='3S', then=Value(10)),
            When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)), When(posto='S2', then=Value(13)),When(posto='REC', then=Value(14)),
            default=Value(99), output_field=IntegerField(),
        ) 
        qs = super().get_queryset().filter(unidade_prestacao_servico__isnull=False).exclude(unidade_prestacao_servico='').annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_completo')
        if query:
            q_objects = Q(nome_completo__icontains=query) | \
                        Q(nome_guerra__icontains=query) | \
                        Q(saram__icontains=query) | \
                        Q(unidade_prestacao_servico__icontains=query) | \
                        Q(portaria_prestacao__icontains=query) | \
                        Q(sigad_prestacao__icontains=query)
            qs = qs.filter(q_objects)
        return qs

    def get_template_names(self):
        if self.request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return ['Secao_pessoal/psv_list_partial.html']
        return ['Secao_pessoal/prestacao_servico_list.html']

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['current_query'] = self.request.GET.get('q', '')
        return context

@s1_required
@require_POST
def historico_inspsau_delete(request, pk):
    if not request.user.is_superuser:
        messages.error(request, "Você não tem permissão para executar esta ação.")
        return redirect('Secao_pessoal:historico_inspsau_list')
    
    historico_item = get_object_or_404(HistoricoInspsau, pk=pk)
    try:
        militar_nome = historico_item.militar.nome_guerra
        historico_item.delete()
        messages.success(request, f"Registro do histórico de {militar_nome} excluído com sucesso.")
    except Exception as e:
        messages.error(request, f"Ocorreu um erro ao excluir o registro: {e}")
        
    return redirect('Secao_pessoal:historico_inspsau_list')
_POSTO_MAP = {
    'ASPIRANTE': 'ASP', 'ASP': 'ASP',
    '2T': '2T', '2º TENENTE': '2T', '2 TENENTE': '2T', 'SEGUNDO TENENTE': '2T',
    '1T': '1T', '1º TENENTE': '1T', '1 TENENTE': '1T', 'PRIMEIRO TENENTE': '1T',
    'CAP': 'CP', 'CAPITAO': 'CP', 'CAPITÃO': 'CP', 'CP': 'CP',
    'MAJ': 'MJ', 'MAJOR': 'MJ', 'MJ': 'MJ',
    'TC': 'TC', 'TEN CEL': 'TC', 'TENENTE CORONEL': 'TC', 'TEN-CEL': 'TC',
    'CEL': 'CL', 'CORONEL': 'CL', 'CL': 'CL',
    'SO': 'SO', 'SARGENTO': 'SO',
    '1S': '1S', '1º SARGENTO': '1S', 'PRIMEIRO SARGENTO': '1S',
    '2S': '2S', '2º SARGENTO': '2S', 'SEGUNDO SARGENTO': '2S',
    '3S': '3S', '3º SARGENTO': '3S', 'TERCEIRO SARGENTO': '3S',
    'CB': 'CB', 'CABO': 'CB',
    'SD': 'S1', 'S1': 'S1', 'S2': 'S2', 'REC': 'REC', 'SOLDADO': 'S1',
}
_POSTO_TEN = {'TEN', 'TENENTE', '1TEN', '2TEN', 'TEN.', '1T/2T', 'TENENTE.', 'TEN.1', 'TEN.2'}

# Prefixos reconhecidos no nome de guerra (ordem: mais longos primeiro para evitar partial matches)
_POSTOS_PREFIXO = sorted(
    list(_POSTO_MAP.keys()) + list(_POSTO_TEN),
    key=lambda x: -len(x)
)


def _strip_posto_de_nome(nome_guerra):
    """Remove prefix de posto do nome_guerra. Retorna (nome_limpo, posto_detectado_ou_None)."""
    if not nome_guerra:
        return nome_guerra, None
    ng_upper = nome_guerra.upper().strip()
    for p in _POSTOS_PREFIXO:
        p_upper = p.upper()
        if ng_upper.startswith(p_upper + ' ') or ng_upper == p_upper:
            stripped = nome_guerra[len(p):].strip()
            posto_norm = _POSTO_MAP.get(normalize_name(p.replace('.', '').replace('-', '')))
            return (stripped.upper() if stripped else nome_guerra.upper()), posto_norm
    return nome_guerra.upper(), None


def _normalizar_posto(valor):
    """Retorna (posto_normalizado, é_tenente_ambiguo)."""
    v = normalize_name(valor.strip())
    # remove pontos e hifens para comparação
    v_clean = v.replace('.', '').replace('-', '').strip()
    if v_clean in _POSTO_TEN or v in _POSTO_TEN:
        return None, True  # TEN ambíguo
    return _POSTO_MAP.get(v_clean) or _POSTO_MAP.get(v) or valor.strip(), False


def _parse_date(val):
    """Tenta converter string para date, retorna None em caso de falha."""
    from datetime import datetime as _dt
    if not val or str(val).strip() in ('', 'nan', 'NaT', 'None'):
        return None
    for fmt in ('%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y', '%m/%d/%Y'):
        try:
            return _dt.strptime(str(val).strip(), fmt).date()
        except ValueError:
            pass
    return None


def _extrair_ano_turma(val):
    """Normaliza o campo TURMA: se for data (ex: '2001-12-15 00:00:00'), extrai só o ano."""
    v = str(val).strip()
    if not v or v in ('', 'nan', 'NaT', 'None'):
        return ''
    if len(v) >= 4 and v[:4].isdigit() and (len(v) == 4 or not v[4].isdigit()):
        return v[:4]
    return v


# ── Adicionar PSV ───────────────────────────────────────────────────────────
@s1_required
def adicionar_psv(request):
    rank_order = Case(
        When(posto='CL', then=Value(0)), When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
        When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)), When(posto='ASP', then=Value(6)), When(posto='SO', then=Value(7)),
        When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)), When(posto='3S', then=Value(10)),
        When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)), When(posto='S2', then=Value(13)), When(posto='REC', then=Value(14)),
        default=Value(99), output_field=IntegerField(),
    )
    militares = Efetivo.objects.exclude(situacao__iexact='Baixado').annotate(rank_order=rank_order).order_by('rank_order', 'turma', 'nome_completo')

    if request.method == 'POST':
        militar_id = request.POST.get('militar_psv')
        if not militar_id:
            messages.error(request, 'Selecione o militar.')
            return redirect('Secao_pessoal:adicionar_psv')
        try:
            militar = Efetivo.objects.get(id=militar_id)
        except Efetivo.DoesNotExist:
            messages.error(request, 'Militar não encontrado.')
            return redirect('Secao_pessoal:adicionar_psv')

        militar.unidade_prestacao_servico = request.POST.get('unidade', '').strip() or None
        militar.sigad_prestacao           = request.POST.get('sigad', '').strip() or None
        militar.portaria_prestacao        = request.POST.get('portaria', '').strip() or None
        militar.boletim_prestacao         = request.POST.get('boletim', '').strip() or None
        militar.data_inicio_prestacao     = _parse_date(request.POST.get('data_inicio', ''))
        militar.data_vencimento_prestacao = _parse_date(request.POST.get('data_vencimento', ''))
        militar.data_portaria_prestacao   = _parse_date(request.POST.get('data_portaria', ''))
        militar.data_boletim_prestacao    = _parse_date(request.POST.get('data_boletim', ''))
        _oms_gsd = {'BINFAE-GL', 'GSD-GL', 'BINFAE GL', 'GSD GL'}
        om_origem = (militar.om or '').strip().upper()
        militar.situacao = 'PSV GSD-GL' if om_origem in _oms_gsd else 'PSV'
        militar.save()
        messages.success(request, f'Prestação de serviço de {militar.posto} {militar.nome_guerra} registrada com sucesso.')
        return redirect('Secao_pessoal:prestacao_servico')

    return render(request, 'Secao_pessoal/adicionar_psv.html', {'militares': militares})


# ── Retornar PSV → OM Origem ────────────────────────────────────────────────
@s1_required
@require_POST
def retornar_psv(request, pk):
    militar = get_object_or_404(Efetivo, pk=pk)
    militar.situacao = 'ATIVA'
    militar.unidade_prestacao_servico = None
    militar.data_inicio_prestacao     = None
    militar.data_vencimento_prestacao = None
    militar.sigad_prestacao           = None
    militar.portaria_prestacao        = None
    militar.boletim_prestacao         = None
    militar.data_portaria_prestacao   = None
    militar.data_boletim_prestacao    = None
    militar.save()
    messages.success(request, f'{militar.posto} {militar.nome_guerra} retornou à OM de origem.')
    return redirect('Secao_pessoal:prestacao_servico')


# ── Exportar PSV ────────────────────────────────────────────────────────────
@s1_required
def exportar_psv(request):
    import io
    qs = Efetivo.objects.filter(
        Q(unidade_prestacao_servico__isnull=False) & ~Q(unidade_prestacao_servico='')
    ).order_by('posto', 'nome_completo')

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'PSV'
    headers = ['Posto', 'Espec', 'Quad', 'Nome Completo', 'Nome de Guerra', 'SARAM', 'Turma', 'OM', 'Início', 'Término', 'Sigad', 'Boletim', 'Observações']
    ws.append(headers)
    for m in qs:
        ws.append([
            m.posto, m.especializacao, m.quad, m.nome_completo, m.nome_guerra,
            m.saram, m.turma,
            m.unidade_prestacao_servico or '',
            m.data_inicio_prestacao.strftime('%d/%m/%Y') if m.data_inicio_prestacao else '',
            m.data_vencimento_prestacao.strftime('%d/%m/%Y') if m.data_vencimento_prestacao else '',
            m.sigad_prestacao or '',
            m.boletim_prestacao or '',
            m.observacao or '',
        ])
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    resp = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    resp['Content-Disposition'] = 'attachment; filename="PSV_GSD_GL.xlsx"'
    return resp


# ── Importar PSV ────────────────────────────────────────────────────────────
@s1_required
def importar_psv(request):
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        sincronizar = request.POST.get('sincronizar') == '1'
        try:
            df = pd.read_excel(excel_file, dtype=str)
            df.columns = df.columns.str.strip()
            df.fillna('', inplace=True)

            # Normaliza nomes de colunas
            COL_MAP = {
                'posto': ['POSTO', 'PST', 'PST.'],
                'espec': ['ESPEC', 'ESP', 'ESPECIALIDADE', 'ESPECIALIZACAO'],
                'quad': ['QUAD', 'QUADRO'],
                'nome_completo': ['NOME COMPLETO', 'NOME'],
                'nome_guerra': ['NOME DE GUERRA', 'NOME GUERRA', 'GUERRA'],
                'saram': ['SARAM'],
                'turma': ['TURMA'],
                'om': ['OM', 'UNIDADE'],
                'inicio': ['INICIO', 'INÍCIO', 'DATA INICIO', 'DATA INÍCIO'],
                'termino': ['TERMINO', 'TÉRMINO', 'DATA TERMINO', 'DATA TÉRMINO', 'VENCIMENTO'],
                'sigad': ['SIGAD'],
                'boletim': ['BOLETIM'],
                'observacoes': ['OBSERVACOES', 'OBSERVAÇÕES', 'OBS'],
            }
            def find_col(df, candidates):
                for c in df.columns:
                    normalized = normalize_name(c.strip().replace('.', '').replace('/', ''))
                    for cand in candidates:
                        if normalized == normalize_name(cand):
                            return c
                return None

            col = {k: find_col(df, v) for k, v in COL_MAP.items()}

            rows = []
            ambiguous = []

            for _, row in df.iterrows():
                saram_raw = str(row.get(col['saram'], '')).strip() if col['saram'] else ''
                nome_completo = str(row.get(col['nome_completo'], '')).strip() if col['nome_completo'] else ''
                if not saram_raw and not nome_completo:
                    continue

                posto_raw = str(row.get(col['posto'], '')).strip() if col['posto'] else ''
                posto_norm, is_ten = _normalizar_posto(posto_raw) if posto_raw else ('', False)

                r = {
                    'saram': saram_raw,
                    'nome_completo': nome_completo,
                    'nome_guerra': str(row.get(col['nome_guerra'], '')).strip() if col['nome_guerra'] else '',
                    'posto_raw': posto_raw,
                    'posto_norm': posto_norm,
                    'espec': str(row.get(col['espec'], '')).strip() if col['espec'] else '',
                    'quad': str(row.get(col['quad'], '')).strip() if col['quad'] else '',
                    'turma': str(row.get(col['turma'], '')).strip() if col['turma'] else '',
                    'om': str(row.get(col['om'], '')).strip() if col['om'] else '',
                    'inicio': str(row.get(col['inicio'], '')).strip() if col['inicio'] else '',
                    'termino': str(row.get(col['termino'], '')).strip() if col['termino'] else '',
                    'sigad': str(row.get(col['sigad'], '')).strip() if col['sigad'] else '',
                    'boletim': str(row.get(col['boletim'], '')).strip() if col['boletim'] else '',
                    'observacoes': str(row.get(col['observacoes'], '')).strip() if col['observacoes'] else '',
                    'is_ten': is_ten,
                }
                if is_ten:
                    ambiguous.append(r)
                else:
                    rows.append(r)

            if ambiguous:
                import json as _json
                request.session['importar_psv_rows'] = rows
                request.session['importar_psv_ambiguous'] = ambiguous
                request.session['importar_psv_sincronizar'] = sincronizar
                return redirect('Secao_pessoal:esclarecer_importacao_psv')

            return _completar_importacao_psv(request, rows, sincronizar)

        except Exception as e:
            messages.error(request, f'Erro na importação: {e}')
            return redirect('Secao_pessoal:prestacao_servico')

    return render(request, 'Secao_pessoal/importar_psv.html')


@s1_required
def esclarecer_importacao_psv(request):
    ambiguous = request.session.get('importar_psv_ambiguous', [])
    if not ambiguous:
        return redirect('Secao_pessoal:prestacao_servico')

    if request.method == 'POST':
        rows = request.session.get('importar_psv_rows', [])
        sincronizar = request.session.get('importar_psv_sincronizar', False)
        for i, r in enumerate(ambiguous):
            posto_escolhido = request.POST.get(f'posto_{i}', '')
            r['posto_norm'] = posto_escolhido
            r['is_ten'] = False
            rows.append(r)
        del request.session['importar_psv_ambiguous']
        del request.session['importar_psv_rows']
        request.session.pop('importar_psv_sincronizar', None)
        return _completar_importacao_psv(request, rows, sincronizar)

    return render(request, 'Secao_pessoal/esclarecer_importacao.html', {'ambiguous': ambiguous, 'fonte': 'psv'})


def _completar_importacao_psv(request, rows, sincronizar):
    from datetime import date as _date
    atualizados = 0
    nao_encontrados = []
    pks_na_planilha = set()

    for r in rows:
        saram_db = None
        if r['saram']:
            try:
                saram_db = int(float(r['saram']))
            except (ValueError, TypeError):
                pass

        efetivo = None
        if saram_db:
            efetivo = Efetivo.objects.filter(saram=saram_db).first()
        if not efetivo and r['nome_completo']:
            chave = normalize_name(r['nome_completo'])
            efetivo = next(
                (e for e in Efetivo.objects.all() if normalize_name(e.nome_completo) == chave),
                None
            )
        if not efetivo:
            nao_encontrados.append(r.get('nome_completo') or r.get('saram') or '?')
            continue

        efetivo.unidade_prestacao_servico = r['om'] or None
        efetivo.data_inicio_prestacao     = _parse_date(r['inicio'])
        efetivo.data_vencimento_prestacao = _parse_date(r['termino'])
        efetivo.sigad_prestacao           = r['sigad'] or None
        efetivo.boletim_prestacao         = r['boletim'] or None
        if r['posto_norm']:
            efetivo.posto = r['posto_norm']
        if r['espec']:
            efetivo.especializacao = r['espec']
        if r['turma']:
            efetivo.turma = r['turma']
        _oms_gsd = {'BINFAE-GL', 'GSD-GL', 'BINFAE GL', 'GSD GL'}
        om_origem = (efetivo.om or '').strip().upper()
        efetivo.situacao = 'PSV GSD-GL' if om_origem in _oms_gsd else 'PSV'
        efetivo.save()
        pks_na_planilha.add(efetivo.pk)
        atualizados += 1

    if sincronizar and pks_na_planilha:
        Efetivo.objects.filter(
            Q(unidade_prestacao_servico__isnull=False) & ~Q(unidade_prestacao_servico='')
        ).exclude(pk__in=pks_na_planilha).update(
            unidade_prestacao_servico=None, data_inicio_prestacao=None,
            data_vencimento_prestacao=None, sigad_prestacao=None,
            boletim_prestacao=None,
        )

    msg = f'{atualizados} militar(es) atualizado(s).'
    if nao_encontrados:
        msg += f' Não encontrado(s) no sistema: {", ".join(nao_encontrados[:5])}{"..." if len(nao_encontrados) > 5 else ""}.'
    messages.success(request, msg)
    return redirect('Secao_pessoal:prestacao_servico')


@s1_required
@require_GET
def api_search_militares(request):
    """
    Endpoint de API para a busca manual de militares no modal de INSPSAU.
    """
    query = request.GET.get('q', '')
    if not query or len(query) < 2:
        return JsonResponse([], safe=False)

    militares = Efetivo.objects.filter(
        Q(nome_completo__icontains=query) |
        Q(nome_guerra__icontains=query) |
        Q(saram__icontains=query)
    ).order_by('posto', 'nome_guerra')[:15]

    data = [{'id': m.id, 'posto': m.posto, 'nome_guerra': m.nome_guerra, 'nome_completo': m.nome_completo} for m in militares]

    return JsonResponse(data, safe=False)


@s1_required
@require_GET
def download_modelo_efetivo(request):
    from django.contrib.staticfiles import finders
    path = finders.find('Secao_pessoal/templates_pdf/modelo_importacao_efetivo.xlsx')
    if not path:
        from django.http import Http404
        raise Http404
    with open(path, 'rb') as f:
        content = f.read()
    resp = HttpResponse(content, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    resp['Content-Disposition'] = 'attachment; filename="MODELO ATUALIZACAO DE EFETIVO.xlsx"'
    return resp


@s1_required
def download_modelo_desligamento(request):
    """Gera e serve o modelo Excel para importação de desligamentos."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Desligamentos"
    headers = ['SARAM', 'NOME COMPLETO', 'PST.', 'QUAD.', 'ESP.', 'NOME DE GUERRA', 'SITUAÇÃO']
    hf = Font(bold=True, color='FFFFFF')
    hfill = PatternFill(start_color='4F81BD', end_color='4F81BD', fill_type='solid')
    ws.append(headers)
    for cell in ws[1]:
        cell.font = hf
        cell.fill = hfill
        cell.alignment = Alignment(horizontal='center')
    ws.append(['', 'EXEMPLO DA SILVA', 'S2', 'QPPM', 'QAE', 'EXEMPLO', 'DESLIGADO'])
    for col in ws.columns:
        width = max(len(str(c.value)) if c.value else 0 for c in col) + 4
        ws.column_dimensions[col[0].column_letter].width = width
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=modelo_desligamento.xlsx'
    wb.save(response)
    return response


@s1_required
def exportar_desligados(request):
    if request.method != 'POST':
        return redirect('Secao_pessoal:desligados_list')
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Desligados"
    hf = Font(bold=True, color='FFFFFF')
    hfill = PatternFill(start_color='4F81BD', end_color='4F81BD', fill_type='solid')
    ac = Alignment(horizontal='center', vertical='center')
    thin = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    headers = ['POSTO', 'NOME DE GUERRA', 'NOME COMPLETO', 'SARAM', 'ESPECIALIZAÇÃO', 'SETOR', 'OBSERVAÇÃO']
    ws.append(headers)
    for cell in ws[1]:
        cell.font = hf; cell.fill = hfill; cell.alignment = ac
    rank_order = Case(
        When(posto='TC', then=Value(1)), When(posto='MJ', then=Value(2)), When(posto='CP', then=Value(3)),
        When(posto='1T', then=Value(4)), When(posto='2T', then=Value(5)), When(posto='ASP', then=Value(6)),
        When(posto='SO', then=Value(7)), When(posto='1S', then=Value(8)), When(posto='2S', then=Value(9)),
        When(posto='3S', then=Value(10)), When(posto='CB', then=Value(11)), When(posto='S1', then=Value(12)),
        When(posto='S2', then=Value(13)), When(posto='REC', then=Value(14)),
        default=Value(99), output_field=IntegerField(),
    )
    qs = Efetivo.all_objects.filter(situacao__iexact='Baixado').annotate(rank_order=rank_order).order_by('rank_order', 'nome_guerra')
    q = request.POST.get('q', '').strip()
    if q:
        qs = qs.filter(Q(nome_completo__icontains=q) | Q(nome_guerra__icontains=q))
    for mil in qs:
        ws.append([mil.posto, mil.nome_guerra, mil.nome_completo, mil.saram or '', mil.especializacao or '', mil.setor or '', mil.observacao or ''])
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.border = thin; cell.alignment = Alignment(vertical='center')
    for col in ws.columns:
        width = max(len(str(c.value)) if c.value else 0 for c in col) + 2
        ws.column_dimensions[col[0].column_letter].width = width
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=desligados_exportados.xlsx'
    wb.save(response)
    return response


@s1_required
def importar_desligamento(request):
    if request.method != 'POST' or not request.FILES.get('excel_file'):
        return redirect('Secao_pessoal:desligados_list')
    excel_file = request.FILES['excel_file']
    try:
        df = pd.read_excel(excel_file, dtype=str, header=2)
        df.columns = df.columns.str.strip()
        _cols = {re.sub(r'\s+', ' ', normalize_name(str(c))).strip() for c in df.columns}
        if 'SARAM' not in _cols and 'NOME COMPLETO' not in _cols:
            df = pd.read_excel(excel_file, dtype=str, header=0)
            df.columns = df.columns.str.strip()
        colunas_norm = {re.sub(r'\s+', ' ', re.sub(r'[.\/]', '', normalize_name(str(c)))).strip(): c for c in df.columns}
        rename_map = {}
        for canonico, aliases in [
            ('SARAM', ['SARAM']),
            ('NOME COMPLETO', ['NOME COMPLETO', 'NOME']),
            ('SITUAÇÃO', ['SITUACAO']),
            ('PST.', ['PST', 'POSTO']),
        ]:
            for alias in aliases:
                real = colunas_norm.get(alias)
                if real:
                    rename_map[real] = canonico
                    break
        df.rename(columns=rename_map, inplace=True)
        df.fillna('', inplace=True)

        ignorados = []
        nao_encontrados = []
        ativos_para_confirmar = []
        atualizados = 0

        existentes_por_nome = {}
        for ef in Efetivo.all_objects.all():
            chave = normalize_name(ef.nome_completo).strip()
            if chave:
                existentes_por_nome[chave] = ef

        for _, row in df.iterrows():
            saram_val = str(row.get('SARAM', '')).strip()
            nome_val = str(row.get('NOME COMPLETO', '')).strip()
            sit_val = str(row.get('SITUAÇÃO', '')).strip()
            if not saram_val and not nome_val:
                continue
            if sit_val.upper() != 'DESLIGADO':
                ignorados.append({'nome': nome_val, 'saram': saram_val, 'situacao': sit_val})
                continue
            obj = None
            saram_db = None
            if saram_val:
                try:
                    saram_db = int(float(saram_val))
                    obj = Efetivo.all_objects.filter(saram=saram_db).first()
                except ValueError:
                    pass
            if obj is None and nome_val:
                obj = existentes_por_nome.get(normalize_name(nome_val).strip())
            if obj is None:
                nao_encontrados.append({'nome': nome_val, 'saram': saram_val})
                continue
            sit_atual = (obj.situacao or '').strip().lower()
            if sit_atual in ('ativo', 'ativa', 'ativo/a'):
                ativos_para_confirmar.append({'pk': obj.pk, 'nome': obj.nome_completo, 'posto': obj.posto, 'situacao_atual': obj.situacao})
            else:
                obj.situacao = 'Baixado'
                obj.deleted = False
                obj.save()
                atualizados += 1

        if ignorados or ativos_para_confirmar:
            request.session['desl_ignorados'] = ignorados
            request.session['desl_ativos'] = ativos_para_confirmar
            request.session['desl_atualizados'] = atualizados
            request.session['desl_nao_encontrados'] = nao_encontrados
            return redirect('Secao_pessoal:esclarecer_desligamento')

        msgs = [f'{atualizados} militar(es) desligado(s).']
        if nao_encontrados:
            msgs.append(f'{len(nao_encontrados)} nome(s) não encontrado(s) no efetivo (ignorados).')
        messages.success(request, ' '.join(msgs))
        return redirect('Secao_pessoal:desligados_list')
    except Exception as e:
        messages.error(request, f'Erro na importação: {str(e)}')
        return redirect('Secao_pessoal:desligados_list')


@s1_required
def esclarecer_desligamento(request):
    ignorados = request.session.get('desl_ignorados', [])
    ativos = request.session.get('desl_ativos', [])
    atualizados_base = request.session.get('desl_atualizados', 0)
    nao_encontrados = request.session.get('desl_nao_encontrados', [])

    if not ignorados and not ativos:
        return redirect('Secao_pessoal:desligados_list')

    if request.method == 'POST':
        acao = request.POST.get('acao')
        novos_desligados = 0

        converter_pks_str = request.POST.getlist('converter_ignorado')
        ignorados_converter = []
        for idx_str in converter_pks_str:
            try:
                idx = int(idx_str)
                if 0 <= idx < len(ignorados):
                    ignorados_converter.append(ignorados[idx])
            except (ValueError, IndexError):
                pass
        for item in ignorados_converter:
            saram_val = item.get('saram', '')
            nome_val = item.get('nome', '')
            obj = None
            if saram_val:
                try:
                    obj = Efetivo.all_objects.filter(saram=int(float(saram_val))).first()
                except (ValueError, TypeError):
                    pass
            if obj is None and nome_val:
                chave = normalize_name(nome_val).strip()
                for ef in Efetivo.all_objects.all():
                    if normalize_name(ef.nome_completo).strip() == chave:
                        obj = ef
                        break
            if obj:
                obj.situacao = 'Baixado'
                obj.deleted = False
                obj.save()
                novos_desligados += 1

        if acao == 'desligar_todos':
            pks = [item['pk'] for item in ativos]
        else:
            pks = [int(pk) for pk in request.POST.getlist('desligar_pk') if pk.isdigit()]

        for pk in pks:
            try:
                obj = Efetivo.all_objects.get(pk=pk)
                obj.situacao = 'Baixado'
                obj.deleted = False
                obj.save()
                novos_desligados += 1
            except Efetivo.DoesNotExist:
                pass

        total = atualizados_base + novos_desligados
        for key in ('desl_ignorados', 'desl_ativos', 'desl_atualizados', 'desl_nao_encontrados'):
            request.session.pop(key, None)
        messages.success(request, f'{total} militar(es) desligado(s) no total.')
        return redirect('Secao_pessoal:desligados_list')

    return render(request, 'Secao_pessoal/esclarecer_desligamento.html', {
        'ignorados': list(enumerate(ignorados)),
        'ativos': ativos,
        'nao_encontrados': nao_encontrados,
        'atualizados_base': atualizados_base,
    })


@s1_required
@require_GET
def download_modelo_psv(request):
    from django.contrib.staticfiles import finders
    path = finders.find('Secao_pessoal/templates_pdf/modelo_importacao_psv.xlsx')
    if not path:
        from django.http import Http404
        raise Http404
    with open(path, 'rb') as f:
        content = f.read()
    resp = HttpResponse(content, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    resp['Content-Disposition'] = 'attachment; filename="MODELO ATUALIZACAO DE PSV.xlsx"'
    return resp


@s1_required
@require_GET
def efetivo_filter_options(request):
    """Retorna valores distintos dos campos filtráveis, filtrados pelo contexto da view."""
    view = request.GET.get('view', 'geral')
    if view == 'operacional':
        base_qs = Efetivo.objects.filter(Q(situacao__iexact='Ativo') | Q(situacao__iexact='Ativa') | Q(situacao__iexact='PSV GSD-GL'))
    elif view == 'junta':
        base_qs = Efetivo.objects.filter(situacao__iexact='De Junta')
    elif view == 'desligados':
        base_qs = Efetivo.objects.filter(situacao__iexact='Baixado')
    elif view == 'psv':
        base_qs = Efetivo.objects.filter(
            Q(unidade_prestacao_servico__isnull=False) & ~Q(unidade_prestacao_servico='')
        )
    else:  # geral
        base_qs = Efetivo.objects.exclude(situacao__iexact='Baixado')

    def distinct_sorted(field):
        vals = list(base_qs.values_list(field, flat=True).distinct())
        non_empty = sorted({v for v in vals if v}, key=lambda x: x.lower())
        has_empty = any(v is None or v == '' for v in vals)
        result = non_empty
        if has_empty:
            result = result + ['__vazio__']
        return result

    return JsonResponse({
        'posto':          distinct_sorted('posto'),
        'especializacao': distinct_sorted('especializacao'),
        'situacao':       distinct_sorted('situacao'),
        'observacao':     distinct_sorted('observacao'),
        'setor':          distinct_sorted('setor'),
    })

@s1_required
def importar_fq(request):
    if request.method == 'POST':
        data_str = request.POST.get('data')
        arquivo = request.FILES.get('documento_fq')

        if not data_str or not arquivo:
            messages.error(request, "Por favor, forneça a data e o documento da FQ.")
            return redirect('Secao_pessoal:importar_fq')
            
        try:
            data_chamada = datetime.strptime(data_str, '%Y-%m-%d').date()
        except ValueError:
            data_chamada = date.today()

        try:
            # Salva o arquivo num ficheiro temporário para processar
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(arquivo.name)[1]) as temp_file:
                for chunk in arquivo.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name

            content = ""
            if temp_file_path.lower().endswith('.pdf'):
                doc = fitz.open(temp_file_path)
                for page in doc:
                    page_text = page.get_text()
                    if len(page_text.strip()) > 50:
                        content += page_text + "\n\n"
                    else:
                        try:
                            pix = page.get_pixmap(dpi=300)
                            img = Image.open(io.BytesIO(pix.tobytes("png")))
                            if pytesseract:
                                content += pytesseract.image_to_string(img, lang='por') + "\n\n"
                        except Exception:
                            pass
                doc.close()
            else:
                # Se for imagem
                try:
                    img = Image.open(temp_file_path)
                    if pytesseract:
                        content = pytesseract.image_to_string(img, lang='por')
                except Exception:
                    pass

            os.remove(temp_file_path)

            if not content.strip():
                messages.error(request, "Não foi possível extrair texto legível do documento.")
                return redirect('Secao_pessoal:importar_fq')

            resultado_ia = analisar_fq_documento(content)
            faltosos_marcados = 0
            nao_encontrados = []

            for militar_ia in resultado_ia.faltosos:
                militar = None
                if militar_ia.saram:
                    saram_limpo = re.sub(r'\D', '', str(militar_ia.saram))
                    if saram_limpo:
                        militar = Efetivo.objects.filter(saram=int(saram_limpo)).first()
                
                if not militar and militar_ia.nome_guerra:
                    candidatos = Efetivo.objects.filter(nome_guerra__icontains=militar_ia.nome_guerra)
                    if candidatos.count() == 1:
                        militar = candidatos.first()
                    elif candidatos.count() > 1 and militar_ia.posto:
                        candidatos_posto = candidatos.filter(posto__icontains=militar_ia.posto)
                        if candidatos_posto.count() == 1:
                            militar = candidatos_posto.first()
                        else:
                            militar = candidatos.first()

                if militar:
                    ChamadaRegistro.objects.update_or_create(
                        data=data_chamada,
                        militar=militar,
                        defaults={'status': 'F'}
                    )
                    faltosos_marcados += 1
                else:
                    nao_encontrados.append(f"{militar_ia.posto} {militar_ia.nome_guerra}")

            if faltosos_marcados > 0:
                msg = f"Sucesso! {faltosos_marcados} falta(s) registrada(s) na chamada do dia {data_chamada.strftime('%d/%m/%Y')}."
                if nao_encontrados:
                    msg += f" (Militares não encontrados: {', '.join(nao_encontrados)})"
                messages.success(request, msg)
            else:
                if nao_encontrados:
                    messages.warning(request, f"Faltas identificadas pela IA, mas nenhum militar correspondente foi encontrado no sistema: {', '.join(nao_encontrados)}")
                else:
                    messages.info(request, "A Inteligência Artificial analisou o documento e não encontrou nenhum militar com falta.")
                    
        except Exception as e:
            messages.error(request, f"Erro ao analisar o documento FQ: {e}")

        return redirect('Secao_pessoal:importar_fq')

    # Busca as últimas 50 faltas registradas no sistema para exibir no histórico
    historico_faltas = ChamadaRegistro.objects.filter(status='F').select_related('militar').order_by('-data', 'militar__nome_guerra')[:50]
    return render(request, 'Secao_pessoal/importar_fq.html', {'historico_faltas': historico_faltas})


@s1_required
def desercao(request):
    if request.method == 'POST':
        militar_id = request.POST.get('militar_id')
        data_inicio_str = request.POST.get('data_inicio')
        
        if militar_id and data_inicio_str:
            try:
                militar = Efetivo.objects.get(id=militar_id)
                data_inicio = datetime.strptime(data_inicio_str, '%Y-%m-%d').date()
                
                # Criar 8 faltas consecutivas para tornar o militar um desertor
                for i in range(8):
                    data_falta = data_inicio + timedelta(days=i)
                    ChamadaRegistro.objects.update_or_create(
                        data=data_falta,
                        militar=militar,
                        defaults={'status': 'F'}
                    )
                messages.success(request, f"O militar {militar.posto} {militar.nome_guerra} foi marcado como desertor manualmente (8 faltas lançadas a partir de {data_inicio.strftime('%d/%m/%Y')}).")
            except Exception as e:
                messages.error(request, f"Erro ao adicionar desertor: {e}")
        return redirect('Secao_pessoal:desercao')

    militares = Efetivo.objects.exclude(situacao__iexact='Baixado')
    
    desertores = []
    alertas = []

    for militar in militares:
        # Busca as últimas 30 chamadas para otimização
        chamadas = ChamadaRegistro.objects.filter(militar=militar).order_by('-data')[:30]
        
        faltas_consecutivas = 0
        ultima_data_falta = None
        
        for chamada in chamadas:
            if chamada.status == 'F':
                faltas_consecutivas += 1
                if faltas_consecutivas == 1:
                    ultima_data_falta = chamada.data
            else:
                # Quebra a sequência se tiver presença (P), Missão (M), etc.
                break
                
        if faltas_consecutivas >= 8:
            desertores.append({
                'militar': militar,
                'faltas': faltas_consecutivas,
                'data': ultima_data_falta
            })
        elif faltas_consecutivas >= 4:
            alertas.append({
                'militar': militar,
                'faltas': faltas_consecutivas,
                'data': ultima_data_falta
            })

    # Ordenar por mais faltas
    alertas.sort(key=lambda x: x['faltas'], reverse=True)
    desertores.sort(key=lambda x: x['faltas'], reverse=True)

    context = {
        'desertores': desertores,
        'alertas': alertas,
        'militares': militares,
    }
    return render(request, 'Secao_pessoal/desercao.html', context)